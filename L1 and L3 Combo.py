
import pandas as pd
import numpy as np
from itertools import combinations
from snowflake.snowpark.context import get_active_session
import time
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
from scipy.optimize import minimize
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === ROUNDING AND PRECISION HELPERS ===
def round_half_up(value, decimals=2):
    """
    Round using 'round half up' to match PRF official tool.
    Handles floating-point precision issues by converting to Decimal early.
    Python's built-in round() uses banker's rounding (12.675 -> 12.67).
    PRF Tool uses round half up (12.675 -> 12.68).
    """
    if value is None or (isinstance(value, float) and (np.isnan(value) or np.isinf(value))):
        return 0.0

    # Convert to Decimal to avoid floating-point precision issues
    d = Decimal(str(value))

    # Create quantization string
    if decimals == 0:
        quantize_to = Decimal('1')
    else:
        quantize_to = Decimal('0.' + '0' * (decimals - 1) + '1')

    # Round half up and convert back to float
    return float(d.quantize(quantize_to, rounding=ROUND_HALF_UP))

def calculate_protection(county_base_value, coverage_level, productivity_factor, decimals=2):
    """
    Calculate dollar protection with proper precision and rounding.
    Uses Decimal arithmetic to avoid floating-point errors.

    This prevents issues like 16.90 * 0.75 * 1.0 = 12.674999999999999
    which would incorrectly round to 12.67 instead of 12.68.
    """
    # Convert all inputs to Decimal for precise calculation
    cbv = Decimal(str(county_base_value))
    cov = Decimal(str(coverage_level))
    prod = Decimal(str(productivity_factor))

    # Perform calculation
    result = cbv * cov * prod

    # Round using round half up
    if decimals == 0:
        quantize_to = Decimal('1')
    else:
        quantize_to = Decimal('0.' + '0' * (decimals - 1) + '1')

    return float(result.quantize(quantize_to, rounding=ROUND_HALF_UP))

# === GLOBAL CONSTANT ===
# The 11 valid PRF intervals
INTERVAL_ORDER_11 = ['Jan-Feb', 'Feb-Mar', 'Mar-Apr', 'Apr-May', 'May-Jun',
                     'Jun-Jul', 'Jul-Aug', 'Aug-Sep', 'Sep-Oct', 'Oct-Nov', 'Nov-Dec']
MONTH_TO_INTERVAL = {
    1: 'Jan-Feb', 2: 'Feb-Mar', 3: 'Mar-Apr', 4: 'Apr-May',
    5: 'May-Jun', 6: 'Jun-Jul', 7: 'Jul-Aug', 8: 'Aug-Sep',
    9: 'Sep-Oct', 10: 'Oct-Nov', 11: 'Nov-Dec'
}
# --- Reverse mapping for the UI ---
INTERVAL_TO_MONTH_NUM = {name: month for month, name in MONTH_TO_INTERVAL.items()}

# === MARKET VIEW CONSTANTS ===
HISTORICAL_CONTEXT_MAP = {
    'Dry': {'min': float('-inf'), 'max': -0.25},
    'Normal': {'min': -0.25, 'max': 0.25},
    'Wet': {'min': 0.25, 'max': float('inf')}
}

TREND_MAP = {
    'Get Drier': {'min': float('-inf'), 'max': -0.05},
    'Stay Stable': {'min': -0.05, 'max': 0.05},
    'Get Wetter': {'min': 0.05, 'max': float('inf')}
}

# === KING RANCH PRESET CONFIGURATION (CORRECTED COUNTY MAPPINGS) ===
KING_RANCH_PRESET = {
    'grids': [9128, 9129, 8829, 9130, 7929, 8230, 8228, 8229],
    'counties': {
        'Kleberg': [9128, 9129, 8829, 9130],
        'Kenedy': [7929, 8230],
        'Brooks': [8228, 8229]
    },
    'acres': {
        9128: 56662,
        9129: 56662,
        8829: 56662,
        9130: 56662,
        7929: 86774,
        8230: 86774,
        8228: 26386,
        8229: 26386
    },
    'allocations': {
        9128: {'Feb-Mar': 20, 'Apr-May': 20, 'Jun-Jul': 20, 'Aug-Sep': 20, 'Oct-Nov': 20},
        9129: {'Jan-Feb': 17, 'Mar-Apr': 17, 'May-Jun': 17, 'Jul-Aug': 17, 'Sep-Oct': 16, 'Nov-Dec': 16},
        8829: {'Jan-Feb': 17, 'Mar-Apr': 17, 'May-Jun': 17, 'Jul-Aug': 16, 'Sep-Oct': 16, 'Nov-Dec': 17},
        9130: {'Feb-Mar': 20, 'Apr-May': 20, 'Jun-Jul': 20, 'Aug-Sep': 20, 'Oct-Nov': 20},
        7929: {'Jan-Feb': 17, 'Mar-Apr': 16, 'May-Jun': 16, 'Jul-Aug': 17, 'Sep-Oct': 17, 'Nov-Dec': 17},
        8230: {'Feb-Mar': 20, 'Apr-May': 20, 'Jun-Jul': 20, 'Aug-Sep': 20, 'Oct-Nov': 20},
        8228: {'Feb-Mar': 20, 'Apr-May': 20, 'Jun-Jul': 20, 'Aug-Sep': 20, 'Oct-Nov': 20},
        8229: {'Jan-Feb': 17, 'Mar-Apr': 17, 'May-Jun': 17, 'Jul-Aug': 16, 'Sep-Oct': 16, 'Nov-Dec': 17}
    }
}

# === KING RANCH INCREMENTAL PRESET (Expansion grids for Challenger) ===
KING_RANCH_INCREMENTAL_PRESET = {
    'grids': [7930, 7931, 8231, 8528, 8529, 8828, 8829, 8830, 8831, 9131],
    'counties': {
        'Kenedy': [7930, 7931, 8231, 8829],
        'Brooks': [8528, 8529],
        'Kleberg': [8828, 8830, 8831, 9131]
    },
    'acres': {
        7930: 40137,
        7931: 40137,
        8231: 40137,
        8528: 26271,
        8529: 26271,
        8828: 50627,
        8829: 40137,
        8830: 50627,
        8831: 50627,
        9131: 50627
    }
}

# === GRID ID HELPER FUNCTION ===
def extract_numeric_grid_id(grid_id):
    """
    Extract numeric grid ID from formatted string.
    Examples:
        "9128 (Jim Wells - TX)" -> 9128
        "8230 (Brooks - TX)" -> 8230
        9128 -> 9128 (handles plain integers)
    """
    if isinstance(grid_id, str):
        # Extract first part before parenthesis
        return int(grid_id.split('(')[0].strip())
    else:
        # Already numeric
        return int(grid_id)

def extract_county_from_grid_id(grid_id):
    """
    Extract county name from formatted grid ID.
    Examples:
        "9128 (Kleberg - TX)" -> "Kleberg"
        "8230 (Brooks - TX)" -> "Brooks"
    """
    if isinstance(grid_id, str) and '(' in grid_id:
        # Extract content between parentheses
        county_part = grid_id.split('(')[1].split('-')[0].strip()
        return county_part
    return None

def get_safe_acres(source_acres, gid, default_acres, min_acres=1):
    """
    Get acres from source with a minimum floor to prevent 0 values.

    This is needed because MVO optimization may set some grids to 0 acres,
    but st.number_input with min_value=1 will fail if given value=0.

    Args:
        source_acres: Dictionary of grid_id -> acres from source portfolio
        gid: The grid ID to look up
        default_acres: Default value if grid not found
        min_acres: Minimum allowed acres (default 1)

    Returns:
        Integer acres value, guaranteed to be >= min_acres
    """
    acres = source_acres.get(gid, default_acres)
    if acres < min_acres:
        # If acres is 0 or near-zero, use default instead
        return max(min_acres, int(default_acres))
    return int(acres)

st.set_page_config(layout="wide", page_title="PRF Backtesting Tool")

# =============================================================================
# === 1. CACHED DATA-LOADING FUNCTIONS (FOR PERFORMANCE) ===
# =============================================================================

@st.cache_data(ttl=3600)
def load_distinct_grids(_session):
    """Fetches the list of all available Grid IDs with county names from PRF_COUNTY_BASE_VALUES."""
    query = """
        SELECT DISTINCT GRID_ID
        FROM CAPITAL_MARKETS_SANDBOX.PUBLIC.PRF_COUNTY_BASE_VALUES
        ORDER BY GRID_ID
    """
    df = _session.sql(query).to_pandas()
    return df['GRID_ID'].tolist()

@st.cache_data(ttl=3600)
def load_all_indices(_session, grid_id):
    """Fetches all historical rainfall data for a single grid, including ENSO phase."""
    # Extract numeric portion from formatted grid ID (e.g., "9128 (Jim Wells - TX)" -> 9128)
    numeric_grid_id = extract_numeric_grid_id(grid_id)

    all_indices_query = f"""
        SELECT
            YEAR, INTERVAL_NAME, INDEX_VALUE, INTERVAL_CODE, INTERVAL_MAPPING_TS_TEXT,
            OPTICAL_MAPPING_CPC
        FROM RAIN_INDEX_PLATINUM_ENHANCED
        WHERE GRID_ID = {numeric_grid_id}
        ORDER BY YEAR, INTERVAL_CODE
    """
    df = _session.sql(all_indices_query).to_pandas()
    df['INDEX_VALUE'] = pd.to_numeric(df['INDEX_VALUE'], errors='coerce')
    # FILTER OUT ROWS WITH NO RAINFALL DATA (for incomplete years like 2025)
    df = df.dropna(subset=['INDEX_VALUE'])
    return df

@st.cache_data(ttl=3600, show_spinner=False)
def filter_indices_by_scenario(all_indices_df, scenario, start_year=1948, end_year=2024):
    """
    Filter indices dataframe by scenario selection.
    Cached for performance when re-running with same parameters.
    """
    if scenario == 'All Years (except Current Year)':
        return all_indices_df[all_indices_df['YEAR'] < 2025]
    elif scenario == 'ENSO Phase: La Nina':
        if 'OPTICAL_MAPPING_CPC' in all_indices_df.columns:
            return all_indices_df[(all_indices_df['OPTICAL_MAPPING_CPC'] == 'La Nina') & (all_indices_df['YEAR'] < 2025)]
        return all_indices_df[all_indices_df['YEAR'] < 2025]
    elif scenario == 'ENSO Phase: El Nino':
        if 'OPTICAL_MAPPING_CPC' in all_indices_df.columns:
            return all_indices_df[(all_indices_df['OPTICAL_MAPPING_CPC'] == 'El Nino') & (all_indices_df['YEAR'] < 2025)]
        return all_indices_df[all_indices_df['YEAR'] < 2025]
    elif scenario == 'ENSO Phase: Neutral':
        if 'OPTICAL_MAPPING_CPC' in all_indices_df.columns:
            return all_indices_df[(all_indices_df['OPTICAL_MAPPING_CPC'] == 'Neutral') & (all_indices_df['YEAR'] < 2025)]
        return all_indices_df[all_indices_df['YEAR'] < 2025]
    elif scenario == 'Analog Years':
        # Filter to only analog years from session state
        analog_years = st.session_state.get('ps_analog_years', [])
        analog_year_list = [y['year'] for y in analog_years]
        if analog_year_list:
            return all_indices_df[all_indices_df['YEAR'].isin(analog_year_list)]
        return all_indices_df[all_indices_df['YEAR'] < 2025]
    else:  # Select my own interval
        return all_indices_df[(all_indices_df['YEAR'] >= start_year) & (all_indices_df['YEAR'] <= end_year)]


@st.cache_data(ttl=3600)
def load_zscore_data(_session, grid_id):
    """Load Z-score and ENSO data for market view filtering."""
    numeric_grid_id = extract_numeric_grid_id(grid_id)

    query = f"""
        SELECT
            YEAR,
            INTERVAL_CODE,
            INTERVAL_NAME,
            SEQUENTIAL_Z_SCORE_HISTORICAL_RECORD,
            SEQUENTIAL_Z_SCORE_5P,
            SEQUENTIAL_Z_SCORE_11P,
            OPTICAL_MAPPING_CPC
        FROM RAIN_INDEX_PLATINUM_ENHANCED
        WHERE GRID_ID = {numeric_grid_id}
        ORDER BY YEAR, INTERVAL_CODE
    """
    df = _session.sql(query).to_pandas()
    df['SEQUENTIAL_Z_SCORE_HISTORICAL_RECORD'] = pd.to_numeric(
        df['SEQUENTIAL_Z_SCORE_HISTORICAL_RECORD'], errors='coerce'
    )
    df['SEQUENTIAL_Z_SCORE_5P'] = pd.to_numeric(
        df['SEQUENTIAL_Z_SCORE_5P'], errors='coerce'
    )
    df['SEQUENTIAL_Z_SCORE_11P'] = pd.to_numeric(
        df['SEQUENTIAL_Z_SCORE_11P'], errors='coerce'
    )
    return df


def calculate_portfolio_aggregated_analog_years(session, selected_grids, regime, hist_context, trend):
    """
    Calculate analog years using PORTFOLIO-AGGREGATED methodology.

    Instead of filtering each grid independently, this:
    1. For each historical year, calculates the AVERAGE Z-score across ALL selected grids
    2. Determines the dominant ENSO phase across the portfolio
    3. Calculates portfolio-average trajectory
    4. Filters years where the PORTFOLIO AVERAGE matches the market view

    This ensures apples-to-apples comparison - all strategies are tested on the SAME years.

    Returns:
        List of dicts with keys: year, dominant_phase, phase_agreement, portfolio_avg_z,
        portfolio_trajectory, grids_with_data
    """
    from collections import Counter

    # Load Z-score data for all grids
    all_grid_data = {}
    for gid in selected_grids:
        try:
            zscore_df = load_zscore_data(session, gid)
            if not zscore_df.empty:
                all_grid_data[gid] = zscore_df
        except Exception as e:
            continue

    if len(all_grid_data) == 0:
        return []

    # Get all years present in any grid
    all_years = set()
    for gid, df in all_grid_data.items():
        all_years.update(df['YEAR'].unique())

    # Filter to complete years (exclude current/incomplete)
    all_years = [y for y in all_years if y < 2025]

    matching_years = []

    for year in sorted(all_years):
        year_grid_data = []
        year_phases = []
        year_trajectories = []

        for gid, df in all_grid_data.items():
            year_df = df[df['YEAR'] == year]
            if year_df.empty:
                continue

            # Calculate grid's average Z-score for the year
            zscore_vals = year_df['SEQUENTIAL_Z_SCORE_HISTORICAL_RECORD'].dropna()
            if len(zscore_vals) == 0:
                continue

            grid_avg_z = zscore_vals.mean()
            year_grid_data.append({
                'grid': gid,
                'avg_z': grid_avg_z
            })

            # Get dominant ENSO phase for this grid-year
            phases = year_df['OPTICAL_MAPPING_CPC'].dropna()
            if len(phases) > 0:
                phase_counts = Counter(phases)
                dominant_phase = phase_counts.most_common(1)[0][0]
                year_phases.append(dominant_phase)

            # Calculate trajectory using 11P/5P methodology
            # SOY 11P = Start of Year 11-period Z-score (from Jan-Feb interval - stable baseline)
            # EOY 5P = End of Year 5-period Z-score (from Nov-Dec interval - recent trend)
            # This captures intra-year momentum: where the year started vs where it's trending

            year_df_sorted = year_df.sort_values('INTERVAL_CODE')

            if len(year_df_sorted) >= 11:
                first_interval = year_df_sorted.iloc[0]   # Jan-Feb
                last_interval = year_df_sorted.iloc[-1]   # Nov-Dec

                z_11p_start = first_interval.get('SEQUENTIAL_Z_SCORE_11P') if 'SEQUENTIAL_Z_SCORE_11P' in first_interval.index else None
                z_5p_end = last_interval.get('SEQUENTIAL_Z_SCORE_5P') if 'SEQUENTIAL_Z_SCORE_5P' in last_interval.index else None

                if pd.notna(z_11p_start) and pd.notna(z_5p_end):
                    trajectory = z_5p_end - z_11p_start
                    year_trajectories.append(trajectory)

        # Require at least half the grids to have data
        min_grids_required = max(1, len(selected_grids) // 2)
        if len(year_grid_data) < min_grids_required:
            continue

        # Calculate portfolio averages
        portfolio_avg_z = np.mean([g['avg_z'] for g in year_grid_data])
        portfolio_trajectory = np.mean(year_trajectories) if year_trajectories else 0

        # Determine dominant ENSO phase across portfolio (requires majority)
        if year_phases:
            phase_counts = Counter(year_phases)
            dominant_phase, phase_count = phase_counts.most_common(1)[0]
            phase_agreement = phase_count / len(year_phases)

            # Require majority agreement on phase
            if phase_agreement < 0.5:
                dominant_phase = 'Mixed'
        else:
            dominant_phase = 'Unknown'
            phase_agreement = 0

        # Apply filters
        # ENSO regime filter
        if regime == 'Some La Nina':
            # Check if any grid-year has at least one La Nina interval
            has_la_nina = False
            for gid, df in all_grid_data.items():
                year_df = df[df['YEAR'] == year]
                if not year_df.empty:
                    phases = year_df['OPTICAL_MAPPING_CPC'].dropna()
                    if 'La Nina' in phases.values:
                        has_la_nina = True
                        break
            if not has_la_nina:
                continue

        elif regime == 'Some El Nino':
            # Check if any grid-year has at least one El Nino interval
            has_el_nino = False
            for gid, df in all_grid_data.items():
                year_df = df[df['YEAR'] == year]
                if not year_df.empty:
                    phases = year_df['OPTICAL_MAPPING_CPC'].dropna()
                    if 'El Nino' in phases.values:
                        has_el_nino = True
                        break
            if not has_el_nino:
                continue

        elif regime != 'Any':
            # Existing strict matching logic for La Nina, El Nino, Neutral
            if regime == 'La Nina' and dominant_phase != 'La Nina':
                continue
            elif regime == 'El Nino' and dominant_phase != 'El Nino':
                continue
            elif regime == 'Neutral' and dominant_phase != 'Neutral':
                continue

        # Historical context filter (based on portfolio average Z)
        if hist_context != 'Any':
            context_bounds = HISTORICAL_CONTEXT_MAP.get(hist_context, None)
            if context_bounds:
                if not (context_bounds['min'] <= portfolio_avg_z < context_bounds['max']):
                    continue

        # Trajectory filter
        if trend != 'Any':
            trend_bounds = TREND_MAP.get(trend, None)
            if trend_bounds:
                if not (trend_bounds['min'] <= portfolio_trajectory < trend_bounds['max']):
                    continue

        # Count ENSO intervals across all grids for this year
        la_nina_count = 0
        el_nino_count = 0
        grids_counted = 0
        for gid, df in all_grid_data.items():
            year_df = df[df['YEAR'] == year]
            if not year_df.empty:
                phases = year_df['OPTICAL_MAPPING_CPC'].dropna()
                la_nina_count += (phases == 'La Nina').sum()
                el_nino_count += (phases == 'El Nino').sum()
                grids_counted += 1

        # Calculate average per grid (more intuitive than total)
        avg_la_nina = la_nina_count / grids_counted if grids_counted > 0 else 0
        avg_el_nino = el_nino_count / grids_counted if grids_counted > 0 else 0

        # Year matches all criteria
        matching_years.append({
            'year': year,
            'dominant_phase': dominant_phase,
            'phase_agreement': phase_agreement,
            'portfolio_avg_z': portfolio_avg_z,
            'portfolio_trajectory': portfolio_trajectory,
            'grids_with_data': len(year_grid_data),
            'la_nina_intervals': avg_la_nina,  # Now average per grid
            'el_nino_intervals': avg_el_nino   # Now average per grid
        })

    return matching_years


@st.cache_data(ttl=3600)
def load_county_base_value(_session, grid_id):
    """Fetches the average county base value for the grid using GRID_ID."""
    base_value_query = f"""
        SELECT AVG(COUNTY_BASE_VALUE)
        FROM PRF_COUNTY_BASE_VALUES
        WHERE GRID_ID = '{grid_id}'
    """
    return float(_session.sql(base_value_query).to_pandas().iloc[0,0])

@st.cache_data(ttl=3600)
def get_current_rate_year(_session):
    """Finds the most recent year in the premium rates table."""
    return int(_session.sql("SELECT MAX(YEAR) FROM PRF_PREMIUM_RATES").to_pandas().iloc[0,0])

@st.cache_data(ttl=3600)
def load_premium_rates(_session, grid_id, use, coverage_levels_list, year):
    """Fetches premium rates for all specified coverage levels."""
    # Extract numeric portion from formatted grid ID
    numeric_grid_id = extract_numeric_grid_id(grid_id)
    
    all_premiums = {}
    for cov_level in coverage_levels_list:
        cov_string = f"{cov_level:.0%}"
        premium_query = f"""
            SELECT INDEX_INTERVAL_NAME, PREMIUMRATE 
            FROM PRF_PREMIUM_RATES 
            WHERE GRID_ID = {numeric_grid_id}
              AND INTENDED_USE = '{use}'
              AND COVERAGE_LEVEL = '{cov_string}'
              AND YEAR = {year}
        """
        prem_df = _session.sql(premium_query).to_pandas()
        prem_df['PREMIUMRATE'] = pd.to_numeric(prem_df['PREMIUMRATE'], errors='coerce')
        all_premiums[cov_level] = prem_df.set_index('INDEX_INTERVAL_NAME')['PREMIUMRATE'].to_dict()
    return all_premiums

@st.cache_data(ttl=3600)
def load_subsidies(_session, plan_code, coverage_levels_list):
    """Fetches subsidy percentages for all specified coverage levels."""
    all_subsidies = {}
    for cov_level in coverage_levels_list:
        subsidy_query = f"""
            SELECT SUBSIDY_PERCENT 
            FROM SUBSIDYPERCENT_YTD_PLATINUM 
            WHERE INSURANCE_PLAN_CODE = {plan_code}
              AND COVERAGE_LEVEL_PERCENT = {cov_level}
            LIMIT 1
        """
        all_subsidies[cov_level] = float(_session.sql(subsidy_query).to_pandas().iloc[0,0])
    return all_subsidies

# =============================================================================
# === 2. GLOBAL HELPER FUNCTIONS ===
# =============================================================================

def is_adjacent(interval1, interval2):
    """Check if two intervals are adjacent, with wrap-around"""
    try:
        idx1 = INTERVAL_ORDER_11.index(interval1)
        idx2 = INTERVAL_ORDER_11.index(interval2)
    except ValueError:
        return False # Interval not in list
    
    diff = abs(idx1 - idx2)
    # Check for direct adjacency (diff == 1) or wrap-around (diff == 10)
    return diff == 1 or diff == (len(INTERVAL_ORDER_11) - 1)

def has_adjacent_intervals_in_list(intervals_list):
    """Check if any intervals in the list are adjacent (excluding Nov-Dec/Jan-Feb wrap)"""
    for i in range(len(intervals_list)):
        for j in range(i+1, len(intervals_list)):
            interval1 = intervals_list[i]
            interval2 = intervals_list[j]
            
            # Check if they're adjacent
            if is_adjacent(interval1, interval2):
                # Allow Nov-Dec and Jan-Feb together (wrap-around exception)
                if (interval1 == 'Nov-Dec' and interval2 == 'Jan-Feb') or \
                   (interval1 == 'Jan-Feb' and interval2 == 'Nov-Dec'):
                    continue  # This is allowed
                else:
                    return True  # Adjacent intervals found - not allowed
    return False  # No adjacent intervals found

def generate_allocations(intervals_to_use, num_intervals):
    """
    Generate allocation percentages for N intervals, respecting all rules:
    - Only whole number percentages (1% increments)
    - Each interval: 0% OR 10-50%
    - Total must equal exactly 100%
    - Max 50% per interval

    Returns allocations as decimals (0.20 = 20%)
    """
    allocations = []

    if num_intervals == 1:
        # 1-interval: 50% max (can't reach 100% with one interval at 50% max)
        # This is actually invalid per rules - would need 2+ intervals
        allocations.append({intervals_to_use[0]: 0.50})

    elif num_intervals == 2:
        # 2-interval: Only 50/50 split is valid
        allocations.append({intervals_to_use[0]: 0.50, intervals_to_use[1]: 0.50})

    elif num_intervals == 3:
        # 3-interval patterns (whole numbers summing to 100%)
        splits = [
            (34, 33, 33),  # Equal-ish (100%)
            (50, 25, 25),  # Max one (100%)
            (40, 30, 30),  # Moderate (100%)
            (45, 30, 25),  # Graduated (100%)
            (50, 30, 20),  # Heavy concentration (100%)
            (40, 35, 25),  # Balanced (100%)
        ]
        for s in splits:
            allocations.append({intervals_to_use[i]: s[i]/100.0 for i in range(3)})

    elif num_intervals == 4:
        # 4-interval patterns (whole numbers summing to 100%)
        splits = [
            (25, 25, 25, 25),  # Equal (100%)
            (50, 20, 15, 15),  # Max one (100%)
            (40, 20, 20, 20),  # Moderate (100%)
            (35, 25, 25, 15),  # Graduated (100%)
            (30, 30, 20, 20),  # Two primaries (100%)
            (40, 25, 20, 15),  # Tiered (100%)
            (35, 30, 20, 15),  # Balanced (100%)
        ]
        for s in splits:
            allocations.append({intervals_to_use[i]: s[i]/100.0 for i in range(4)})

    elif num_intervals == 5:
        # 5-interval patterns (whole numbers summing to 100%)
        splits = [
            (20, 20, 20, 20, 20),  # Equal (100%)
            (50, 15, 15, 10, 10),  # Max one (100%)
            (30, 20, 20, 15, 15),  # Moderate (100%)
            (40, 15, 15, 15, 15),  # Higher concentration (100%)
            (25, 25, 20, 15, 15),  # Graduated (100%)
            (35, 20, 15, 15, 15),  # Two-tier (100%)
            (30, 25, 20, 15, 10),  # Descending (100%)
        ]
        for s in splits:
            allocations.append({intervals_to_use[i]: s[i]/100.0 for i in range(5)})

    elif num_intervals == 6:
        # 6-interval patterns (whole numbers summing to 100%)
        splits = [
            (17, 17, 17, 17, 16, 16),  # Equal-ish (100%)
            (50, 10, 10, 10, 10, 10),  # Max one (100%)
            (30, 15, 15, 15, 15, 10),  # Moderate (100%)
            (40, 12, 12, 12, 12, 12),  # Higher concentration (100%)
            (25, 20, 15, 15, 15, 10),  # Graduated (100%)
            (35, 15, 15, 15, 10, 10),  # Two-tier (100%)
            (20, 20, 15, 15, 15, 15),  # Balanced (100%)
        ]
        for s in splits:
            allocations.append({intervals_to_use[i]: s[i]/100.0 for i in range(6)})

    return allocations

def is_valid_allocation(alloc_dict):
    """
    Check if allocation meets all rules:
    - Whole number percentages (1% increments)
    - Each interval: 0% OR 10-50%
    - Total equals 100%
    - Max 50% per interval
    """
    total = sum(alloc_dict.values())
    if abs(total - 1.0) > 0.001:
        return False

    for interval, pct in alloc_dict.items():
        # Check max 50%
        if pct > 0.50:
            return False
        # Check 10% minimum (must be 0% or >= 10%)
        if pct > 0 and pct < 0.10:
            return False
        # Check whole number (allow small floating point errors)
        pct_as_percent = pct * 100
        if abs(pct_as_percent - round(pct_as_percent)) > 0.001:
            return False

    return True


def generate_marginal_variations(base_allocation_dict):
    """
    Generate subtle variations of an existing King Ranch allocation.
    Returns list of allocation dictionaries (as decimals, not percentages).
    All variations maintain whole number percentages and 10% minimum rule.
    """
    variations = []

    # Convert percentages to decimals if needed, round to whole percentages
    base_alloc = {}
    for k, v in base_allocation_dict.items():
        if v > 0:
            decimal_val = v / 100.0 if v > 1 else v
            # Round to nearest whole percentage
            decimal_val = round(decimal_val * 100) / 100.0
            base_alloc[k] = decimal_val

    # 1. Original allocation (normalized to decimals)
    if is_valid_allocation(base_alloc):
        variations.append(base_alloc.copy())

    # 2. Shift forward by one month
    shifted_forward = {}
    for interval, pct in base_alloc.items():
        idx = INTERVAL_ORDER_11.index(interval)
        new_idx = (idx + 1) % len(INTERVAL_ORDER_11)
        new_interval = INTERVAL_ORDER_11[new_idx]
        shifted_forward[new_interval] = pct

    if not has_adjacent_intervals_in_list(list(shifted_forward.keys())) and is_valid_allocation(shifted_forward):
        variations.append(shifted_forward)

    # 3. Shift backward by one month
    shifted_backward = {}
    for interval, pct in base_alloc.items():
        idx = INTERVAL_ORDER_11.index(interval)
        new_idx = (idx - 1) % len(INTERVAL_ORDER_11)
        new_interval = INTERVAL_ORDER_11[new_idx]
        shifted_backward[new_interval] = pct

    if not has_adjacent_intervals_in_list(list(shifted_backward.keys())) and is_valid_allocation(shifted_backward):
        variations.append(shifted_backward)

    # 4. Shift forward by two months
    shifted_forward_2 = {}
    for interval, pct in base_alloc.items():
        idx = INTERVAL_ORDER_11.index(interval)
        new_idx = (idx + 2) % len(INTERVAL_ORDER_11)
        new_interval = INTERVAL_ORDER_11[new_idx]
        shifted_forward_2[new_interval] = pct

    if not has_adjacent_intervals_in_list(list(shifted_forward_2.keys())) and is_valid_allocation(shifted_forward_2):
        variations.append(shifted_forward_2)

    # 5. Shift backward by two months
    shifted_backward_2 = {}
    for interval, pct in base_alloc.items():
        idx = INTERVAL_ORDER_11.index(interval)
        new_idx = (idx - 2) % len(INTERVAL_ORDER_11)
        new_interval = INTERVAL_ORDER_11[new_idx]
        shifted_backward_2[new_interval] = pct

    if not has_adjacent_intervals_in_list(list(shifted_backward_2.keys())) and is_valid_allocation(shifted_backward_2):
        variations.append(shifted_backward_2)

    # 6. Minor percentage adjustments (redistribute ±1% between intervals)
    # Only produce variations that maintain 10% minimum rule
    if len(base_alloc) >= 2:
        intervals_list = list(base_alloc.keys())
        for i in range(len(intervals_list)):
            for j in range(i + 1, len(intervals_list)):
                # Create variation: take 1% from interval i, give to interval j
                variation = base_alloc.copy()
                new_val_i = variation[intervals_list[i]] - 0.01
                new_val_j = variation[intervals_list[j]] + 0.01

                # Check if new values meet 10% minimum rule
                if (new_val_i == 0 or new_val_i >= 0.10) and new_val_j <= 0.50:
                    variation[intervals_list[i]] = new_val_i
                    variation[intervals_list[j]] = new_val_j
                    if is_valid_allocation(variation):
                        variations.append(variation)

    return variations

def generate_incremental_variations(base_allocation_dict):
    """
    Generate incremental percentage adjustments within existing King Ranch allocation.
    Only does small fine-tuning: ±1%, ±2%, ±3%, ±4%, ±5% between intervals.
    Does NOT change which intervals are selected.
    Returns list of allocation dictionaries (as decimals, not percentages).
    All variations maintain whole number percentages and 10% minimum rule.
    """
    variations = []

    # Convert percentages to decimals if needed, round to whole percentages
    base_alloc = {}
    for k, v in base_allocation_dict.items():
        if v > 0:
            decimal_val = v / 100.0 if v > 1 else v
            # Round to nearest whole percentage
            decimal_val = round(decimal_val * 100) / 100.0
            base_alloc[k] = decimal_val

    # Get the intervals that are allocated (non-zero)
    active_intervals = [k for k, v in base_alloc.items() if v > 0]
    num_intervals = len(active_intervals)

    if num_intervals == 0:
        return variations

    # 1. Original allocation
    if is_valid_allocation(base_alloc):
        variations.append(base_alloc.copy())

    # 2. Small fine-tuning adjustments only: ±1%, ±2%, ±3%, ±4%, ±5%
    small_adjustments = [0.01, 0.02, 0.03, 0.04, 0.05]

    for adj in small_adjustments:
        for i in range(num_intervals):
            for j in range(i + 1, num_intervals):
                # Take from i, give to j
                var1 = {k: 0.0 for k in INTERVAL_ORDER_11}
                for k, v in base_alloc.items():
                    var1[k] = v

                new_val_i = var1[active_intervals[i]] - adj
                new_val_j = var1[active_intervals[j]] + adj

                # Check 10% minimum rule (must be 0% or >= 10%)
                if (new_val_i == 0 or new_val_i >= 0.10) and new_val_j <= 0.50:
                    var1[active_intervals[i]] = new_val_i
                    var1[active_intervals[j]] = new_val_j
                    if is_valid_allocation(var1):
                        variations.append(var1.copy())

                # Take from j, give to i
                var2 = {k: 0.0 for k in INTERVAL_ORDER_11}
                for k, v in base_alloc.items():
                    var2[k] = v

                new_val_j2 = var2[active_intervals[j]] - adj
                new_val_i2 = var2[active_intervals[i]] + adj

                # Check 10% minimum rule (must be 0% or >= 10%)
                if (new_val_j2 == 0 or new_val_j2 >= 0.10) and new_val_i2 <= 0.50:
                    var2[active_intervals[j]] = new_val_j2
                    var2[active_intervals[i]] = new_val_i2
                    if is_valid_allocation(var2):
                        variations.append(var2.copy())

    return variations

# =============================================================================
# === PORTFOLIO OPTIMIZATION FUNCTIONS (Vectorized) ===
# =============================================================================

import random

# Minimum allocation per interval (10%)
MIN_ALLOCATION = 0.10
# Number of intervals (11 bi-monthly periods)
INTERVAL_RANGE = 11

def generate_random_valid_allocation(min_allocation=MIN_ALLOCATION, interval_range=INTERVAL_RANGE, interval_count_range=(2, 6)):
    """
    Generate a random valid allocation as a numpy array.
    Rules: 10-50% per interval, non-adjacent, total = 100%.
    Returns numpy array of length 11.

    Args:
        min_allocation: Minimum allocation per interval (default 0.10)
        interval_range: Total number of intervals (default 11)
        interval_count_range: Tuple of (min, max) number of active intervals (default (2, 6))
    """
    weights = np.zeros(interval_range)

    # Randomly select intervals within the specified range
    min_count, max_count = interval_count_range
    num_intervals = random.randint(min_count, max_count)
    available = list(range(interval_range))
    selected = []

    for _ in range(num_intervals):
        if not available:
            break
        idx = random.choice(available)
        selected.append(idx)

        # Remove this index and adjacent ones
        to_remove = [idx]
        if idx > 0:
            to_remove.append(idx - 1)
        if idx < interval_range - 1:
            to_remove.append(idx + 1)
        # Handle wrap-around (but we allow Nov-Dec/Jan-Feb adjacency)
        available = [i for i in available if i not in to_remove]

    if len(selected) < 2:
        # Fallback to safe default
        selected = [2, 4, 8]  # Mar-Apr, May-Jun, Sep-Oct

    # Generate random percentages
    remaining = 100
    for i, idx in enumerate(selected[:-1]):
        max_pct = min(50, remaining - min_allocation * 100 * (len(selected) - i - 1))
        min_pct = max(min_allocation * 100, remaining - 50 * (len(selected) - i - 1))

        if min_pct > max_pct:
            pct = min_pct
        else:
            pct = random.randint(int(min_pct), int(max_pct))

        weights[idx] = pct / 100.0
        remaining -= pct

    # Last interval gets remainder
    if remaining >= 10 and remaining <= 50:
        weights[selected[-1]] = remaining / 100.0
    else:
        weights[selected[-1]] = 0.20
        # Normalize
        if weights.sum() > 0:
            weights = weights / weights.sum()

    return weights

def generate_full_coverage_allocation(grid_index, min_allocation=MIN_ALLOCATION, interval_range=INTERVAL_RANGE):
    """
    Generate allocation for full calendar coverage mode.
    Uses staggered Pattern A (6 intervals) and Pattern B (5 intervals)
    to cover all 11 intervals while maintaining non-adjacency within each grid.

    Args:
        grid_index: Index of the grid (0 or 1 to determine pattern)
        min_allocation: Minimum allocation per interval (default 0.10)
        interval_range: Total number of intervals (default 11)

    Returns:
        numpy array of length 11 with valid allocation
    """
    weights = np.zeros(interval_range)

    # Pattern A: Even positions (0, 2, 4, 6, 8, 10) - 6 intervals
    # Pattern B: Odd positions (1, 3, 5, 7, 9) - 5 intervals
    if grid_index % 2 == 0:
        selected = [0, 2, 4, 6, 8, 10]  # Pattern A - 6 intervals
    else:
        selected = [1, 3, 5, 7, 9]  # Pattern B - 5 intervals

    # Generate random percentages within bounds (10-50% each)
    remaining = 100
    for i, idx in enumerate(selected[:-1]):
        max_pct = min(50, remaining - min_allocation * 100 * (len(selected) - i - 1))
        min_pct = max(min_allocation * 100, remaining - 50 * (len(selected) - i - 1))

        if min_pct > max_pct:
            pct = min_pct
        else:
            pct = random.randint(int(min_pct), int(max_pct))

        weights[idx] = pct / 100.0
        remaining -= pct

    # Last interval gets remainder
    if remaining >= 10 and remaining <= 50:
        weights[selected[-1]] = remaining / 100.0
    else:
        weights[selected[-1]] = 0.20
        # Normalize
        if weights.sum() > 0:
            weights = weights / weights.sum()

    return weights

def generate_naive_weights(min_allocation=MIN_ALLOCATION, interval_range=INTERVAL_RANGE):
    """
    Generate naive equal distribution across non-adjacent intervals.
    Returns numpy array of length 11.
    """
    weights = np.zeros(interval_range)

    # Select spread-out intervals: Jan-Feb(0), Mar-Apr(2), May-Jun(4), Jul-Aug(6), Sep-Oct(8)
    selected_indices = [0, 2, 4, 6, 8]

    # Equal distribution
    pct_each = 1.0 / len(selected_indices)
    for idx in selected_indices:
        weights[idx] = pct_each

    # Round to 1% increments
    weights = np.round(weights / 0.01) * 0.01

    # Fix rounding
    diff = 1.0 - weights.sum()
    if abs(diff) > 0.001:
        weights[selected_indices[0]] += diff

    return weights

def generate_marginal_candidate(base_weights, min_allocation=MIN_ALLOCATION, interval_range=INTERVAL_RANGE):
    """
    Perturbs base_weights by shifting small amounts (5%) between intervals
    or adjacent time periods, maintaining validity.
    """
    candidate = base_weights.copy()
    active_indices = np.where(candidate > 0.001)[0]

    # Fallback if empty
    if len(active_indices) == 0:
        return generate_random_valid_allocation(min_allocation, interval_range)

    # Type A: Weight Shift (Move 5% from one active interval to another)
    if random.random() < 0.6 and len(active_indices) > 1:
        idx1, idx2 = np.random.choice(active_indices, 2, replace=False)
        shift = 0.05
        # Check bounds
        if candidate[idx1] - shift >= min_allocation and candidate[idx2] + shift <= 0.50:
            candidate[idx1] -= shift
            candidate[idx2] += shift

    # Type B: Time Shift (Move 5% to a neighbor)
    else:
        idx = np.random.choice(active_indices)
        # Check left/right neighbors
        neighbors = []
        if idx > 0:
            neighbors.append(idx - 1)
        if idx < interval_range - 1:
            neighbors.append(idx + 1)

        if neighbors:
            target = np.random.choice(neighbors)
            shift = 0.05

            # If target is currently 0, we must move at least min_allocation to activate it
            if candidate[target] < 0.001:
                shift = min_allocation

            # Check bounds and constraints
            if candidate[idx] - shift >= min_allocation and candidate[target] + shift <= 0.50:
                # Check adjacency rules - ensure target doesn't violate adjacency with other active intervals
                other_active = [i for i in active_indices if i != idx]
                is_adjacent_violation = False
                for other_idx in other_active:
                    if abs(target - other_idx) == 1:
                        # Allow Nov-Dec(10) and Jan-Feb(0) adjacency exception
                        if not ((target == 0 and other_idx == 10) or (target == 10 and other_idx == 0)):
                            is_adjacent_violation = True
                            break

                if not is_adjacent_violation:
                    candidate[idx] -= shift
                    candidate[target] += shift

    # Normalize and Round
    candidate = np.clip(candidate, 0, 0.50)
    if candidate.sum() > 0:
        candidate = candidate / candidate.sum()
        candidate = np.round(candidate / 0.01) * 0.01  # Maintain 1% increments
        # Fix rounding errors
        diff = 1.0 - candidate.sum()
        if abs(diff) > 0.001:
            candidate[np.argmax(candidate)] += diff

    return candidate

def calculate_vectorized_roi(weights_batch, index_matrix, premium_rates_array,
                              coverage_level, subsidy, total_protection):
    """
    Vectorized ROI calculation for a batch of weight candidates.

    Args:
        weights_batch: numpy array of shape (n_candidates, 11) - allocation weights
        index_matrix: numpy array of shape (n_years, 11) - index values per year/interval
        premium_rates_array: numpy array of shape (11,) - premium rates per interval
        coverage_level: float (e.g., 0.80)
        subsidy: float (e.g., 0.59)
        total_protection: float - total policy protection amount

    Returns:
        numpy array of shape (n_candidates,) - cumulative ROI for each candidate
    """
    n_candidates = weights_batch.shape[0]
    n_years = index_matrix.shape[0]

    # Calculate protection per interval for each candidate: (n_candidates, 11)
    interval_protection = weights_batch * total_protection

    # Calculate premium per interval: (n_candidates, 11)
    total_premium = interval_protection * premium_rates_array
    producer_premium = total_premium * (1 - subsidy)

    # Sum producer premium across intervals for each candidate: (n_candidates,)
    annual_premium = producer_premium.sum(axis=1)

    # Total premium over all years: (n_candidates,)
    total_premium_all_years = annual_premium * n_years

    # Calculate trigger level
    trigger = coverage_level * 100

    # Calculate shortfall for each year/interval: (n_years, 11)
    shortfall_pct = np.maximum(0, (trigger - index_matrix) / trigger)

    # Calculate indemnity: broadcast (n_candidates, 11) * (n_years, 11) -> need to iterate or reshape
    # For each candidate, sum indemnity across all years and intervals
    # indemnity[c, y, i] = shortfall_pct[y, i] * interval_protection[c, i]

    # Reshape for broadcasting: (1, n_years, 11) and (n_candidates, 1, 11)
    shortfall_expanded = shortfall_pct[np.newaxis, :, :]  # (1, n_years, 11)
    protection_expanded = interval_protection[:, np.newaxis, :]  # (n_candidates, 1, 11)

    # Indemnity: (n_candidates, n_years, 11)
    indemnity = shortfall_expanded * protection_expanded

    # Sum across years and intervals: (n_candidates,)
    total_indemnity = indemnity.sum(axis=(1, 2))

    # Calculate ROI
    roi = np.where(
        total_premium_all_years > 0,
        (total_indemnity - total_premium_all_years) / total_premium_all_years,
        -1.0
    )

    return roi

@st.cache_data(ttl=3600, show_spinner=False)
def run_fast_optimization_core(
    _session, grid_id, start_year, end_year, plan_code, productivity_factor,
    acres, intended_use, coverage_level, iterations, search_mode,
    require_full_coverage=False, interval_range_opt=(2, 6), grid_index=0
):
    """
    Core optimization function with caching.

    Args:
        _session: Snowflake session (excluded from cache key with underscore prefix)
        grid_id: Grid identifier
        start_year, end_year: Year range for backtest
        plan_code: Insurance plan code
        productivity_factor: Productivity multiplier
        acres: Number of acres
        intended_use: Use type (Grazing, Haying)
        coverage_level: Coverage level (e.g., 0.80)
        iterations: Number of iterations for global search
        search_mode: 'global' or 'marginal'
        require_full_coverage: If True, use Pattern A/B for full calendar coverage
        interval_range_opt: Tuple of (min, max) active intervals when not using full coverage
        grid_index: Grid index for determining Pattern A/B in full coverage mode

    Returns:
        Tuple of (best_allocation_dict, best_roi, strategies_tested)
    """
    # Load data
    county_base_value = load_county_base_value(_session, grid_id)
    all_indices_df = load_all_indices(_session, grid_id)
    all_indices_df = all_indices_df[
        (all_indices_df['YEAR'] >= start_year) &
        (all_indices_df['YEAR'] <= end_year)
    ]

    current_rate_year = get_current_rate_year(_session)
    premium_rates = load_premium_rates(_session, grid_id, intended_use, [coverage_level], current_rate_year)[coverage_level]
    subsidy = load_subsidies(_session, plan_code, [coverage_level])[coverage_level]

    dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
    total_protection = dollar_protection * acres

    # Build index matrix: (n_years, 11)
    years = sorted(all_indices_df['YEAR'].unique())
    n_years = len(years)
    index_matrix = np.zeros((n_years, INTERVAL_RANGE))

    for y_idx, year in enumerate(years):
        year_data = all_indices_df[all_indices_df['YEAR'] == year]
        for interval_idx, interval_name in enumerate(INTERVAL_ORDER_11):
            row = year_data[year_data['INTERVAL_NAME'] == interval_name]
            if not row.empty:
                index_matrix[y_idx, interval_idx] = float(row['INDEX_VALUE'].iloc[0])
            else:
                index_matrix[y_idx, interval_idx] = 100.0  # Default to no shortfall

    # Build premium rates array: (11,)
    premium_rates_array = np.array([
        premium_rates.get(interval, 0) for interval in INTERVAL_ORDER_11
    ])

    # Generate candidates based on search mode and diversification constraints
    candidates = []

    if require_full_coverage:
        # Full calendar coverage mode - use Pattern A or B based on grid index
        for _ in range(iterations):
            candidates.append(generate_full_coverage_allocation(grid_index))
    elif search_mode == 'marginal':
        # Start from naive allocation and perturb
        naive_weights = generate_naive_weights()
        candidates.append(naive_weights)

        # Generate marginal variations
        for _ in range(iterations):
            # Pick a random existing candidate as base
            base = candidates[random.randint(0, len(candidates) - 1)]
            new_candidate = generate_marginal_candidate(base.copy())
            candidates.append(new_candidate)
    else:
        # Global search - random valid allocations with custom interval range
        for _ in range(iterations):
            candidates.append(generate_random_valid_allocation(interval_count_range=interval_range_opt))

    # Convert to batch array: (n_candidates, 11)
    weights_batch = np.array(candidates)

    # Vectorized ROI calculation
    roi_scores = calculate_vectorized_roi(
        weights_batch, index_matrix, premium_rates_array,
        coverage_level, subsidy, total_protection
    )

    # Find best
    best_idx = np.argmax(roi_scores)
    best_weights = weights_batch[best_idx]
    best_roi = roi_scores[best_idx]

    # Convert weights array to dictionary
    best_allocation = {}
    for idx, interval in enumerate(INTERVAL_ORDER_11):
        if best_weights[idx] > 0.001:
            best_allocation[interval] = round(best_weights[idx], 2)

    return best_allocation, float(best_roi), len(candidates)


def run_analog_year_optimization(
    session, grid_id, analog_years, plan_code, productivity_factor,
    acres, intended_use, coverage_level, iterations, interval_range_opt,
    objective='cumulative_roi'
):
    """
    Run optimization for a single grid using ONLY the specified analog years.

    IMPROVED: Uses interval-weighted search instead of pure random.
    Scores intervals by average shortfall during analog years, then focuses
    the search on high-shortfall intervals (similar to Challenger 1 approach).

    Args:
        session: Snowflake session
        grid_id: Grid identifier
        analog_years: List of year integers (e.g., [1954, 1971, 1988, ...])
        plan_code: Insurance plan code
        productivity_factor: Productivity multiplier
        acres: Number of acres
        intended_use: Use type (Grazing, Haying)
        coverage_level: Coverage level (e.g., 0.80)
        iterations: Number of iterations for optimization
        interval_range_opt: Tuple of (min, max) active intervals
        objective: Optimization objective - 'cumulative_roi', 'median_roi', 'profitable_pct', or 'risk_adj_ret'

    Returns:
        Tuple of (best_allocation_dict, best_roi, strategies_tested)
    """
    if not analog_years or len(analog_years) == 0:
        return {}, 0.0, 0

    # Load data
    county_base_value = load_county_base_value(session, grid_id)
    all_indices_df = load_all_indices(session, grid_id)

    # Filter to ONLY analog years
    analog_indices_df = all_indices_df[all_indices_df['YEAR'].isin(analog_years)]

    if analog_indices_df.empty:
        return {}, 0.0, 0

    current_rate_year = get_current_rate_year(session)
    premium_rates = load_premium_rates(session, grid_id, intended_use, [coverage_level], current_rate_year)[coverage_level]
    subsidy = load_subsidies(session, plan_code, [coverage_level])[coverage_level]

    dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
    total_protection = dollar_protection * acres

    # === KEY IMPROVEMENT: Score intervals by analog-year shortfall ===
    trigger = coverage_level * 100
    interval_scores = {}

    for interval in INTERVAL_ORDER_11:
        interval_data = analog_indices_df[analog_indices_df['INTERVAL_NAME'] == interval]['INDEX_VALUE']
        if len(interval_data) > 0:
            # Calculate average shortfall (how much below trigger)
            # Higher shortfall = more indemnity potential = better interval to insure
            avg_index = interval_data.mean()
            avg_shortfall = max(0, trigger - avg_index)

            # Also consider frequency of payouts
            payout_frequency = (interval_data < trigger).sum() / len(interval_data)

            # Combined score: shortfall magnitude * frequency
            # This favors intervals that both pay out often AND pay out big
            interval_scores[interval] = avg_shortfall * (0.5 + payout_frequency)
        else:
            interval_scores[interval] = 0

    # Sort intervals by score (highest shortfall first)
    sorted_intervals = sorted(interval_scores.items(), key=lambda x: x[1], reverse=True)

    # Determine search depth based on iterations
    # More iterations = explore more intervals
    if iterations >= 7000:
        search_depth = 9  # Near-exhaustive
    elif iterations >= 3000:
        search_depth = 7  # Thorough
    elif iterations >= 1000:
        search_depth = 6  # Standard
    else:
        search_depth = 5  # Fast

    top_intervals = [x[0] for x in sorted_intervals[:search_depth]]

    # === Generate candidates using INTELLIGENT combinatorial search ===
    candidates = []
    min_intervals, max_intervals = interval_range_opt

    # Phase 1: Systematic combinations of top intervals (like Challenger 1)
    for num_intervals in range(min_intervals, min(max_intervals + 1, len(top_intervals) + 1)):
        for combo in combinations(top_intervals, num_intervals):
            # Skip adjacent intervals (except Nov-Dec/Jan-Feb wrap)
            if has_adjacent_intervals_in_list(list(combo)):
                continue

            # Generate multiple allocation patterns for this combination
            combo_allocations = generate_allocations(list(combo), num_intervals)
            candidates.extend(combo_allocations)

    # Phase 2: Add weighted random candidates that favor high-score intervals
    # This explores beyond the strict top-N but still biased toward good intervals
    total_score = sum(max(0.1, score) for score in interval_scores.values())
    interval_weights = {k: max(0.1, v) / total_score for k, v in interval_scores.items()}

    remaining_iterations = max(0, iterations - len(candidates))

    for _ in range(remaining_iterations):
        # Weighted random selection of intervals
        num_intervals = random.randint(min_intervals, max_intervals)

        # Select intervals with probability proportional to their score
        available = list(INTERVAL_ORDER_11)
        selected = []

        for _ in range(num_intervals):
            if not available:
                break

            # Calculate weights for available intervals
            weights = [interval_weights.get(i, 0.1) for i in available]
            total_w = sum(weights)
            if total_w == 0:
                break

            probs = [w / total_w for w in weights]

            # Weighted random choice
            r = random.random()
            cumsum = 0
            chosen_idx = 0
            for i, p in enumerate(probs):
                cumsum += p
                if r <= cumsum:
                    chosen_idx = i
                    break

            chosen = available[chosen_idx]
            selected.append(chosen)

            # Remove chosen and adjacent intervals from available
            to_remove = [chosen]
            chosen_global_idx = INTERVAL_ORDER_11.index(chosen)
            if chosen_global_idx > 0:
                to_remove.append(INTERVAL_ORDER_11[chosen_global_idx - 1])
            if chosen_global_idx < len(INTERVAL_ORDER_11) - 1:
                to_remove.append(INTERVAL_ORDER_11[chosen_global_idx + 1])

            available = [i for i in available if i not in to_remove]

        if len(selected) >= 2:
            # Generate allocation for selected intervals
            alloc = generate_weighted_allocation(selected, interval_scores)
            if alloc and is_valid_allocation(alloc):
                candidates.append(alloc)

    # Deduplicate candidates
    unique_candidates = []
    seen = set()
    for candidate in candidates:
        # Handle both dict and array formats
        if isinstance(candidate, dict):
            key = tuple(sorted((k, round(v, 2)) for k, v in candidate.items() if v > 0))
        else:
            key = tuple(round(v, 2) for v in candidate)
        if key not in seen:
            seen.add(key)
            unique_candidates.append(candidate)

    if len(unique_candidates) == 0:
        # Fallback to naive allocation
        return generate_naive_weights_as_dict(), 0.0, 0

    # === Vectorized ROI calculation ===
    # Build index matrix for analog years only
    years = sorted(analog_indices_df['YEAR'].unique())
    n_years = len(years)

    if n_years == 0:
        return {}, 0.0, 0

    index_matrix = np.zeros((n_years, INTERVAL_RANGE))

    for y_idx, year in enumerate(years):
        year_data = analog_indices_df[analog_indices_df['YEAR'] == year]
        for interval_idx, interval_name in enumerate(INTERVAL_ORDER_11):
            row = year_data[year_data['INTERVAL_NAME'] == interval_name]
            if not row.empty:
                index_matrix[y_idx, interval_idx] = float(row['INDEX_VALUE'].iloc[0])
            else:
                index_matrix[y_idx, interval_idx] = 100.0  # Default to no shortfall

    # Build premium rates array
    premium_rates_array = np.array([
        premium_rates.get(interval, 0) for interval in INTERVAL_ORDER_11
    ])

    # Convert candidates to batch array (handle both dict and array formats)
    weights_batch = np.zeros((len(unique_candidates), INTERVAL_RANGE))
    for i, candidate in enumerate(unique_candidates):
        if isinstance(candidate, dict):
            for j, interval in enumerate(INTERVAL_ORDER_11):
                weights_batch[i, j] = candidate.get(interval, 0)
        else:
            weights_batch[i] = candidate

    # Vectorized ROI calculation - cumulative ROI for default/fallback
    roi_scores = calculate_vectorized_roi(
        weights_batch, index_matrix, premium_rates_array,
        coverage_level, subsidy, total_protection
    )

    # For non-cumulative objectives, calculate per-year ROI for each candidate
    if objective != 'cumulative_roi':
        n_candidates = weights_batch.shape[0]

        # Calculate protection and premium per interval for each candidate
        interval_protection = weights_batch * total_protection
        total_premium_per_interval = interval_protection * premium_rates_array
        producer_premium_per_interval = total_premium_per_interval * (1 - subsidy)
        annual_premium = producer_premium_per_interval.sum(axis=1)  # (n_candidates,)

        trigger = coverage_level * 100
        shortfall_pct = np.maximum(0, (trigger - index_matrix) / trigger)  # (n_years, 11)

        # Calculate per-year ROI for each candidate
        # yearly_roi[c, y] = (indemnity[c, y] - premium[c]) / premium[c]
        yearly_roi = np.zeros((n_candidates, n_years))

        for c_idx in range(n_candidates):
            prem = annual_premium[c_idx]
            if prem <= 0:
                yearly_roi[c_idx, :] = -1.0
                continue

            for y_idx in range(n_years):
                year_indemnity = np.sum(shortfall_pct[y_idx, :] * interval_protection[c_idx, :])
                yearly_roi[c_idx, y_idx] = (year_indemnity - prem) / prem

        # Calculate objective scores based on selection
        if objective == 'median_roi':
            objective_scores = np.median(yearly_roi, axis=1)
        elif objective == 'profitable_pct':
            objective_scores = np.mean(yearly_roi > 0, axis=1)  # Fraction of profitable years
        elif objective == 'risk_adj_ret':
            # Risk-adjusted return = mean / std (Sharpe-like ratio)
            means = np.mean(yearly_roi, axis=1)
            stds = np.std(yearly_roi, axis=1)
            stds = np.where(stds == 0, 1e-9, stds)  # Avoid division by zero
            objective_scores = means / stds
        else:
            objective_scores = roi_scores  # Fallback to cumulative

        best_idx = np.argmax(objective_scores)
    else:
        best_idx = np.argmax(roi_scores)

    best_weights = weights_batch[best_idx]
    best_roi = roi_scores[best_idx]

    # Convert weights array to dictionary
    best_allocation = {}
    for idx, interval in enumerate(INTERVAL_ORDER_11):
        if best_weights[idx] > 0.001:
            best_allocation[interval] = round(best_weights[idx], 2)

    return best_allocation, float(best_roi), len(unique_candidates)


def generate_weighted_allocation(selected_intervals, interval_scores):
    """
    Generate an allocation for selected intervals, weighting by their scores.
    Higher-scoring intervals get more allocation (within 10-50% bounds).

    Args:
        selected_intervals: List of interval names to allocate to
        interval_scores: Dict of interval -> score

    Returns:
        Dict of interval -> weight (as decimal, e.g., 0.25 for 25%)
    """
    if len(selected_intervals) == 0:
        return None

    # Get scores for selected intervals
    scores = [max(0.1, interval_scores.get(i, 0.1)) for i in selected_intervals]
    total_score = sum(scores)

    # Calculate raw proportional weights
    raw_weights = [s / total_score for s in scores]

    # Apply constraints: each must be 10-50%, total must be 100%
    # Start with proportional, then clamp and redistribute
    weights = []
    for w in raw_weights:
        clamped = max(0.10, min(0.50, w))
        weights.append(clamped)

    # Normalize to sum to 1.0
    total = sum(weights)
    if total > 0:
        weights = [w / total for w in weights]

    # Round to whole percentages
    weights = [round(w * 100) / 100 for w in weights]

    # Fix rounding errors
    diff = 1.0 - sum(weights)
    if abs(diff) > 0.001:
        # Add/subtract from largest weight
        max_idx = weights.index(max(weights))
        weights[max_idx] += diff

    # Final validation
    for w in weights:
        if w > 0 and (w < 0.10 or w > 0.50):
            # Fallback to equal distribution
            equal_weight = round(1.0 / len(selected_intervals), 2)
            weights = [equal_weight] * len(selected_intervals)
            diff = 1.0 - sum(weights)
            weights[0] += diff
            break

    # Build allocation dict
    allocation = {interval: 0.0 for interval in INTERVAL_ORDER_11}
    for i, interval in enumerate(selected_intervals):
        allocation[interval] = weights[i]

    return allocation


def generate_naive_weights_as_dict():
    """
    Generate naive equal distribution as a dictionary.
    Fallback when no valid candidates found.
    """
    selected_indices = [0, 2, 4, 6, 8]  # Jan-Feb, Mar-Apr, May-Jun, Jul-Aug, Sep-Oct
    allocation = {interval: 0.0 for interval in INTERVAL_ORDER_11}
    pct_each = 0.20
    for idx in selected_indices:
        allocation[INTERVAL_ORDER_11[idx]] = pct_each
    return allocation


def run_weather_mvo_optimization(
    session, weather_grids, weather_acres, analog_years,
    coverage_level, productivity_factor, intended_use, plan_code,
    interval_range_opt, optimization_iterations, risk_aversion=1.0, max_turnover=0.20
):
    """
    Two-stage weather-aware portfolio optimization with MVO.

    Stage 1: Optimize intervals per grid using analog years
    Stage 2: Optimize acre distribution using MVO on analog year correlations

    Args:
        session: Snowflake session
        weather_grids: List of grid IDs
        weather_acres: Dict of grid_id -> starting acres
        analog_years: List of year integers
        coverage_level: Coverage level (e.g., 0.75)
        productivity_factor: Productivity multiplier
        intended_use: 'Grazing' or 'Haying'
        plan_code: Insurance plan code
        interval_range_opt: Tuple of (min, max) intervals
        optimization_iterations: Number of iterations per grid
        risk_aversion: MVO risk aversion parameter (higher = more conservative)
        max_turnover: Maximum relative change for each grid's allocation (0.20 = ±20%)

    Returns:
        Dict with keys:
            'allocations': grid_id -> {interval: weight}
            'optimized_acres': grid_id -> acres (after MVO)
            'initial_acres': grid_id -> acres (before MVO)
            'analog_roi_correlation': DataFrame of correlations during analog years
            'stage1_stats': Per-grid optimization stats
    """
    if not analog_years or len(weather_grids) == 0:
        return None

    # === STAGE 1: Interval Optimization ===
    allocations = {}
    stage1_stats = {}

    for gid in weather_grids:
        grid_acres = weather_acres.get(gid, 1000)

        best_alloc, best_roi, tested = run_analog_year_optimization(
            session=session,
            grid_id=gid,
            analog_years=analog_years,
            plan_code=plan_code,
            productivity_factor=productivity_factor,
            acres=grid_acres,
            intended_use=intended_use,
            coverage_level=coverage_level,
            iterations=optimization_iterations,
            interval_range_opt=interval_range_opt
        )

        allocations[gid] = best_alloc
        stage1_stats[gid] = {'roi': best_roi, 'tested': tested}

    # === STAGE 2: MVO Acre Optimization ===
    # Build ROI series for each grid during ANALOG YEARS ONLY

    analog_roi_data = []

    for gid in weather_grids:
        allocation = allocations.get(gid, {})
        if not allocation:
            continue

        try:
            subsidy = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
            county_base_value = load_county_base_value(session, gid)
            current_rate_year = get_current_rate_year(session)
            premium_rates = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]

            dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)

            all_indices_df = load_all_indices(session, gid)
            # Filter to ONLY analog years
            all_indices_df = all_indices_df[all_indices_df['YEAR'].isin(analog_years)]

            for year in analog_years:
                year_data = all_indices_df[all_indices_df['YEAR'] == year]
                if year_data.empty:
                    continue

                year_indemnity = 0
                year_premium = 0

                for interval, pct in allocation.items():
                    if pct == 0:
                        continue

                    index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                    index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

                    premium_rate = premium_rates.get(interval, 0)
                    # Normalize to 1 acre for correlation calculation
                    interval_protection = int(round_half_up(dollar_protection * 1 * pct, 0))
                    total_prem = int(round_half_up(interval_protection * premium_rate, 0))
                    prem_subsidy = int(round_half_up(total_prem * subsidy, 0))
                    producer_premium = total_prem - prem_subsidy

                    trigger = coverage_level * 100
                    shortfall = max(0, (trigger - index_value) / trigger)
                    raw_indemnity = shortfall * interval_protection
                    # Convert to int immediately to ensure exact integer arithmetic when summing
                    indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                    year_indemnity += indemnity
                    year_premium += producer_premium

                # Both indemnity and premium are already exact integer sums

                roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                analog_roi_data.append({'year': year, 'grid': gid, 'roi': roi})

        except Exception as e:
            continue

    if len(analog_roi_data) == 0:
        # Fallback: return Stage 1 results with original acres
        return {
            'allocations': allocations,
            'optimized_acres': weather_acres.copy(),
            'initial_acres': weather_acres.copy(),
            'analog_roi_correlation': pd.DataFrame(),
            'stage1_stats': stage1_stats
        }

    analog_roi_df = pd.DataFrame(analog_roi_data)

    # Build correlation matrix from analog years
    pivot_df = analog_roi_df.pivot_table(
        values='roi',
        index='year',
        columns='grid'
    )

    mean_rois = pivot_df.mean()
    cov_matrix = pivot_df.cov()
    analog_roi_correlation = pivot_df.corr()

    # MVO Optimization
    grid_list = [gid for gid in weather_grids if gid in mean_rois.index]

    if len(grid_list) < 2:
        # Not enough grids for MVO
        return {
            'allocations': allocations,
            'optimized_acres': weather_acres.copy(),
            'initial_acres': weather_acres.copy(),
            'analog_roi_correlation': analog_roi_correlation,
            'stage1_stats': stage1_stats
        }

    n = len(grid_list)
    means = np.array([mean_rois.get(gid, 0) for gid in grid_list])

    cov = np.zeros((n, n))
    for i, gi in enumerate(grid_list):
        for j, gj in enumerate(grid_list):
            if gi in cov_matrix.index and gj in cov_matrix.columns:
                cov[i, j] = cov_matrix.loc[gi, gj]

    # Initial weights from input acres
    total_acres = sum(weather_acres.get(gid, 0) for gid in grid_list)
    if total_acres == 0:
        total_acres = len(grid_list) * 1000  # Default
    initial_weights = np.array([weather_acres.get(gid, 0) / total_acres for gid in grid_list])

    def neg_utility(weights):
        portfolio_return = np.dot(weights, means)
        portfolio_variance = np.dot(weights, np.dot(cov, weights))
        utility = portfolio_return - risk_aversion * portfolio_variance
        return -utility

    from scipy.optimize import minimize

    constraints = [
        {'type': 'eq', 'fun': lambda w: np.sum(w) - 1.0}
    ]

    # Bounds: ±max_turnover relative to initial weights
    bounds = []
    for i in range(n):
        w_base = initial_weights[i]
        lower = max(0.0, w_base * (1 - max_turnover))
        upper = min(1.0, w_base * (1 + max_turnover))
        bounds.append((lower, upper))

    result = minimize(
        neg_utility,
        initial_weights,
        method='SLSQP',
        bounds=bounds,
        constraints=constraints,
        options={'maxiter': 1000, 'ftol': 1e-9}
    )

    if result.success:
        optimal_weights = result.x
    else:
        optimal_weights = initial_weights

    # Convert weights back to acres
    optimized_acres = {}
    for i, gid in enumerate(grid_list):
        optimized_acres[gid] = optimal_weights[i] * total_acres

    # Include any grids that weren't in MVO
    for gid in weather_grids:
        if gid not in optimized_acres:
            optimized_acres[gid] = weather_acres.get(gid, 0)

    return {
        'allocations': allocations,
        'optimized_acres': optimized_acres,
        'initial_acres': weather_acres.copy(),
        'analog_roi_correlation': analog_roi_correlation,
        'stage1_stats': stage1_stats
    }


# =============================================================================
# === ACRE OPTIMIZATION FUNCTIONS (Two-Stage Optimization) ===
# =============================================================================

def calculate_yearly_roi_for_grid(
    session, grid_id, year, allocation, coverage_level,
    productivity_factor, intended_use, plan_code, acres=1
):
    """
    Calculate ROI for a single grid in a single year.
    Returns (roi, indemnity, premium) as normalized per-acre values.
    """
    try:
        county_base_value = load_county_base_value(session, grid_id)
        all_indices_df = load_all_indices(session, grid_id)
        year_data = all_indices_df[all_indices_df['YEAR'] == year]

        if year_data.empty:
            return 0, 0, 0

        current_rate_year = get_current_rate_year(session)
        premium_rates = load_premium_rates(session, grid_id, intended_use, [coverage_level], current_rate_year)[coverage_level]
        subsidy = load_subsidies(session, plan_code, [coverage_level])[coverage_level]

        dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
        total_protection = dollar_protection * acres

        total_indemnity = 0
        total_producer_premium = 0

        for interval, pct in allocation.items():
            if pct == 0:
                continue

            index_row = year_data[year_data['INTERVAL_NAME'] == interval]
            index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

            premium_rate = premium_rates.get(interval, 0)
            interval_protection = int(round_half_up(total_protection * pct, 0))
            total_premium = int(round_half_up(interval_protection * premium_rate, 0))
            premium_subsidy = int(round_half_up(total_premium * subsidy, 0))
            producer_premium = total_premium - premium_subsidy

            trigger = coverage_level * 100
            shortfall_pct = max(0, (trigger - index_value) / trigger)
            raw_indemnity = shortfall_pct * interval_protection
            # Convert to int immediately to ensure exact integer arithmetic when summing
            indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

            total_indemnity += indemnity
            total_producer_premium += producer_premium

        # Both indemnity and premium are already exact integer sums

        roi = (total_indemnity - total_producer_premium) / total_producer_premium if total_producer_premium > 0 else 0

        return roi, total_indemnity, total_producer_premium

    except Exception as e:
        return 0, 0, 0


def calculate_average_interval_weights(grid_results):
    """
    Average the best interval allocation across all grids.
    Returns numpy array of weights for each of the 11 intervals.
    """
    all_allocations = []
    for gid, data in grid_results.items():
        alloc = data['best_strategy']['allocation']
        # Convert to array format
        weights = np.array([alloc.get(interval, 0) for interval in INTERVAL_ORDER_11])
        all_allocations.append(weights)

    if len(all_allocations) == 0:
        return np.zeros(11)

    return np.mean(all_allocations, axis=0)


def calculate_annual_premium_cost(
    session, selected_grids, grid_acres, grid_results,
    productivity_factor, intended_use, plan_code
):
    """
    Calculate total annual premium cost using current rates.
    Returns: (total_cost, grid_breakdown_dict)
    """
    total_cost = 0
    grid_breakdown = {}

    try:
        current_rate_year = get_current_rate_year(session)

        for gid in selected_grids:
            if gid not in grid_results:
                continue

            acres = grid_acres.get(gid, 0)
            if acres <= 0:
                continue

            best_strategy = grid_results[gid]['best_strategy']
            allocation = best_strategy['allocation']
            coverage_level = best_strategy['coverage_level']

            county_base_value = load_county_base_value(session, gid)
            premium_rates = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
            subsidy = load_subsidies(session, plan_code, [coverage_level])[coverage_level]

            dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
            total_protection = dollar_protection * acres

            grid_premium = 0
            for interval, pct in allocation.items():
                if pct == 0:
                    continue
                premium_rate = premium_rates.get(interval, 0)
                interval_protection = int(round_half_up(total_protection * pct, 0))
                total_premium = int(round_half_up(interval_protection * premium_rate, 0))
                premium_subsidy = int(round_half_up(total_premium * subsidy, 0))
                producer_premium = total_premium - premium_subsidy
                grid_premium += producer_premium

            grid_breakdown[gid] = grid_premium
            total_cost += grid_premium

    except Exception as e:
        pass

    return total_cost, grid_breakdown


def apply_budget_constraint(grid_acres, total_cost, budget_limit, allow_scale_up=False):
    """
    Scale acres proportionally to fit budget.

    Args:
        grid_acres: Dict of {grid_id: acres}
        total_cost: Current total premium cost
        budget_limit: Maximum budget allowed
        allow_scale_up: If True, scale up acres when under budget to utilize full budget

    Returns: (scaled_grid_acres_dict, scale_factor)
    """
    if total_cost == 0:
        return grid_acres.copy(), 1.0

    if total_cost > budget_limit:
        # Over budget - scale DOWN (with 0.05% buffer to ensure we stay under)
        scale_factor = (budget_limit * 0.9995) / total_cost
        scaled_acres = {gid: acres * scale_factor for gid, acres in grid_acres.items()}
        return scaled_acres, scale_factor
    elif allow_scale_up and total_cost < budget_limit:
        # Under budget and scale-up enabled - scale UP to fill budget (with 0.05% buffer)
        scale_factor = (budget_limit * 0.9995) / total_cost
        scaled_acres = {gid: acres * scale_factor for gid, acres in grid_acres.items()}
        return scaled_acres, scale_factor
    else:
        # Within budget (or under budget but scale-up disabled)
        return grid_acres.copy(), 1.0


def optimize_grid_allocation(
    base_data_df, grid_results, initial_acres_per_grid,
    annual_budget, session, productivity_factor, intended_use, plan_code,
    selected_grids, risk_aversion=1.0, max_turnover=0.20, allow_scale_up=False
):
    """
    Two-stage optimization with proper turnover constraint handling:

    Stage 1: Budget Scaling - Scale all grids proportionally to meet budget
    Stage 2: MVO Rebalancing - Optimize within ±max_turnover of scaled baseline

    The key insight: turnover bounds apply to the BUDGET-SCALED baseline, not original acres.
    This ensures budget compliance while preserving the user's portfolio distribution intent.

    Args:
        max_turnover: Maximum relative change allowed for each grid's allocation (0.20 = ±20%)
                      Applied AFTER budget scaling, relative to the scaled baseline.

    Returns: (optimized_acres_dict, roi_correlation_df, optimization_info)
    """
    optimization_info = {
        'initial_total_acres': 0,
        'budget_scaled_acres': 0,
        'budget_scale_factor': 1.0,
        'mvo_applied': False
    }

    try:
        # Build ROI correlation matrix from historical data
        pivot_df = base_data_df.pivot_table(
            values='roi',
            index='year',
            columns='grid'
        )

        # Calculate mean ROI and covariance for each grid
        mean_rois = pivot_df.mean()
        cov_matrix = pivot_df.cov()

        # Normalize to correlation for display
        roi_correlation = pivot_df.corr()

        n_grids = len(selected_grids)

        if n_grids == 0:
            return initial_acres_per_grid.copy(), pd.DataFrame(), optimization_info

        # Calculate cost per acre for each grid
        cost_per_acre = {}
        for gid in selected_grids:
            if gid not in grid_results:
                continue

            # Calculate cost for 1 acre
            test_acres = {g: 0 for g in selected_grids}
            test_acres[gid] = 1
            cost, _ = calculate_annual_premium_cost(
                session, [gid], test_acres, grid_results,
                productivity_factor, intended_use, plan_code
            )
            cost_per_acre[gid] = cost

        # === STAGE 1: Budget Scaling ===
        # Scale all grids proportionally to meet the budget constraint FIRST
        total_initial_acres = sum(initial_acres_per_grid.values())
        optimization_info['initial_total_acres'] = total_initial_acres

        # Calculate cost at initial allocation
        initial_cost, _ = calculate_annual_premium_cost(
            session, selected_grids, initial_acres_per_grid, grid_results,
            productivity_factor, intended_use, plan_code
        )

        # Determine budget scale factor
        if initial_cost > annual_budget:
            # Over budget - scale DOWN proportionally (with 0.05% buffer)
            budget_scale = (annual_budget * 0.9995) / initial_cost
        elif allow_scale_up and initial_cost < annual_budget:
            # Under budget and scale-up enabled - scale UP to fill budget
            budget_scale = (annual_budget * 0.9995) / initial_cost
        else:
            # Within budget and scale-up disabled - keep as is
            budget_scale = 1.0

        optimization_info['budget_scale_factor'] = budget_scale

        # Calculate budget-adjusted baseline (proportional scaling)
        budget_adjusted_baseline = {
            gid: acres * budget_scale
            for gid, acres in initial_acres_per_grid.items()
        }
        budget_adjusted_total = total_initial_acres * budget_scale
        optimization_info['budget_scaled_acres'] = budget_adjusted_total

        # === STAGE 2: MVO Rebalancing ===
        # MVO optimizes weights, but bounds are relative to BUDGET-ADJUSTED baseline

        # Convert grid indices to ordered list
        grid_list = [gid for gid in selected_grids if gid in grid_results and gid in mean_rois.index]

        if len(grid_list) == 0:
            return budget_adjusted_baseline, roi_correlation, optimization_info

        # Get mean ROIs and covariance submatrix for our grids
        means = np.array([mean_rois.get(gid, 0) for gid in grid_list])

        # Build covariance matrix
        n = len(grid_list)
        cov = np.zeros((n, n))
        for i, gi in enumerate(grid_list):
            for j, gj in enumerate(grid_list):
                if gi in cov_matrix.index and gj in cov_matrix.columns:
                    cov[i, j] = cov_matrix.loc[gi, gj]

        # Cost per acre array
        costs = np.array([cost_per_acre.get(gid, 1) for gid in grid_list])

        # Initial weights from BUDGET-ADJUSTED baseline (not original!)
        initial_weights = np.array([
            budget_adjusted_baseline.get(gid, 0) / budget_adjusted_total
            for gid in grid_list
        ]) if budget_adjusted_total > 0 else np.ones(n) / n

        def neg_utility(weights):
            """Negative utility function (we minimize, so negate for maximization)"""
            portfolio_return = np.dot(weights, means)
            portfolio_variance = np.dot(weights, np.dot(cov, weights))
            utility = portfolio_return - risk_aversion * portfolio_variance
            return -utility

        def budget_constraint(weights):
            """Budget constraint: total cost <= budget"""
            acres = weights * budget_adjusted_total
            total_cost = np.dot(acres, costs)
            return annual_budget - total_cost

        # Constraints
        constraints = [
            {'type': 'eq', 'fun': lambda w: np.sum(w) - 1.0},  # Weights sum to 1
            {'type': 'ineq', 'fun': budget_constraint}  # Budget constraint
        ]

        # Bounds: ±max_turnover relative to BUDGET-ADJUSTED weights
        bounds = []
        for i in range(n):
            w_base = initial_weights[i]
            lower = max(0.0, w_base * (1 - max_turnover))
            upper = min(1.0, w_base * (1 + max_turnover))
            bounds.append((lower, upper))

        # Optimize
        result = minimize(
            neg_utility,
            initial_weights,
            method='SLSQP',
            bounds=bounds,
            constraints=constraints,
            options={'maxiter': 1000, 'ftol': 1e-9}
        )

        if result.success:
            optimal_weights = result.x
            optimization_info['mvo_applied'] = True
        else:
            # Fallback to budget-adjusted proportional allocation
            optimal_weights = initial_weights

        # Convert weights to acres using budget_adjusted_total
        optimized_acres = {}
        for i, gid in enumerate(grid_list):
            optimized_acres[gid] = optimal_weights[i] * budget_adjusted_total

        # Add any grids that weren't in the optimization with their budget-scaled value
        for gid in selected_grids:
            if gid not in optimized_acres:
                optimized_acres[gid] = budget_adjusted_baseline.get(gid, 0)

        # === STAGE 3: Post-MVO Budget Fill-up (when auto-fill enabled) ===
        # After MVO rebalances, the total cost may be under budget if allocation shifted
        # from higher-premium grids to lower-premium grids. Scale up to hit 99.95% of budget.
        if allow_scale_up:
            # Calculate current cost after MVO rebalancing
            mvo_cost, _ = calculate_annual_premium_cost(
                session, selected_grids, optimized_acres, grid_results,
                productivity_factor, intended_use, plan_code
            )

            target_budget = annual_budget * 0.9995  # Target 99.95% utilization

            if mvo_cost > 0 and mvo_cost < target_budget * 0.999:  # Under-utilized (below 99.9%)
                # Calculate scale factor to hit target budget
                scale_factor = target_budget / mvo_cost

                # Define max acres per grid (e.g., 3x initial allocation to prevent extreme concentration)
                max_acres_per_grid = {
                    gid: initial_acres_per_grid.get(gid, 0) * 3.0
                    for gid in selected_grids
                }

                # Apply scaling with iterative redistribution to handle max constraints
                scaled_acres = {gid: acres * scale_factor for gid, acres in optimized_acres.items()}

                # Cap at max acres
                for gid in scaled_acres:
                    if max_acres_per_grid.get(gid, float('inf')) > 0:
                        scaled_acres[gid] = min(scaled_acres[gid], max_acres_per_grid[gid])

                # Iterative redistribution: if some grids hit max, redistribute remaining budget to others
                for iteration in range(10):  # Max iterations to prevent infinite loop
                    new_cost, _ = calculate_annual_premium_cost(
                        session, selected_grids, scaled_acres, grid_results,
                        productivity_factor, intended_use, plan_code
                    )

                    if new_cost >= annual_budget * 0.9990:  # Close enough (99.90%)
                        break

                    # Find grids not at max that can absorb more acres
                    can_scale_grids = [
                        gid for gid in selected_grids
                        if scaled_acres.get(gid, 0) < max_acres_per_grid.get(gid, float('inf')) * 0.99
                    ]

                    if not can_scale_grids:
                        break  # All grids at max

                    # Calculate additional scale factor for remaining budget
                    remaining_budget = target_budget - new_cost
                    if remaining_budget <= 0 or new_cost <= 0:
                        break

                    # Conservative step: scale up non-maxed grids proportionally
                    additional_scale = 1 + (remaining_budget / new_cost) * 0.5

                    for gid in can_scale_grids:
                        proposed = scaled_acres[gid] * additional_scale
                        scaled_acres[gid] = min(proposed, max_acres_per_grid.get(gid, proposed))

                optimized_acres = scaled_acres
                optimization_info['stage3_applied'] = True
                optimization_info['stage3_scale_factor'] = scale_factor

        return optimized_acres, roi_correlation, optimization_info

    except Exception as e:
        # Return initial allocation on error
        return initial_acres_per_grid.copy(), pd.DataFrame(), optimization_info


def optimize_without_budget(
    base_data_df, grid_results, max_total_acres,
    selected_grids, risk_aversion=1.0, max_turnover=0.20, initial_acres_per_grid=None
):
    """
    Optimize acre distribution for pure risk-adjusted return without budget constraint.
    Uses mean-variance optimization (Markowitz).

    Args:
        max_turnover: Maximum relative change allowed for each grid's allocation (0.20 = ±20%)
        initial_acres_per_grid: Optional baseline allocation for turnover bounds. If provided,
                                turnover bounds are relative to this allocation's weights.
                                If None, uses equal weights as baseline.

    Returns: optimized_acres_dict
    """
    try:
        # Build ROI data from historical
        pivot_df = base_data_df.pivot_table(
            values='roi',
            index='year',
            columns='grid'
        )

        mean_rois = pivot_df.mean()
        cov_matrix = pivot_df.cov()

        grid_list = [gid for gid in selected_grids if gid in grid_results and gid in mean_rois.index]

        if len(grid_list) == 0:
            # Fallback to uniform
            return {gid: max_total_acres / len(selected_grids) for gid in selected_grids}

        n = len(grid_list)
        means = np.array([mean_rois.get(gid, 0) for gid in grid_list])

        cov = np.zeros((n, n))
        for i, gi in enumerate(grid_list):
            for j, gj in enumerate(grid_list):
                if gi in cov_matrix.index and gj in cov_matrix.columns:
                    cov[i, j] = cov_matrix.loc[gi, gj]

        # Initial weights: use provided baseline if available, else equal distribution
        if initial_acres_per_grid is not None:
            total_baseline = sum(initial_acres_per_grid.get(gid, 0) for gid in grid_list)
            if total_baseline > 0:
                initial_weights = np.array([
                    initial_acres_per_grid.get(gid, 0) / total_baseline
                    for gid in grid_list
                ])
            else:
                initial_weights = np.ones(n) / n
        else:
            initial_weights = np.ones(n) / n

        def neg_utility(weights):
            portfolio_return = np.dot(weights, means)
            portfolio_variance = np.dot(weights, np.dot(cov, weights))
            utility = portfolio_return - risk_aversion * portfolio_variance
            return -utility

        constraints = [
            {'type': 'eq', 'fun': lambda w: np.sum(w) - 1.0}
        ]

        # Bounds: constrain each weight to ±max_turnover relative change from baseline
        bounds = []
        for i in range(n):
            w_base = initial_weights[i]
            lower = max(0.0, w_base * (1 - max_turnover))
            upper = min(1.0, w_base * (1 + max_turnover))
            bounds.append((lower, upper))

        result = minimize(
            neg_utility,
            initial_weights,
            method='SLSQP',
            bounds=bounds,
            constraints=constraints,
            options={'maxiter': 1000, 'ftol': 1e-9}
        )

        if result.success:
            optimal_weights = result.x
        else:
            optimal_weights = initial_weights

        optimized_acres = {}
        for i, gid in enumerate(grid_list):
            optimized_acres[gid] = optimal_weights[i] * max_total_acres

        for gid in selected_grids:
            if gid not in optimized_acres:
                optimized_acres[gid] = 0

        return optimized_acres

    except Exception as e:
        # Fallback to uniform
        return {gid: max_total_acres / len(selected_grids) for gid in selected_grids}


def generate_strategy_report_docx(
    comparison_df, stress_scenario, start_year, end_year, coverage_level,
    productivity_factor, intended_use, analog_years_count, weather_config
):
    """
    Generate a Word document with strategy comparison report.
    Returns BytesIO buffer with the document.
    """
    doc = Document()

    # Set landscape orientation
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # Title
    title = doc.add_heading('PRF Strategy Comparison Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with metadata
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").bold = False
    subtitle.add_run(f"Scenario: {stress_scenario}\n")
    subtitle.add_run(f"Year Range: {start_year} - {end_year}")

    doc.add_paragraph()

    # --- Performance Comparison Table ---
    doc.add_heading('Performance Comparison', level=1)

    # Create table from comparison_df
    table = doc.add_table(rows=1, cols=len(comparison_df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(comparison_df.columns):
        hdr_cells[idx].text = col
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for _, row in comparison_df.iterrows():
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            row_cells[idx].text = str(val)

    doc.add_paragraph()

    # --- Helper function to add allocation table ---
    def add_allocation_table(doc, title, results_data, results_key):
        """Add an allocation table for a strategy. Returns (coverage_count, dropped_grids)."""
        if results_key not in st.session_state or not st.session_state[results_key]:
            doc.add_heading(f'{title} - No data available', level=1)
            return 0, []

        data = st.session_state[results_key]
        allocations = data.get('allocations', {})
        acres = data.get('acres', {})
        grids = data.get('grids', [])

        if not grids:
            doc.add_heading(f'{title} - No grids', level=1)
            return 0, []

        # Filter to only grids with acres > 0
        active_grids = [gid for gid in grids if acres.get(gid, 0) > 0]
        dropped_grids = [gid for gid in grids if acres.get(gid, 0) <= 0]

        if not active_grids:
            doc.add_heading(f'{title} - All grids have 0 acres', level=1)
            return 0, dropped_grids

        doc.add_heading(title, level=1)

        # Create table: Grid | Jan-Feb | Feb-Mar | ... | Nov-Dec | Acres
        headers = ['Grid'] + INTERVAL_ORDER_11 + ['Acres']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        for idx, col in enumerate(headers):
            hdr_cells[idx].text = col
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Data rows - only include active grids (acres > 0)
        total_acres = 0
        for gid in active_grids:
            row_cells = table.add_row().cells
            row_cells[0].text = str(gid)
            alloc = allocations.get(gid, {})
            for idx, interval in enumerate(INTERVAL_ORDER_11):
                pct = alloc.get(interval, 0)
                row_cells[idx + 1].text = f"{pct*100:.0f}%" if pct > 0 else "-"
            grid_acres = acres.get(gid, 0)
            row_cells[len(INTERVAL_ORDER_11) + 1].text = f"{grid_acres:,.0f}"
            total_acres += grid_acres

        # Add TOTAL row
        total_row_cells = table.add_row().cells
        total_row_cells[0].text = "TOTAL"
        for paragraph in total_row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for idx in range(len(INTERVAL_ORDER_11)):
            total_row_cells[idx + 1].text = "--"
        total_row_cells[len(INTERVAL_ORDER_11) + 1].text = f"{total_acres:,.0f}"
        for paragraph in total_row_cells[len(INTERVAL_ORDER_11) + 1].paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Add COVERAGE row
        coverage_row_cells = table.add_row().cells
        coverage_row_cells[0].text = "COVERAGE"
        for paragraph in coverage_row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Check each interval for coverage (only considering active grids)
        coverage_count = 0
        for idx, interval in enumerate(INTERVAL_ORDER_11):
            has_coverage = False
            for gid in active_grids:
                alloc = allocations.get(gid, {})
                grid_acres = acres.get(gid, 0)
                interval_pct = alloc.get(interval, 0)
                if grid_acres > 0 and interval_pct > 0:
                    has_coverage = True
                    break
            if has_coverage:
                coverage_row_cells[idx + 1].text = "X"
                coverage_count += 1
            else:
                coverage_row_cells[idx + 1].text = "--"

        # Show coverage count in Acres column
        coverage_row_cells[len(INTERVAL_ORDER_11) + 1].text = f"{coverage_count}/11"
        for paragraph in coverage_row_cells[len(INTERVAL_ORDER_11) + 1].paragraphs:
            for run in paragraph.runs:
                run.bold = True

        doc.add_paragraph()

        return coverage_count, dropped_grids

    # --- Add allocation tables for each strategy ---
    champ_coverage, champ_dropped = add_allocation_table(doc, 'Champion Allocation', st.session_state, 'champion_results')
    chall1_coverage, chall1_dropped = add_allocation_table(doc, 'Challenger 1 Allocation', st.session_state, 'challenger_results')
    chall2_coverage, chall2_dropped = add_allocation_table(doc, 'Challenger 2 Allocation', st.session_state, 'weather_challenger_results')
    chall3_coverage, chall3_dropped = add_allocation_table(doc, 'Challenger 3 Allocation', st.session_state, 'weather_challenger_3_results')

    # --- Footnotes Section ---
    doc.add_heading('Footnotes', level=1)

    footnotes = doc.add_paragraph()
    footnotes.add_run("Custom Range: ").bold = True
    footnotes.add_run(f"{start_year} - {end_year}\n")

    # Analog years criteria
    if weather_config:
        criteria_parts = []
        if weather_config.get('enso_regime', 'Any') != 'Any':
            criteria_parts.append(weather_config.get('enso_regime'))
        if weather_config.get('historical_context', 'Any') != 'Any':
            criteria_parts.append(weather_config.get('historical_context'))
        if weather_config.get('trajectory', 'Any') != 'Any':
            criteria_parts.append(weather_config.get('trajectory'))

        footnotes.add_run("Analog Years Criteria: ").bold = True
        footnotes.add_run(f"{' + '.join(criteria_parts) if criteria_parts else 'All Conditions'}\n")

    if analog_years_count:
        footnotes.add_run("Number of Analog Years: ").bold = True
        footnotes.add_run(f"{analog_years_count}\n")

    # Coverage Summary
    footnotes.add_run("\nCoverage Summary:\n").bold = True
    footnotes.add_run(f"- Champion: {champ_coverage} of 11 intervals covered\n")
    footnotes.add_run(f"- Challenger 1: {chall1_coverage} of 11 intervals covered\n")
    footnotes.add_run(f"- Challenger 2: {chall2_coverage} of 11 intervals covered\n")
    footnotes.add_run(f"- Challenger 3: {chall3_coverage} of 11 intervals covered\n")

    # Dropped Grids section (0 acres) - always show all strategies
    footnotes.add_run("\nDropped Grids (0 acres):\n").bold = True

    def format_dropped_grids(dropped_list):
        """Format dropped grid list for footnotes with county info."""
        if not dropped_list:
            return "None"
        # Format each grid with its county info
        formatted = []
        for g in dropped_list:
            numeric_id = extract_numeric_grid_id(g)
            county = extract_county_from_grid_id(g)
            if county:
                formatted.append(f"{numeric_id} ({county} - TX)")
            else:
                formatted.append(str(numeric_id))
        return ', '.join(formatted)

    footnotes.add_run(f"  - Champion: {format_dropped_grids(champ_dropped)}\n")
    footnotes.add_run(f"  - Challenger 1: {format_dropped_grids(chall1_dropped)}\n")
    footnotes.add_run(f"  - Challenger 2: {format_dropped_grids(chall2_dropped)}\n")
    footnotes.add_run(f"  - Challenger 3: {format_dropped_grids(chall3_dropped)}\n")

    footnotes.add_run("\nCoverage Level: ").bold = True
    footnotes.add_run(f"{coverage_level:.0%}\n")

    footnotes.add_run("Productivity Factor: ").bold = True
    footnotes.add_run(f"{productivity_factor:.2f}\n")

    footnotes.add_run("Intended Use: ").bold = True
    footnotes.add_run(f"{intended_use}\n")

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def render_allocation_inputs(key_prefix):
    """Creates the 11-row data editor for interval allocation."""
    st.subheader("Interval Allocation")

    # Check if there's preset allocation data for this key_prefix
    preset_key = f"{key_prefix}_preset_allocation"
    if preset_key in st.session_state:
        preset_alloc = st.session_state[preset_key]
        # Convert from decimal to percentage format - round to whole numbers
        alloc_data = {interval: round(preset_alloc.get(interval, 0.0) * 100) for interval in INTERVAL_ORDER_11}
    else:
        # Default allocation
        alloc_data = {
            'Jan-Feb': 50, 'Feb-Mar': 0, 'Mar-Apr': 50, 'Apr-May': 0,
            'May-Jun': 0, 'Jun-Jul': 0, 'Jul-Aug': 0, 'Aug-Sep': 0,
            'Sep-Oct': 0, 'Oct-Nov': 0, 'Nov-Dec': 0
        }

    df_alloc = pd.DataFrame(list(alloc_data.items()), columns=['Interval', 'Percent of Value'])

    st.caption("Whole numbers only (1% increments). Each interval: 0% OR 10-50%. Total must equal 100%. No adjacent intervals (except Nov-Dec/Jan-Feb wrap).")

    edited_df = st.data_editor(
        df_alloc,
        key=f"{key_prefix}_alloc_editor",
        num_rows="fixed",
        use_container_width=True,
        column_config={
            "Interval": st.column_config.TextColumn("Interval", disabled=True, width="medium"),
            "Percent of Value": st.column_config.NumberColumn("Percent of Value (%)", min_value=0, max_value=50, step=1, format="%d%%")
        }
    )

    # --- Validation ---
    alloc_dict = pd.Series(edited_df['Percent of Value'].values, index=edited_df['Interval']).to_dict()

    # Round to integers to ensure whole numbers
    alloc_dict = {k: round(v) for k, v in alloc_dict.items()}

    total_pct = sum(alloc_dict.values())
    max_pct = max(alloc_dict.values())

    is_valid = True

    # Check for whole numbers
    for interval, pct in alloc_dict.items():
        if pct != int(pct):
            st.error(f"All allocations must be whole numbers. {interval} has {pct}%")
            is_valid = False
            break

    # Check for 0% or 10-50% range
    for interval, pct in alloc_dict.items():
        if pct > 0 and pct < 10:
            st.error(f"Each interval must be 0% OR between 10-50%. {interval} has {pct}% (below 10% minimum)")
            is_valid = False
            break

    if abs(total_pct - 100) > 0.01:
        st.error(f"Allocation must total 100%. Current total: {total_pct:.0f}%")
        is_valid = False

    if max_pct > 50:
        st.error(f"No interval can exceed 50%.")
        is_valid = False

    for i in range(len(INTERVAL_ORDER_11) - 1):  # Stops before Nov-Dec
        if alloc_dict[INTERVAL_ORDER_11[i]] > 0 and alloc_dict[INTERVAL_ORDER_11[i+1]] > 0:
            st.error(f"Cannot allocate to adjacent intervals: {INTERVAL_ORDER_11[i]} and {INTERVAL_ORDER_11[i+1]}")
            is_valid = False
            break

    if is_valid:
        st.success(f"Valid. Total: {total_pct:.0f}%")

    alloc_dict_float = {k: v / 100.0 for k, v in alloc_dict.items()}

    return alloc_dict_float, is_valid


# =============================================================================
# === 3. TAB 1: TFC DECISION SUPPORT TOOL (S2) ===
# =============================================================================
def render_tab2(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code):
    st.subheader("Parameters")
    
    col1, col2 = st.columns(2)
    sample_year = col1.selectbox("Historical Rainfall Year", list(range(1948, 2026)), index=77, key="s2_year")
    coverage_level = col2.selectbox("Coverage Level", [0.70, 0.75, 0.80, 0.85, 0.90], index=2, format_func=lambda x: f"{x:.0%}", key="s2_coverage")
    
    with st.expander("Step 2: Set Interval Allocations", expanded=True):
        pct_of_value_alloc, is_valid = render_allocation_inputs("s2")
    
    st.divider()

    if 'tab2_run' not in st.session_state:
        st.session_state.tab2_run = False

    if st.button("Run Calculation", key="s2_run_button", disabled=not is_valid):
        st.session_state.tab2_run = True
        try:
            with st.spinner("Calculating..."):
                # --- 1. FETCH DATA ---
                coverage_level_string = f"{coverage_level:.0%}"
                subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                county_base_value = load_county_base_value(session, grid_id)
                current_rate_year = get_current_rate_year(session)
                premium_rates_df = load_premium_rates(session, grid_id, intended_use, [coverage_level], current_rate_year)[coverage_level]
                
                # Extract numeric grid ID for rainfall data query
                numeric_grid_id = extract_numeric_grid_id(grid_id)
                actuals_query = f"""
                    SELECT INTERVAL_NAME, INDEX_VALUE 
                    FROM RAIN_INDEX_PLATINUM_ENHANCED 
                    WHERE GRID_ID = {numeric_grid_id} AND YEAR = {sample_year}
                """
                actuals_df = session.sql(actuals_query).to_pandas().set_index('INTERVAL_NAME')
                actuals_df['INDEX_VALUE'] = pd.to_numeric(actuals_df['INDEX_VALUE'], errors='coerce')

                dollar_amount_of_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                total_policy_protection = dollar_amount_of_protection * total_insured_acres

                roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                roi_df['Percent of Value'] = roi_df.index.map(pct_of_value_alloc)
                roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce')
                roi_df['Premium Rate Per $100'] = roi_df['PREMIUM_RATE'] * 100
                roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                trigger_level = coverage_level * 100
                shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)
                roi_df['ROI %'] = np.where(roi_df['Producer Premium'] > 0, (roi_df['Estimated Indemnity'] - roi_df['Producer Premium']) / roi_df['Producer Premium'], 0)

            # === 3. SAVE RESULTS TO SESSION STATE ===
            st.session_state.tab2_results = {
                "roi_df": roi_df, "grid_id": grid_id, "sample_year": sample_year, "current_rate_year": current_rate_year,
                "intended_use": intended_use, "coverage_level": coverage_level, "productivity_factor": productivity_factor,
                "total_insured_acres": total_insured_acres, "county_base_value": county_base_value,
                "dollar_amount_of_protection": dollar_amount_of_protection, "total_policy_protection": total_policy_protection,
                "subsidy_percent": subsidy_percent
            }
            # Clear other tab results
            st.session_state.tab4_results = None
            st.session_state.tab5_results = None

        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.exception(e)
            st.session_state.tab2_results = None

    # === 4. DISPLAY RESULTS (if they exist) ===
    if 'tab2_results' in st.session_state and st.session_state.tab2_results:
        try:
            r = st.session_state.tab2_results
            st.header(f"ROI Calculation - Grid {r['grid_id']}, Year {r['sample_year']}")
            st.caption(f"Coverage: {r['coverage_level']:.0%} | Productivity: {r['productivity_factor']:.0%} | Use: {r['intended_use']} | Acres: {r['total_insured_acres']:,}")
            
            c1, c2 = st.columns(2)
            with c1.container(border=True):
                st.subheader("Protection")
                st.text(f"Use: {r['intended_use']}")
                st.text(f"Coverage: {r['coverage_level']:.0%}")
                st.text(f"Productivity: {r['productivity_factor']:.0%}")
                st.text(f"Acres: {r['total_insured_acres']:,}")
            
            with c2.container(border=True):
                st.subheader("Policy")
                st.text(f"Base Value: ${r['county_base_value']:,.2f}")
                st.text(f"Protection: ${r['dollar_amount_of_protection']:,.2f}")
                st.text(f"Total Protection: ${r['total_policy_protection']:,.0f}")
                st.text(f"Subsidy: {r['subsidy_percent']:.1%}")

            st.subheader("Protection Table")
            
            # Add CSV download button
            csv_df = r['roi_df'].copy()
            csv_df['Percent of Value'] = csv_df['Percent of Value'] * 100
            csv_columns = ['Percent of Value', 'Policy Protection Per Unit', 'Premium Rate Per $100', 
                           'Total Premium', 'Premium Subsidy', 'Producer Premium', 
                           'Actual Index Value', 'Estimated Indemnity', 'ROI %']
            csv_export = csv_df[csv_columns].to_csv()
            
            st.download_button(
                label="📥 Export CSV",
                data=csv_export,
                file_name=f"protection_grid_{extract_numeric_grid_id(r['grid_id'])}_year_{r['sample_year']}.csv",
                mime="text/csv",
            )
            
            display_df = pd.DataFrame(index=r['roi_df'].index)
            display_df['% Value'] = r['roi_df']['Percent of Value'].apply(lambda x: f"{x*100:.0f}" if x > 0 else 'N/A')
            display_df['Protection'] = r['roi_df']['Policy Protection Per Unit'].apply(lambda x: f"${x:,.0f}" if x > 0 else 'N/A')
            display_df['Rate/$100'] = r['roi_df']['Premium Rate Per $100'].apply(lambda x: f"{x:.2f}" if x > 0 else 'N/A')
            display_df['Premium'] = r['roi_df']['Total Premium'].apply(lambda x: f"${x:,.0f}" if x > 0 else 'N/A')
            display_df['Subsidy'] = r['roi_df']['Premium Subsidy'].apply(lambda x: f"${x:,.0f}" if x > 0 else 'N/A')
            display_df['Producer'] = r['roi_df']['Producer Premium'].apply(lambda x: f"${x:,.0f}" if x > 0 else 'N/A')
            display_df['Index'] = r['roi_df'].apply(
                lambda row: 'N/A' if pd.isna(row['Actual Index Value']) 
                else (f"{row['Actual Index Value']:.1f}" if row['Percent of Value'] > 0 or row['Actual Index Value'] > 0 
                else 'N/A'),
                axis=1
            )
            display_df['Indemnity'] = r['roi_df']['Estimated Indemnity'].apply(
                lambda x: f"${x:,.0f}" if pd.notna(x) and x > 0 else 'N/A'
            )
            display_df['ROI %'] = r['roi_df'].apply(
                lambda row: 'N/A' if pd.isna(row['Estimated Indemnity']) or row['ROI %'] == 0 
                else f"{row['ROI %']:.2%}",
                axis=1
            )
            st.dataframe(display_df, use_container_width=True)

            # Totals
            total_producer_prem = r['roi_df']['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
            total_indemnity = r['roi_df']['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum(skipna=True)
            net_return = total_indemnity - total_producer_prem
            
            st.subheader("Totals")
            
            has_missing_data = r['roi_df']['Actual Index Value'].isna().any()
            if has_missing_data:
                st.info("ℹ️ Some intervals have incomplete data")
            
            c1, c2, c3 = st.columns(3)
            c1.metric("Producer Premium", f"${total_producer_prem:,.0f}")
            c2.metric("Total Indemnity", f"${total_indemnity:,.0f}")
            c3.metric("Net Return", f"${net_return:,.0f}")
            
            if total_producer_prem > 0:
                st.metric("ROI", f"{net_return / total_producer_prem:.2%}")
        except Exception as e:
            st.error(f"Error displaying results: {e}")
            st.session_state.tab2_results = None
    elif not st.session_state.tab2_run:
        st.info("Select parameters and click 'Run Calculation'")

# =============================================================================
# === 5. PORTFOLIO STRATEGY TAB (Champion vs Challenger) ===
# =============================================================================

def run_portfolio_backtest(
    session, selected_grids, grid_allocations, grid_acres,
    start_year, end_year, coverage_level, productivity_factor,
    intended_use, plan_code, scenario='All Years (except Current Year)'
):
    """
    Run a historical backtest for a portfolio of grids.
    Returns: (portfolio_results_df, grid_results_dict, metrics_dict)
    """
    grid_results = {}
    portfolio_yearly = {}

    for gid in selected_grids:
        try:
            subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
            county_base_value = load_county_base_value(session, gid)
            current_rate_year = get_current_rate_year(session)
            premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]

            dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
            total_protection = dollar_protection * grid_acres.get(gid, 0)

            all_indices_df = load_all_indices(session, gid)

            # Apply scenario filter
            filtered_df = filter_indices_by_scenario(all_indices_df, scenario, start_year, end_year)

            allocation = grid_allocations.get(gid, {})
            year_results = []

            for year in sorted(filtered_df['YEAR'].unique()):
                actuals_df = filtered_df[filtered_df['YEAR'] == year].set_index('INTERVAL_NAME')
                if actuals_df.empty:
                    continue

                year_indemnity = 0
                year_premium = 0

                for interval, pct in allocation.items():
                    if pct == 0:
                        continue

                    index_row = actuals_df.loc[interval] if interval in actuals_df.index else None
                    index_value = float(index_row['INDEX_VALUE']) if index_row is not None else 100

                    premium_rate = premium_rates_df.get(interval, 0)
                    interval_protection = int(round_half_up(total_protection * pct, 0))
                    total_premium = int(round_half_up(interval_protection * premium_rate, 0))
                    premium_subsidy = int(round_half_up(total_premium * subsidy_percent, 0))
                    producer_premium = total_premium - premium_subsidy

                    trigger = coverage_level * 100
                    shortfall_pct = max(0, (trigger - index_value) / trigger)
                    raw_indemnity = shortfall_pct * interval_protection
                    # Convert to int immediately to ensure exact integer arithmetic when summing
                    indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                    year_indemnity += indemnity
                    year_premium += producer_premium

                # Both indemnity and premium are already exact integer sums

                year_results.append({
                    'year': year,
                    'indemnity': year_indemnity,
                    'premium': year_premium,
                    'net': year_indemnity - year_premium,
                    'roi': (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                })

                # Aggregate to portfolio level
                if year not in portfolio_yearly:
                    portfolio_yearly[year] = {'indemnity': 0, 'premium': 0}
                portfolio_yearly[year]['indemnity'] += year_indemnity
                portfolio_yearly[year]['premium'] += year_premium

            grid_results[gid] = {
                'year_results': pd.DataFrame(year_results),
                'allocation': allocation,
                'acres': grid_acres.get(gid, 0)
            }

        except Exception as e:
            grid_results[gid] = {'error': str(e)}

    # Build portfolio results
    portfolio_rows = []
    for year, data in sorted(portfolio_yearly.items()):
        net = data['indemnity'] - data['premium']
        roi = net / data['premium'] if data['premium'] > 0 else 0
        portfolio_rows.append({
            'Year': year,
            'Total Indemnity': data['indemnity'],
            'Producer Premium': data['premium'],
            'Net Return': net,
            'ROI': roi
        })

    portfolio_df = pd.DataFrame(portfolio_rows)

    # Calculate aggregate metrics
    if len(portfolio_df) > 0:
        total_indemnity = portfolio_df['Total Indemnity'].sum()
        total_premium = portfolio_df['Producer Premium'].sum()
        cumulative_roi = (total_indemnity - total_premium) / total_premium if total_premium > 0 else 0
        avg_roi = portfolio_df['ROI'].mean()
        std_roi = portfolio_df['ROI'].std()
        risk_adj_return = cumulative_roi / std_roi if std_roi > 0 else 0
        profitable_years = (portfolio_df['Net Return'] > 0).sum()
        profitable_pct = profitable_years / len(portfolio_df) if len(portfolio_df) > 0 else 0
        years_tested = len(portfolio_df)
        avg_annual_premium = total_premium / years_tested if years_tested > 0 else 0

        metrics = {
            'cumulative_roi': cumulative_roi,
            'avg_roi': avg_roi,
            'std_roi': std_roi,
            'risk_adj_return': risk_adj_return,
            'total_indemnity': total_indemnity,
            'total_premium': total_premium,
            'total_profit': total_indemnity - total_premium,
            'profitable_pct': profitable_pct,
            'years_tested': years_tested,
            'avg_annual_premium': avg_annual_premium
        }
    else:
        metrics = {}

    return portfolio_df, grid_results, metrics


def generate_base_data_for_mvo(session, selected_grids, grid_results_with_allocations,
                                start_year, end_year, coverage_level, productivity_factor,
                                intended_use, plan_code):
    """
    Generate the base_data_df required by optimize_grid_allocation.
    This creates historical ROI data for each grid using the optimized allocations.
    Returns: DataFrame with columns ['year', 'grid', 'roi']
    """
    rows = []

    for gid in selected_grids:
        if gid not in grid_results_with_allocations:
            continue

        allocation = grid_results_with_allocations[gid].get('allocation', {})
        if not allocation:
            continue

        try:
            subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
            county_base_value = load_county_base_value(session, gid)
            current_rate_year = get_current_rate_year(session)
            premium_rates = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]

            dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
            # Use 1 acre for normalized ROI calculation
            total_protection = dollar_protection * 1

            all_indices_df = load_all_indices(session, gid)
            all_indices_df = all_indices_df[
                (all_indices_df['YEAR'] >= start_year) &
                (all_indices_df['YEAR'] <= end_year)
            ]

            for year in range(start_year, end_year + 1):
                year_data = all_indices_df[all_indices_df['YEAR'] == year]
                if year_data.empty:
                    continue

                year_indemnity = 0
                year_premium = 0

                for interval, pct in allocation.items():
                    if pct == 0:
                        continue

                    index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                    index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

                    premium_rate = premium_rates.get(interval, 0)
                    interval_protection = int(round_half_up(total_protection * pct, 0))
                    total_prem = int(round_half_up(interval_protection * premium_rate, 0))
                    prem_subsidy = int(round_half_up(total_prem * subsidy_percent, 0))
                    producer_premium = total_prem - prem_subsidy

                    trigger = coverage_level * 100
                    shortfall_pct = max(0, (trigger - index_value) / trigger)
                    raw_indemnity = shortfall_pct * interval_protection
                    # Convert to int immediately to ensure exact integer arithmetic when summing
                    indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                    year_indemnity += indemnity
                    year_premium += producer_premium

                # Both indemnity and premium are already exact integer sums

                roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                rows.append({'year': year, 'grid': gid, 'roi': roi})

        except Exception as e:
            continue

    return pd.DataFrame(rows)


def create_optimized_allocation_table(allocations_dict, grid_list, grid_acres=None,
                                       label="OPTIMIZED AVERAGE"):
    """
    Create a styled allocation table for Champion or Challenger.

    Args:
        allocations_dict: Dict mapping grid_id -> {interval: weight (0-1), ...}
        grid_list: List of grid IDs to include
        grid_acres: Optional dict of grid_id -> acres
        label: Label for the summary row (default "OPTIMIZED AVERAGE")

    Returns:
        Tuple of (styled DataFrame, raw DataFrame)
    """
    # Filter to only grids with acres > 0
    if grid_acres:
        active_grids = [gid for gid in grid_list if grid_acres.get(gid, 0) > 0]
    else:
        active_grids = grid_list

    rows = []
    total_coverage = {interval: 0 for interval in INTERVAL_ORDER_11}
    total_acres = 0

    for gid in active_grids:
        alloc = allocations_dict.get(gid, {})
        row = {'Grid': gid}
        row_sum = 0

        for interval in INTERVAL_ORDER_11:
            # Handle both decimal (0.25) and percentage (25) formats
            pct_raw = alloc.get(interval, 0)
            if isinstance(pct_raw, (int, float)):
                pct = pct_raw * 100 if pct_raw <= 1 else pct_raw
            else:
                pct = 0
            row_sum += pct
            total_coverage[interval] += pct
            row[interval] = f"{pct:.0f}%"  # Always show as percentage, even 0%

        row['Row Sum'] = f"{row_sum:.0f}%"

        # Add acres if provided
        acres = grid_acres.get(gid, 0) if grid_acres else 0
        total_acres += acres
        row['Acres'] = f"{acres:,.0f}"

        rows.append(row)

    # Add summary row (OPTIMIZED AVERAGE or PORTFOLIO AVERAGE)
    avg_row = {'Grid': label}
    avg_row_sum = 0
    grid_count = len(active_grids)

    for interval in INTERVAL_ORDER_11:
        avg_pct = total_coverage[interval] / grid_count if grid_count > 0 else 0
        avg_row_sum += avg_pct
        avg_row[interval] = f"{avg_pct:.0f}%"

    avg_row['Row Sum'] = f"{avg_row_sum:.0f}%"
    avg_row['Acres'] = f"{total_acres:,.0f}"
    rows.append(avg_row)

    df = pd.DataFrame(rows)

    # Style function for green gradient on allocation percentages
    def highlight_allocation_cell(val):
        if isinstance(val, str) and val.endswith('%'):
            try:
                pct = float(val.replace('%', ''))
                if pct >= 40:
                    return 'background-color: #2e7d32; color: white'  # Dark green
                elif pct >= 25:
                    return 'background-color: #4caf50; color: white'  # Medium green
                elif pct >= 10:
                    return 'background-color: #81c784'  # Light green
            except:
                pass
        return ''

    # Apply styling to interval columns
    styled = df.style.applymap(highlight_allocation_cell, subset=INTERVAL_ORDER_11)

    return styled, df


def create_change_analysis_table(champ_alloc, chall_alloc, champ_acres, chall_acres, grid_list):
    """
    Create a styled table showing changes between Champion and Challenger allocations,
    including acreage changes merged into the same table.

    Args:
        champ_alloc: Champion allocations dict {grid_id: {interval: weight, ...}}
        chall_alloc: Challenger allocations dict
        champ_acres: Champion acres dict {grid_id: acres}
        chall_acres: Challenger acres dict
        grid_list: List of grid IDs

    Returns:
        Styled pandas DataFrame with allocation and acreage changes
    """
    # Filter out grids with 0 acres in BOTH portfolios
    active_grids = [
        gid for gid in grid_list
        if champ_acres.get(gid, 0) > 0 or chall_acres.get(gid, 0) > 0
    ]

    rows = []
    total_changes = {interval: 0 for interval in INTERVAL_ORDER_11}
    total_champ_acres = 0
    total_chall_acres = 0

    for gid in active_grids:
        c_alloc = champ_alloc.get(gid, {})
        ch_alloc = chall_alloc.get(gid, {})
        row = {'Grid': gid}

        # Check if grid is only in one portfolio
        in_champ = gid in champ_alloc or gid in champ_acres
        in_chall = gid in chall_alloc or gid in chall_acres

        # Allocation changes for each interval
        for interval in INTERVAL_ORDER_11:
            if not in_champ and in_chall:
                # Grid only in Challenger - show "NEW"
                ch_val = ch_alloc.get(interval, 0)
                ch_pct = ch_val * 100 if isinstance(ch_val, (int, float)) and ch_val <= 1 else (ch_val if isinstance(ch_val, (int, float)) else 0)
                row[interval] = f"+{ch_pct:.0f}%" if ch_pct > 0 else "0%"
                total_changes[interval] += ch_pct
            elif in_champ and not in_chall:
                # Grid only in Champion - show "REMOVED"
                c_val = c_alloc.get(interval, 0)
                c_pct = c_val * 100 if isinstance(c_val, (int, float)) and c_val <= 1 else (c_val if isinstance(c_val, (int, float)) else 0)
                row[interval] = f"-{c_pct:.0f}%" if c_pct > 0 else "0%"
                total_changes[interval] -= c_pct
            else:
                # Grid in both portfolios - show change
                c_val = c_alloc.get(interval, 0)
                ch_val = ch_alloc.get(interval, 0)

                # Convert to percentage if decimal
                c_pct = c_val * 100 if isinstance(c_val, (int, float)) and c_val <= 1 else (c_val if isinstance(c_val, (int, float)) else 0)
                ch_pct = ch_val * 100 if isinstance(ch_val, (int, float)) and ch_val <= 1 else (ch_val if isinstance(ch_val, (int, float)) else 0)

                change = ch_pct - c_pct
                total_changes[interval] += change

                # Format: "+50%", "-50%", or "0%" (NOT "--")
                if change > 0:
                    row[interval] = f"+{change:.0f}%"
                elif change < 0:
                    row[interval] = f"{change:.0f}%"
                else:
                    row[interval] = "0%"

        # Net Change column
        if not in_champ and in_chall:
            row['Net Change'] = "NEW"
        elif in_champ and not in_chall:
            row['Net Change'] = "REMOVED"
        else:
            row['Net Change'] = "0%"

        # Acreage columns
        c_acres = champ_acres.get(gid, 0)
        ch_acres = chall_acres.get(gid, 0)
        acre_change = ch_acres - c_acres

        total_champ_acres += c_acres
        total_chall_acres += ch_acres

        # Format acres - show "N/A" if grid not in that portfolio
        row['Champ Acres'] = f"{c_acres:,.0f}" if in_champ else "N/A"
        row['Chall Acres'] = f"{ch_acres:,.0f}" if in_chall else "N/A"

        # Format acre change with +/- or 0
        if not in_champ and in_chall:
            row['Acre Change'] = f"+{ch_acres:,.0f}"
        elif in_champ and not in_chall:
            row['Acre Change'] = f"-{c_acres:,.0f}"
        elif acre_change > 0:
            row['Acre Change'] = f"+{acre_change:,.0f}"
        elif acre_change < 0:
            row['Acre Change'] = f"{acre_change:,.0f}"
        else:
            row['Acre Change'] = "0"

        rows.append(row)

    # Add PORTFOLIO TOTALS row (average allocation changes, total acres)
    totals_row = {'Grid': 'PORTFOLIO TOTALS'}
    grid_count = len(active_grids)

    for interval in INTERVAL_ORDER_11:
        avg_change = total_changes[interval] / grid_count if grid_count > 0 else 0
        if avg_change > 0:
            totals_row[interval] = f"+{avg_change:.0f}%"
        elif avg_change < 0:
            totals_row[interval] = f"{avg_change:.0f}%"
        else:
            totals_row[interval] = "0%"

    totals_row['Net Change'] = "0%"

    # Acre totals for the summary row
    total_acre_change = total_chall_acres - total_champ_acres
    totals_row['Champ Acres'] = f"{total_champ_acres:,.0f}"
    totals_row['Chall Acres'] = f"{total_chall_acres:,.0f}"

    if total_acre_change > 0:
        totals_row['Acre Change'] = f"+{total_acre_change:,.0f}"
    elif total_acre_change < 0:
        totals_row['Acre Change'] = f"{total_acre_change:,.0f}"
    else:
        totals_row['Acre Change'] = "0"

    rows.append(totals_row)

    df = pd.DataFrame(rows)

    # Style function for red/green/gray gradient on changes
    def highlight_change_cell(val):
        if isinstance(val, str):
            try:
                # Handle percentage values
                if val.endswith('%'):
                    change = float(val.replace('%', '').replace('+', ''))
                # Handle acre values with commas
                elif ',' in val or val.startswith('+') or val.startswith('-'):
                    change = float(val.replace(',', '').replace('+', ''))
                elif val == "0":
                    change = 0
                else:
                    return ''

                # Color coding
                if change == 0:
                    return 'background-color: #e0e0e0'  # Gray for zero
                elif change >= 40:
                    return 'background-color: #2e7d32; color: white'  # Dark green
                elif change >= 25:
                    return 'background-color: #4caf50; color: white'  # Medium green
                elif change >= 10:
                    return 'background-color: #81c784'  # Light green
                elif change > 0:
                    return 'background-color: #c8e6c9'  # Very light green
                elif change <= -40:
                    return 'background-color: #c62828; color: white'  # Dark red
                elif change <= -25:
                    return 'background-color: #e53935; color: white'  # Medium red
                elif change <= -10:
                    return 'background-color: #ef9a9a'  # Light red
                elif change < 0:
                    return 'background-color: #ffcdd2'  # Very light red
            except:
                pass
        return ''

    # Apply styling to interval columns, Net Change, and acre change column
    style_cols = list(INTERVAL_ORDER_11) + ['Net Change', 'Acre Change']
    styled = df.style.applymap(highlight_change_cell, subset=[c for c in style_cols if c in df.columns])

    return styled, df


def create_performance_comparison_table(champ_metrics, chall_metrics):
    """
    Create a performance comparison table with proper formatting.

    Args:
        champ_metrics: Dict with champion performance metrics
        chall_metrics: Dict with challenger performance metrics

    Returns:
        pandas DataFrame for display with st.table
    """
    # Extract metrics
    champ_roi = champ_metrics.get('cumulative_roi', 0)
    chall_roi = chall_metrics.get('cumulative_roi', 0)
    roi_change = chall_roi - champ_roi

    champ_risk_adj = champ_metrics.get('risk_adj_return', 0)
    chall_risk_adj = chall_metrics.get('risk_adj_return', 0)
    risk_adj_change = chall_risk_adj - champ_risk_adj

    champ_premium = champ_metrics.get('avg_annual_premium', 0)
    chall_premium = chall_metrics.get('avg_annual_premium', 0)
    premium_change = chall_premium - champ_premium

    champ_win = champ_metrics.get('profitable_pct', 0)
    chall_win = chall_metrics.get('profitable_pct', 0)
    win_change = chall_win - champ_win

    # Format with + prefix for positive changes
    def format_change_pct(val):
        return f"+{val:.1%}" if val > 0 else f"{val:.1%}"

    def format_change_float(val):
        return f"+{val:.2f}" if val > 0 else f"{val:.2f}"

    def format_change_currency(val):
        return f"+${val:,.0f}" if val > 0 else f"-${abs(val):,.0f}" if val < 0 else "$0"

    comparison_data = {
        'Metric': ['Cumulative ROI', 'Risk-Adjusted Return', 'Est. Annual Premium', 'Win Rate'],
        'Champion': [
            f"{champ_roi:.1%}",
            f"{champ_risk_adj:.2f}",
            f"${champ_premium:,.0f}",
            f"{champ_win:.0%}"
        ],
        'Challenger': [
            f"{chall_roi:.1%}",
            f"{chall_risk_adj:.2f}",
            f"${chall_premium:,.0f}",
            f"{chall_win:.0%}"
        ],
        'Change': [
            format_change_pct(roi_change),
            format_change_float(risk_adj_change),
            format_change_currency(premium_change),
            format_change_pct(win_change)
        ]
    }

    return pd.DataFrame(comparison_data)


def create_3way_comparison_table(champ_metrics, chall1_metrics, weather_metrics):
    """
    Create a 3-way performance comparison table: Champion vs Challenger 1 vs Weather Challenger.

    Args:
        champ_metrics: Dict with champion performance metrics
        chall1_metrics: Dict with challenger 1 performance metrics
        weather_metrics: Dict with weather challenger performance metrics

    Returns:
        pandas DataFrame for display with st.table
    """
    # Extract metrics
    champ_roi = champ_metrics.get('cumulative_roi', 0)
    chall1_roi = chall1_metrics.get('cumulative_roi', 0)
    weather_roi = weather_metrics.get('cumulative_roi', 0)

    champ_risk_adj = champ_metrics.get('risk_adj_return', 0)
    chall1_risk_adj = chall1_metrics.get('risk_adj_return', 0)
    weather_risk_adj = weather_metrics.get('risk_adj_return', 0)

    champ_premium = champ_metrics.get('avg_annual_premium', 0)
    chall1_premium = chall1_metrics.get('avg_annual_premium', 0)
    weather_premium = weather_metrics.get('avg_annual_premium', 0)

    champ_win = champ_metrics.get('profitable_pct', 0)
    chall1_win = chall1_metrics.get('profitable_pct', 0)
    weather_win = weather_metrics.get('profitable_pct', 0)

    comparison_data = {
        'Metric': ['Cumulative ROI', 'Risk-Adjusted Return', 'Est. Annual Premium', 'Win Rate'],
        'Champion': [
            f"{champ_roi:.1%}",
            f"{champ_risk_adj:.2f}",
            f"${champ_premium:,.0f}",
            f"{champ_win:.0%}"
        ],
        'Challenger 1': [
            f"{chall1_roi:.1%}",
            f"{chall1_risk_adj:.2f}",
            f"${chall1_premium:,.0f}",
            f"{chall1_win:.0%}"
        ],
        'Weather Chall 2': [
            f"{weather_roi:.1%}",
            f"{weather_risk_adj:.2f}",
            f"${weather_premium:,.0f}",
            f"{weather_win:.0%}"
        ]
    }

    return pd.DataFrame(comparison_data)


def render_allocation_text_table(allocations_dict, grid_list, grid_acres=None, label="AVERAGE"):
    """
    Render an allocation table as stable text output.

    Args:
        allocations_dict: Dict mapping grid_id -> {interval: weight (0-1), ...}
        grid_list: List of grid IDs to include
        grid_acres: Optional dict of grid_id -> acres
        label: Label for the summary row
    """
    # Filter to only grids with acres > 0
    if grid_acres:
        active_grids = [gid for gid in grid_list if grid_acres.get(gid, 0) > 0]
    else:
        active_grids = grid_list

    # Build data rows
    rows = []
    total_coverage = {interval: 0 for interval in INTERVAL_ORDER_11}
    total_acres = 0

    for gid in active_grids:
        alloc = allocations_dict.get(gid, {})
        row = {'Grid': str(gid)[:20]}  # Truncate long grid names
        row_sum = 0

        for interval in INTERVAL_ORDER_11:
            pct_raw = alloc.get(interval, 0)
            pct = pct_raw * 100 if isinstance(pct_raw, (int, float)) and pct_raw <= 1 else (pct_raw if isinstance(pct_raw, (int, float)) else 0)
            row_sum += pct
            total_coverage[interval] += pct
            row[interval] = f"{pct:.0f}%" if pct > 0 else "--"

        row['Row Sum'] = f"{row_sum:.0f}%"
        acres = grid_acres.get(gid, 0) if grid_acres else 0
        total_acres += acres
        row['Acres'] = f"{acres:,.0f}"
        rows.append(row)

    # Summary row
    grid_count = len(active_grids)
    avg_row = {'Grid': label}
    avg_sum = 0
    for interval in INTERVAL_ORDER_11:
        avg_pct = total_coverage[interval] / grid_count if grid_count > 0 else 0
        avg_sum += avg_pct
        avg_row[interval] = f"{avg_pct:.0f}%"
    avg_row['Row Sum'] = f"{avg_sum:.0f}%"
    avg_row['Acres'] = f"{total_acres:,.0f}"

    # Render header
    header = f"{'Grid':<20}"
    for interval in INTERVAL_ORDER_11:
        short_name = interval[:7]  # "Jan-Feb" etc
        header += f" {short_name:>7}"
    header += f" {'Row Sum':>8} {'Acres':>10}"

    st.text(header)
    st.text("─" * len(header))

    # Render data rows
    for row in rows:
        line = f"{row['Grid']:<20}"
        for interval in INTERVAL_ORDER_11:
            line += f" {row[interval]:>7}"
        line += f" {row['Row Sum']:>8} {row['Acres']:>10}"
        st.text(line)

    # Render summary row
    st.text("═" * len(header))
    line = f"{avg_row['Grid']:<20}"
    for interval in INTERVAL_ORDER_11:
        line += f" {avg_row[interval]:>7}"
    line += f" {avg_row['Row Sum']:>8} {avg_row['Acres']:>10}"
    st.text(line)


def render_change_analysis_text_table(champ_alloc, chall_alloc, champ_acres, chall_acres, grid_list):
    """
    Render a change analysis table as stable text output.

    Args:
        champ_alloc: Champion allocations dict
        chall_alloc: Challenger allocations dict
        champ_acres: Champion acres dict
        chall_acres: Challenger acres dict
        grid_list: List of grid IDs
    """
    # Build data rows
    rows = []
    total_changes = {interval: 0 for interval in INTERVAL_ORDER_11}
    total_champ_acres = 0
    total_chall_acres = 0

    for gid in grid_list:
        c_alloc = champ_alloc.get(gid, {})
        ch_alloc = chall_alloc.get(gid, {})
        row = {'Grid': str(gid)[:20]}

        for interval in INTERVAL_ORDER_11:
            c_val = c_alloc.get(interval, 0)
            ch_val = ch_alloc.get(interval, 0)
            c_pct = c_val * 100 if isinstance(c_val, (int, float)) and c_val <= 1 else (c_val if isinstance(c_val, (int, float)) else 0)
            ch_pct = ch_val * 100 if isinstance(ch_val, (int, float)) and ch_val <= 1 else (ch_val if isinstance(ch_val, (int, float)) else 0)
            change = ch_pct - c_pct
            total_changes[interval] += change

            if change > 0:
                row[interval] = f"+{change:.0f}%"
            elif change < 0:
                row[interval] = f"{change:.0f}%"
            else:
                row[interval] = "0%"

        row['Net'] = "0%"

        c_acres = champ_acres.get(gid, 0)
        ch_acres = chall_acres.get(gid, 0)
        acre_change = ch_acres - c_acres
        total_champ_acres += c_acres
        total_chall_acres += ch_acres

        row['Champ'] = f"{c_acres:,.0f}"
        row['Chall'] = f"{ch_acres:,.0f}"
        row['Acre Change'] = f"{acre_change:+,.0f}" if acre_change != 0 else "0"
        rows.append(row)

    # Totals row
    grid_count = len(grid_list)
    totals_row = {'Grid': 'PORTFOLIO TOTALS'}
    for interval in INTERVAL_ORDER_11:
        avg_change = total_changes[interval] / grid_count if grid_count > 0 else 0
        if avg_change > 0:
            totals_row[interval] = f"+{avg_change:.0f}%"
        elif avg_change < 0:
            totals_row[interval] = f"{avg_change:.0f}%"
        else:
            totals_row[interval] = "0%"
    totals_row['Net'] = "0%"
    total_acre_change = total_chall_acres - total_champ_acres
    totals_row['Champ'] = f"{total_champ_acres:,.0f}"
    totals_row['Chall'] = f"{total_chall_acres:,.0f}"
    totals_row['Acre Change'] = f"{total_acre_change:+,.0f}" if total_acre_change != 0 else "0"

    # Render header
    header = f"{'Grid':<20}"
    for interval in INTERVAL_ORDER_11:
        short_name = interval[:7]
        header += f" {short_name:>7}"
    header += f" {'Net':>5} {'Champ':>10} {'Chall':>10} {'Acre Change':>12}"

    st.text(header)
    st.text("─" * len(header))

    # Render data rows
    for row in rows:
        line = f"{row['Grid']:<20}"
        for interval in INTERVAL_ORDER_11:
            line += f" {row[interval]:>7}"
        line += f" {row['Net']:>5} {row['Champ']:>10} {row['Chall']:>10} {row['Acre Change']:>12}"
        st.text(line)

    # Render totals row
    st.text("═" * len(header))
    line = f"{totals_row['Grid']:<20}"
    for interval in INTERVAL_ORDER_11:
        line += f" {totals_row[interval]:>7}"
    line += f" {totals_row['Net']:>5} {totals_row['Champ']:>10} {totals_row['Chall']:>10} {totals_row['Acre Change']:>12}"
    st.text(line)


def render_portfolio_strategy_tab(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code):
    """
    Unified Portfolio Strategy tab implementing Champion vs Challenger workflow.
    """

    # ==========================================================================
    # CALLBACK FUNCTIONS (Must be defined before widgets that use them)
    # ==========================================================================
    def load_king_ranch_callback():
        """Simple callback that just sets a flag - actual loading happens in main function."""
        st.session_state.ps_kr_load_requested = True

    st.subheader("Portfolio Strategy: Champion vs. Challenger")

    # ==========================================================================
    # DEFERRED KING RANCH LOADING (executed where session is available)
    # ==========================================================================
    if st.session_state.get('ps_kr_load_requested', False):
        try:
            all_grids_for_loading = load_distinct_grids(session)

            target_grid_mapping = {}
            for county, grid_ids in KING_RANCH_PRESET['counties'].items():
                for gid_num in grid_ids:
                    target_grid_mapping[gid_num] = f"{gid_num} ({county} - TX)"

            preset_grid_ids = []
            for numeric_id in KING_RANCH_PRESET['grids']:
                target_str = target_grid_mapping.get(numeric_id, "")
                if target_str in all_grids_for_loading:
                    preset_grid_ids.append(target_str)
                else:
                    for grid_option in all_grids_for_loading:
                        if extract_numeric_grid_id(grid_option) == numeric_id:
                            preset_grid_ids.append(grid_option)
                            break

            # Set session state values
            st.session_state.ps_grids = preset_grid_ids
            st.session_state.productivity_factor = 1.35
            st.session_state.ps_coverage = 0.75  # Set to 75% coverage

            # Set acres for each grid
            for gid in preset_grid_ids:
                numeric_id = extract_numeric_grid_id(gid)
                st.session_state[f"ps_champ_acres_{gid}"] = KING_RANCH_PRESET['acres'].get(numeric_id, total_insured_acres)

            # Set allocations via preset keys (convert percentages to decimals)
            for gid in preset_grid_ids:
                numeric_id = extract_numeric_grid_id(gid)
                alloc = KING_RANCH_PRESET['allocations'][numeric_id]
                alloc_decimal = {interval: float(alloc.get(interval, 0.0)) / 100.0 for interval in INTERVAL_ORDER_11}
                st.session_state[f"ps_champ_{gid}_preset_allocation"] = alloc_decimal

            st.session_state.ps_kr_load_requested = False
            st.success("King Ranch preset loaded! (8 grids, 135% productivity, 75% coverage)")
            st.rerun()  # Force UI refresh to pick up new session state values

        except Exception as e:
            st.error(f"Error loading King Ranch: {e}")
            st.session_state.ps_kr_load_requested = False

    # ==========================================================================
    # SIDEBAR: LOAD KING RANCH (CHAMPION) HANDLER
    # ==========================================================================
    if st.session_state.get('sidebar_kr_champion_requested', False):
        try:
            all_grids_for_loading = load_distinct_grids(session)

            target_grid_mapping = {}
            for county, grid_ids in KING_RANCH_PRESET['counties'].items():
                for gid_num in grid_ids:
                    target_grid_mapping[gid_num] = f"{gid_num} ({county} - TX)"

            preset_grid_ids = []
            for numeric_id in KING_RANCH_PRESET['grids']:
                target_str = target_grid_mapping.get(numeric_id, "")
                if target_str in all_grids_for_loading:
                    preset_grid_ids.append(target_str)
                else:
                    for grid_option in all_grids_for_loading:
                        if extract_numeric_grid_id(grid_option) == numeric_id:
                            preset_grid_ids.append(grid_option)
                            break

            # Set session state values
            st.session_state.ps_grids = preset_grid_ids
            st.session_state.productivity_factor = 1.35
            st.session_state.ps_coverage = 0.75  # Set to 75% coverage

            # Set acres for each grid
            for gid in preset_grid_ids:
                numeric_id = extract_numeric_grid_id(gid)
                st.session_state[f"ps_champ_acres_{gid}"] = KING_RANCH_PRESET['acres'].get(numeric_id, total_insured_acres)

            # Set allocations via preset keys (convert percentages to decimals)
            for gid in preset_grid_ids:
                numeric_id = extract_numeric_grid_id(gid)
                alloc = KING_RANCH_PRESET['allocations'][numeric_id]
                alloc_decimal = {interval: float(alloc.get(interval, 0.0)) / 100.0 for interval in INTERVAL_ORDER_11}
                st.session_state[f"ps_champ_{gid}_preset_allocation"] = alloc_decimal

            st.session_state.sidebar_kr_champion_requested = False
            st.success("King Ranch Champion loaded! (8 grids, 135% productivity, 75% coverage)")
            st.rerun()  # Force UI refresh to pick up new session state values

        except Exception as e:
            st.error(f"Error loading King Ranch Champion: {e}")
            st.session_state.sidebar_kr_champion_requested = False

    # ==========================================================================
    # SIDEBAR: LOAD KING RANCH INCREMENTALS HANDLER
    # ==========================================================================
    if st.session_state.get('sidebar_kr_incrementals_requested', False):
        try:
            all_grids_for_loading = load_distinct_grids(session)

            # Build mapping for incremental grids
            target_grid_mapping = {}
            for county, grid_ids in KING_RANCH_INCREMENTAL_PRESET['counties'].items():
                for gid_num in grid_ids:
                    target_grid_mapping[gid_num] = f"{gid_num} ({county} - TX)"

            # Match grids from preset
            preset_incremental_grid_ids = []
            for numeric_id in KING_RANCH_INCREMENTAL_PRESET['grids']:
                target_str = target_grid_mapping.get(numeric_id, "")
                if target_str in all_grids_for_loading:
                    preset_incremental_grid_ids.append(target_str)
                else:
                    for grid_option in all_grids_for_loading:
                        if extract_numeric_grid_id(grid_option) == numeric_id:
                            preset_incremental_grid_ids.append(grid_option)
                            break

            # MERGE with existing incremental grids (don't replace!)
            # Get existing grids and acres
            existing_incremental_grids = list(st.session_state.get('ps_incremental_grids', []))

            # Add preset grids (skip duplicates)
            for grid in preset_incremental_grid_ids:
                if grid not in existing_incremental_grids:
                    existing_incremental_grids.append(grid)

            # Set merged incremental grids in session state
            st.session_state.ps_incremental_grids = existing_incremental_grids

            # Set acres for each incremental grid from preset (merge with existing)
            for gid in preset_incremental_grid_ids:
                numeric_id = extract_numeric_grid_id(gid)
                st.session_state[f"ps_incr_acres_{gid}"] = KING_RANCH_INCREMENTAL_PRESET['acres'].get(numeric_id, 40000)

            st.session_state.sidebar_kr_incrementals_requested = False
            st.success(f"King Ranch Incrementals loaded! ({len(preset_incremental_grid_ids)} preset grids added, {len(existing_incremental_grids)} total incremental grids)")
            st.rerun()  # Force UI refresh to pick up new session state values

        except Exception as e:
            st.error(f"Error loading King Ranch Incrementals: {e}")
            st.session_state.sidebar_kr_incrementals_requested = False

    # ==========================================================================
    # TOP SECTION: GLOBAL SETTINGS
    # ==========================================================================
    st.markdown("### Global Settings")

    # Grid Selection
    # Note: load_distinct_grids() is cached with @st.cache_data(ttl=3600) to minimize
    # unnecessary database queries and improve performance during page reruns
    try:
        all_grids = load_distinct_grids(session)
    except:
        all_grids = [grid_id]

    default_grids = st.session_state.get('ps_grids', [grid_id])
    # Ensure default_grids only contains valid options
    default_grids = [g for g in default_grids if g in all_grids]
    selected_grids = st.multiselect(
        "Select Grids for Portfolio",
        options=all_grids,
        default=default_grids,
        max_selections=20,
        key="ps_grids"
    )

    if not selected_grids:
        st.warning("Select at least one grid to continue.")
        return

    # Scenario and Policy Parameters
    col1, col2, col3 = st.columns(3)

    with col1:
        scenario_options = [
            'All Years (except Current Year)',
            'ENSO Phase: La Nina',
            'ENSO Phase: El Nino',
            'ENSO Phase: Neutral',
            'Custom Range'
        ]
        selected_scenario = st.selectbox("Scenario", scenario_options, key="ps_scenario")

    with col2:
        coverage_level = st.selectbox(
            "Coverage Level",
            [0.70, 0.75, 0.80, 0.85, 0.90],
            index=2,
            format_func=lambda x: f"{x:.0%}",
            key="ps_coverage"
        )

    with col3:
        st.metric("Productivity Factor", f"{productivity_factor:.0%}")

    # Year range (if custom)
    if selected_scenario == 'Custom Range':
        col1, col2 = st.columns(2)
        start_year = col1.selectbox("Start Year", list(range(1948, 2026)), index=62, key="ps_start")
        end_year = col2.selectbox("End Year", list(range(1948, 2026)), index=76, key="ps_end")
    else:
        start_year = 1948
        end_year = 2024

    st.divider()

    # ==========================================================================
    # HISTORICAL GRID CORRELATIONS (Optional Analysis)
    # ==========================================================================
    with st.expander("📉 Historical Grid Correlations", expanded=False):
        if len(selected_grids) < 2:
            st.info("Select at least two grids to view correlations.")
        else:
            try:
                # Collect index data for all selected grids
                all_index_data = []
                for gid in selected_grids:
                    grid_indices = load_all_indices(session, gid)

                    # Apply scenario filter to match user's selected scenario
                    if selected_scenario == 'Custom Range':
                        filtered_indices = filter_indices_by_scenario(
                            grid_indices, selected_scenario, start_year, end_year
                        )
                    else:
                        filtered_indices = filter_indices_by_scenario(
                            grid_indices, selected_scenario
                        )

                    # Add Grid ID and create time period identifier
                    filtered_indices = filtered_indices.copy()
                    filtered_indices['Grid ID'] = gid
                    filtered_indices['Period'] = filtered_indices['YEAR'].astype(str) + '-' + filtered_indices['INTERVAL_NAME']
                    all_index_data.append(filtered_indices[['Period', 'Grid ID', 'INDEX_VALUE']])

                if all_index_data:
                    # Combine all grid data
                    combined_df = pd.concat(all_index_data, ignore_index=True)

                    # Pivot so rows are time periods and columns are Grid IDs
                    pivot_df = combined_df.pivot_table(
                        values='INDEX_VALUE',
                        index='Period',
                        columns='Grid ID',
                        aggfunc='mean'
                    )

                    # Calculate correlation matrix
                    corr_matrix = pivot_df.corr()

                    # Generate heatmap
                    import matplotlib.pyplot as plt
                    import seaborn as sns

                    fig, ax = plt.subplots(figsize=(10, 6))
                    sns.heatmap(
                        corr_matrix,
                        annot=True,
                        cmap='RdYlGn',
                        fmt=".2f",
                        vmin=-1,
                        vmax=1,
                        center=0,
                        ax=ax,
                        square=True,
                        linewidths=0.5
                    )
                    ax.set_title(f"Grid Index Correlations ({selected_scenario})", fontsize=12)
                    plt.tight_layout()

                    st.pyplot(fig)
                    plt.close(fig)

                    st.caption("Correlations close to 1.0 indicate grids that perform similarly. "
                              "Lower correlations suggest diversification benefits.")
                else:
                    st.warning("No index data available for selected grids.")

            except Exception as e:
                st.error(f"Error generating correlation heatmap: {e}")

    # ==========================================================================
    # MIDDLE SECTION: THE CHAMPION (BASELINE)
    # ==========================================================================
    st.markdown("### The Champion (Baseline)")
    st.caption("Define your manual baseline strategy. This is what the optimizer will try to beat.")

    # Initialize expander state management (prevents auto-opening on rerun)
    if 'ps_acres_expander_opened' not in st.session_state:
        st.session_state.ps_acres_expander_opened = True  # Start open first time only
    if 'ps_alloc_expander_opened' not in st.session_state:
        st.session_state.ps_alloc_expander_opened = False  # Start collapsed

    # Champion Acreage Configuration
    with st.expander("Champion Acreage per Grid", expanded=st.session_state.ps_acres_expander_opened):
        champion_acres = {}
        cols = st.columns(min(4, len(selected_grids)))
        for idx, gid in enumerate(selected_grids):
            with cols[idx % 4]:
                numeric_id = extract_numeric_grid_id(gid)
                default_acres = KING_RANCH_PRESET['acres'].get(numeric_id, total_insured_acres)
                champion_acres[gid] = st.number_input(
                    f"{gid}",
                    min_value=1,
                    value=default_acres,
                    step=10,
                    key=f"ps_champ_acres_{gid}"
                )

    # Champion Interval Allocations
    champion_allocations = {}
    champion_all_valid = True

    with st.expander("Champion Interval Allocations", expanded=st.session_state.ps_alloc_expander_opened):
        for gid in selected_grids:
            st.markdown(f"**{gid}**")
            alloc_dict, is_valid = render_allocation_inputs(f"ps_champ_{gid}")
            champion_allocations[gid] = alloc_dict
            if not is_valid:
                champion_all_valid = False

    # Champion Run Button
    if st.button("Run Champion Backtest", key="ps_run_champion", disabled=not champion_all_valid):
        with st.spinner("Running Champion backtest..."):
            champion_df, champion_grid_results, champion_metrics = run_portfolio_backtest(
                session, selected_grids, champion_allocations, champion_acres,
                start_year, end_year, coverage_level, productivity_factor,
                intended_use, plan_code, selected_scenario
            )

            st.session_state.champion_results = {
                'df': champion_df,
                'grid_results': champion_grid_results,
                'metrics': champion_metrics,
                'allocations': champion_allocations,
                'acres': champion_acres,
                'grids': selected_grids  # Store champion's grid list
            }
            st.success("Champion backtest complete!")

    # Display Champion Results (if available)
    if 'champion_results' in st.session_state and st.session_state.champion_results:
        with st.container(border=True):
            st.markdown("#### 🏆 Champion Results")
            champ = st.session_state.champion_results
            metrics = champ.get('metrics', {})

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Cumulative ROI", f"{metrics.get('cumulative_roi', 0):.1%}")
            col2.metric("Risk-Adj Return", f"{metrics.get('risk_adj_return', 0):.2f}")
            col3.metric("Avg Annual Premium", f"${metrics.get('avg_annual_premium', 0):,.0f}")
            col4.metric("Win Rate", f"{metrics.get('profitable_pct', 0):.0%}")

    st.divider()

    # ==========================================================================
    # BOTTOM SECTION: THE CHALLENGER (OPTIMIZER)
    # ==========================================================================
    st.markdown("### The Challenger (Optimizer)")
    st.caption("Build on your Champion strategy by adding incremental grids.")

    # --------------------------------------------------------------------------
    # Challenger Portfolio Structure: Base (Champion) + Incremental Grids
    # --------------------------------------------------------------------------
    st.markdown("**Incremental Portfolio Expansion:**")

    # Show current base (Champion grids)
    st.caption(f"Base Portfolio: {len(selected_grids)} grid(s) inherited from Champion")

    # Incremental grids: all_grids EXCLUDING the Champion's selected_grids
    available_incremental_grids = [g for g in all_grids if g not in selected_grids]

    # Get default incremental grids from session state
    default_incremental = st.session_state.get('ps_incremental_grids', [])
    default_incremental = [g for g in default_incremental if g in available_incremental_grids]

    incremental_grids = st.multiselect(
        "Add Incremental Grids to Challenger",
        options=available_incremental_grids,
        default=default_incremental,
        max_selections=12,
        help="Select additional grids to expand beyond the Champion's base portfolio. The optimizer will find the best allocations for all grids.",
        key="ps_incremental_grids"
    )

    # Combine: Champion base + Incremental = Challenger portfolio
    challenger_grids = list(selected_grids) + list(incremental_grids)

    if not challenger_grids:
        st.warning("No grids available for Challenger portfolio.")
        return

    # Display summary
    if incremental_grids:
        st.success(f"Challenger Portfolio: {len(selected_grids)} base + {len(incremental_grids)} incremental = {len(challenger_grids)} total grids")
    else:
        st.info("No incremental grids selected. Challenger will optimize the same grids as Champion.")

    # --------------------------------------------------------------------------
    # Challenger Acreage: Inherit Base + Manual Incremental
    # --------------------------------------------------------------------------
    # Start with Champion acres for base grids
    challenger_acres = {}
    for gid in selected_grids:
        challenger_acres[gid] = champion_acres.get(gid, total_insured_acres / len(selected_grids))

    # Render manual inputs ONLY for incremental grids
    if incremental_grids:
        st.markdown("**Incremental Grid Acreage:**")
        st.caption("Set acres for each new grid being added to the Challenger portfolio.")

        default_incremental_acres = total_insured_acres / len(selected_grids)  # Use avg of base as default
        acre_cols = st.columns(min(4, len(incremental_grids)))

        for idx, gid in enumerate(incremental_grids):
            with acre_cols[idx % 4]:
                # Check session state first for preset values, otherwise use default
                preset_acres = st.session_state.get(f"ps_incr_acres_{gid}", int(default_incremental_acres))
                challenger_acres[gid] = st.number_input(
                    f"{gid}",
                    min_value=1,
                    value=preset_acres,
                    step=10,
                    key=f"ps_incr_acres_{gid}"
                )

    # Show summary of all acres
    total_challenger_acres = sum(challenger_acres.values())
    base_acres = sum(champion_acres.get(g, 0) for g in selected_grids)
    incremental_acres_total = total_challenger_acres - base_acres

    if incremental_grids:
        st.caption(f"Total Challenger Acres: {total_challenger_acres:,.0f} ({base_acres:,.0f} base + {incremental_acres_total:,.0f} incremental)")
    else:
        st.caption(f"Challenger inherits Champion acres: {total_challenger_acres:,.0f} total")

    # --------------------------------------------------------------------------
    # Manual Incremental Allocations (Optional)
    # --------------------------------------------------------------------------
    use_manual_incremental_allocations = False
    incremental_allocations = {}
    incremental_all_valid = True

    if incremental_grids:
        use_manual_incremental_allocations = st.checkbox(
            "Manually set interval allocations for incremental grids",
            value=False,
            help="If unchecked, the optimizer will find the best allocations. If checked, you can specify allocations like Champion.",
            key="ps_manual_incremental_alloc"
        )

        if use_manual_incremental_allocations:
            with st.expander("Incremental Grid Interval Allocations", expanded=True):
                st.caption("Set allocations for each incremental grid. Base grids will inherit Champion's allocations.")
                for gid in incremental_grids:
                    st.markdown(f"**{gid}**")
                    alloc_dict, is_valid = render_allocation_inputs(f"ps_incr_alloc_{gid}")
                    incremental_allocations[gid] = alloc_dict
                    if not is_valid:
                        incremental_all_valid = False

    st.markdown("---")

    # === PREMIUM BUDGET (First - Optional) ===
    st.markdown("**Premium Budget**")

    enable_budget = st.checkbox(
        "Set annual premium budget",
        value=False,
        help="Limit total annual premium spending",
        key="ps_challenger_budget"
    )

    annual_budget = 50000
    allow_scale_up = False
    if enable_budget:
        budget_col1, budget_col2 = st.columns(2)
        with budget_col1:
            annual_budget = st.number_input(
                "Maximum Annual Premium ($)",
                min_value=1000,
                value=50000,
                step=1000,
                help="Maximum producer premium (after subsidy) per year",
                key="ps_challenger_budget_amt"
            )
        with budget_col2:
            allow_scale_up = st.checkbox(
                "Auto-fill budget (scale up if under)",
                value=False,
                help="If the optimized strategy comes in under budget, automatically scale up acres to utilize the full budget.",
                key="ps_challenger_autofill"
            )

    st.markdown("---")

    # === INTERVAL STRATEGY (Second) ===
    st.markdown("**Interval Strategy**")
    col1, col2 = st.columns(2)

    with col1:
        use_marginal_search = st.checkbox(
            "Use Marginal Search (Perturb Naive Allocation)",
            value=False,
            key="ps_challenger_marginal"
        )

    if use_marginal_search:
        search_mode = 'marginal'
        search_iterations = 1000
        st.info("Marginal search perturbs naive allocation by ±5%")
    else:
        with col2:
            iteration_map = {'Fast': 500, 'Standard': 3000, 'Thorough': 7000, 'Maximum': 15000}
            search_depth_key = st.select_slider(
                "Search Depth",
                options=list(iteration_map.keys()),
                value='Standard',
                key="ps_challenger_depth"
            )
            search_iterations = iteration_map[search_depth_key]
            search_mode = 'global'

    # Diversification Constraints
    st.markdown("**Diversification Constraints:**")
    div_col1, div_col2 = st.columns(2)

    with div_col1:
        require_full_coverage = st.checkbox(
            "📅 Ensure Full Calendar Coverage",
            value=False,
            help="Creates staggered portfolio with Pattern A (6 intervals) and Pattern B (5 intervals) to cover all 11 intervals while maintaining non-adjacency. Requires at least 2 grids.",
            key="ps_challenger_full_coverage"
        )

    with div_col2:
        if not require_full_coverage:
            interval_range_opt = st.slider(
                "Number of Active Intervals Range",
                min_value=2,
                max_value=6,
                value=(2, 6),
                help="Optimizer can choose ANY number of intervals within this range. Each interval must have 10%-50%.",
                key="ps_challenger_interval_range"
            )
        else:
            interval_range_opt = (11, 11)  # Full coverage uses all intervals via Pattern A/B
            st.info("Full coverage mode: Pattern A (6) + Pattern B (5) = 11 intervals")

    st.caption("Diversification controls how interval allocations are spread across your portfolio. "
               "Full Calendar Coverage ensures year-round protection by using complementary patterns across grids.")

    st.markdown("---")

    # === ACREAGE STRATEGY (Third) ===
    st.markdown("**Acreage Strategy**")

    optimize_acreage = st.checkbox(
        "Optimize acreage distribution (Mean-Variance Optimization)",
        value=False,
        help="Use Markowitz portfolio theory to redistribute acres based on risk-adjusted returns",
        key="ps_challenger_mvo"
    )

    risk_aversion = 1.0
    max_turnover = 0.20
    if optimize_acreage:
        mvo_col1, mvo_col2 = st.columns(2)
        with mvo_col1:
            risk_profile = st.select_slider(
                "Risk Profile",
                options=["Aggressive", "Growth", "Balanced", "Conservative", "Defensive"],
                value="Balanced",
                help="Aggressive = chase highest returns, accept volatility. Defensive = prioritize stability and diversification.",
                key="ps_challenger_risk_profile"
            )

            # Map to actual lambda values for the optimizer
            risk_aversion_map = {
                "Aggressive": 0.5,
                "Growth": 0.75,
                "Balanced": 1.0,
                "Conservative": 1.5,
                "Defensive": 2.0
            }
            risk_aversion = risk_aversion_map[risk_profile]

        with mvo_col2:
            max_turnover = st.slider(
                "Max Turnover",
                min_value=0,
                max_value=100,
                value=20,
                step=5,
                format="%d%%",
                help="How much each grid's allocation can change. 0% = no changes allowed, 100% = full reallocation allowed.",
                key="ps_challenger_turnover"
            ) / 100.0  # Convert back to decimal for calculations
        st.info("MVO will redistribute acres across grids based on historical correlations and returns.")

    st.divider()

    # ==========================================================================
    # TRAIN CHALLENGER BUTTON
    # ==========================================================================
    # Disable button if manual incremental allocations are enabled but invalid
    train_button_disabled = use_manual_incremental_allocations and not incremental_all_valid

    if train_button_disabled:
        st.warning("Please fix the incremental grid allocation errors before training.")

    if st.button("Train Challenger", key="ps_train_challenger", type="primary", disabled=train_button_disabled):

        if 'champion_results' not in st.session_state or not st.session_state.champion_results:
            st.warning("Please run the Champion backtest first!")
        else:
            try:
                # ===== STEP 1: Optimize Intervals =====
                st.write("**Step 1: Optimizing Interval Allocations...**")

                challenger_allocations = {}
                challenger_interval_stats = {}

                # Determine which grids need optimization vs manual allocation
                grids_to_optimize = []
                for gid in challenger_grids:
                    if use_manual_incremental_allocations and gid in incremental_allocations:
                        # Use manual allocation for this incremental grid
                        challenger_allocations[gid] = incremental_allocations[gid]
                        challenger_interval_stats[gid] = {
                            'roi': None,  # Not optimized
                            'tested': 0,
                            'allocation': incremental_allocations[gid],
                            'manual': True
                        }
                    else:
                        grids_to_optimize.append(gid)

                # Show info about manual vs optimized grids
                if use_manual_incremental_allocations and incremental_allocations:
                    st.info(f"Using manual allocations for {len(incremental_allocations)} incremental grid(s). "
                           f"Optimizing {len(grids_to_optimize)} base grid(s).")

                # Optimize remaining grids
                if grids_to_optimize:
                    progress_bar = st.progress(0, text="Starting interval optimization...")

                    for idx, gid in enumerate(grids_to_optimize):
                        progress_bar.progress(
                            (idx + 1) / len(grids_to_optimize),
                            text=f"Optimizing intervals for {gid}..."
                        )

                        # Find the original index for full coverage pattern assignment
                        original_idx = challenger_grids.index(gid)

                        best_alloc, best_roi, tested = run_fast_optimization_core(
                            session, gid, start_year, end_year, plan_code,
                            productivity_factor, challenger_acres[gid], intended_use,
                            coverage_level, search_iterations, search_mode,
                            require_full_coverage=require_full_coverage,
                            interval_range_opt=interval_range_opt,
                            grid_index=original_idx
                        )

                        if best_alloc:
                            challenger_allocations[gid] = best_alloc
                            challenger_interval_stats[gid] = {
                                'roi': best_roi,
                                'tested': tested,
                                'allocation': best_alloc
                            }
                        else:
                            # Fallback to naive allocation (no champion reference for different grids)
                            naive_alloc = {interval: 0.0 for interval in INTERVAL_ORDER_11}
                            for i, interval in enumerate([0, 2, 4, 6, 8]):
                                naive_alloc[INTERVAL_ORDER_11[interval]] = 0.20
                            challenger_allocations[gid] = naive_alloc

                    progress_bar.empty()

                # Report results
                optimized_count = sum(1 for s in challenger_interval_stats.values() if not s.get('manual', False))
                manual_count = sum(1 for s in challenger_interval_stats.values() if s.get('manual', False))
                total_tested = sum(s['tested'] for s in challenger_interval_stats.values())

                if manual_count > 0:
                    st.success(f"Interval allocation complete! {optimized_count} grid(s) optimized ({total_tested:,} strategies tested), "
                              f"{manual_count} grid(s) using manual allocations.")
                else:
                    st.success(f"Interval optimization complete! Tested {total_tested:,} strategies.")

                # ===== STEP 2: Acreage Optimization =====
                # challenger_acres already defined from UI (default or manual override)
                initial_challenger_acres = challenger_acres.copy()

                if optimize_acreage:
                    st.write("**Step 2: Optimizing Acreage Distribution (MVO)...**")

                    # Build proper grid_results structure for MVO using challenger_grids
                    grid_results_for_mvo = {}
                    for gid in challenger_grids:
                        grid_results_for_mvo[gid] = {
                            'best_strategy': {
                                'allocation': challenger_allocations[gid],
                                'coverage_level': coverage_level
                            }
                        }

                    # Regenerate base_data_df properly using challenger_grids
                    base_data_rows = []
                    for gid in challenger_grids:
                        if gid in challenger_interval_stats:
                            # Use the stats we already computed
                            allocation = challenger_allocations[gid]

                            subsidy = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                            county_base_value = load_county_base_value(session, gid)
                            current_rate_year = get_current_rate_year(session)
                            premium_rates = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]

                            dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)

                            all_indices_df = load_all_indices(session, gid)
                            all_indices_df = all_indices_df[
                                (all_indices_df['YEAR'] >= start_year) &
                                (all_indices_df['YEAR'] <= end_year)
                            ]

                            for year in range(start_year, end_year + 1):
                                year_data = all_indices_df[all_indices_df['YEAR'] == year]
                                if year_data.empty:
                                    continue

                                year_indemnity = 0
                                year_premium = 0

                                for interval, pct in allocation.items():
                                    if pct == 0:
                                        continue

                                    index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                                    index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

                                    premium_rate = premium_rates.get(interval, 0)
                                    interval_protection = int(round_half_up(dollar_protection * 1 * pct, 0))  # 1 acre
                                    total_prem = int(round_half_up(interval_protection * premium_rate, 0))
                                    prem_subsidy = int(round_half_up(total_prem * subsidy, 0))
                                    producer_prem = total_prem - prem_subsidy

                                    trigger = coverage_level * 100
                                    shortfall = max(0, (trigger - index_value) / trigger)
                                    raw_indemnity = shortfall * interval_protection
                                    # Convert to int immediately to ensure exact integer arithmetic when summing
                                    indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                                    year_indemnity += indemnity
                                    year_premium += producer_prem

                                # Both indemnity and premium are already exact integer sums

                                roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                                base_data_rows.append({'year': year, 'grid': gid, 'roi': roi})

                    base_data_df = pd.DataFrame(base_data_rows)

                    if enable_budget:
                        # MVO with budget constraint (two-stage: budget scaling then MVO rebalancing)
                        challenger_acres, roi_corr, opt_info = optimize_grid_allocation(
                            base_data_df, grid_results_for_mvo, initial_challenger_acres,
                            annual_budget, session, productivity_factor, intended_use, plan_code,
                            challenger_grids, risk_aversion, max_turnover, allow_scale_up
                        )

                        # Display two-stage optimization feedback
                        scale_factor = opt_info.get('budget_scale_factor', 1.0)
                        initial_total = opt_info.get('initial_total_acres', 0)
                        scaled_total = opt_info.get('budget_scaled_acres', 0)
                        if scale_factor != 1.0:
                            st.info(f"**Stage 1 - Budget scaling:** {initial_total:,.0f} → {scaled_total:,.0f} acres ({scale_factor:.1%} of original)")
                        st.info(f"**Stage 2 - MVO rebalancing:** Grid weights adjusted within ±{max_turnover:.0%} turnover bounds")
                        # Show Stage 3 messaging if it was applied
                        if opt_info.get('stage3_applied', False):
                            stage3_scale = opt_info.get('stage3_scale_factor', 1.0)
                            st.success(f"**Stage 3 - Budget fill-up:** Acres scaled up by {stage3_scale:.1%} to utilize full budget")
                    else:
                        # MVO without budget - use optimize_without_budget
                        total_acres = sum(initial_challenger_acres.values())
                        challenger_acres = optimize_without_budget(
                            base_data_df, grid_results_for_mvo, total_acres,
                            challenger_grids, risk_aversion, max_turnover, initial_challenger_acres
                        )

                    st.success("Acreage optimization complete!")

                elif enable_budget:
                    st.write("**Step 2: Applying Budget Constraint...**")

                    # Calculate current cost with optimized intervals using challenger_grids
                    grid_results_for_cost = {
                        gid: {'best_strategy': {'allocation': challenger_allocations[gid], 'coverage_level': coverage_level}}
                        for gid in challenger_grids
                    }

                    total_cost, _ = calculate_annual_premium_cost(
                        session, challenger_grids, initial_challenger_acres, grid_results_for_cost,
                        productivity_factor, intended_use, plan_code
                    )

                    # Apply proportional scaling (with optional scale-up)
                    challenger_acres, scale_factor = apply_budget_constraint(
                        initial_challenger_acres, total_cost, annual_budget, allow_scale_up=allow_scale_up
                    )

                    if scale_factor < 1.0:
                        st.info(f"Acres scaled DOWN by {scale_factor:.1%} to fit budget")
                    elif scale_factor > 1.0:
                        st.success(f"Acres scaled UP by {scale_factor:.1%} to fill budget")
                    else:
                        st.info("Budget constraint not binding")

                # ===== STEP 3: Final Backtest =====
                st.write("**Step 3: Running Final Challenger Backtest...**")

                challenger_df, challenger_grid_results, challenger_metrics = run_portfolio_backtest(
                    session, challenger_grids, challenger_allocations, challenger_acres,
                    start_year, end_year, coverage_level, productivity_factor,
                    intended_use, plan_code, selected_scenario
                )

                st.session_state.challenger_results = {
                    'df': challenger_df,
                    'grid_results': challenger_grid_results,
                    'metrics': challenger_metrics,
                    'allocations': challenger_allocations,
                    'acres': challenger_acres,
                    'interval_stats': challenger_interval_stats,
                    'grids': challenger_grids  # Store challenger's grid list
                }

                st.success("Challenger training complete!")

            except Exception as e:
                st.error(f"Error training challenger: {e}")
                st.exception(e)

    # ==========================================================================
    # COMPARISON OUTPUT
    # ==========================================================================
    if ('champion_results' in st.session_state and st.session_state.champion_results and
        'challenger_results' in st.session_state and st.session_state.challenger_results):

        st.divider()
        st.markdown("### Results: Champion vs. Challenger")

        champ = st.session_state.champion_results
        chall = st.session_state.challenger_results

        champ_metrics = champ.get('metrics', {})
        chall_metrics = chall.get('metrics', {})

        # Performance Comparison Table (using st.table for clean display)
        st.markdown("#### Performance Comparison")
        comparison_df = create_performance_comparison_table(champ_metrics, chall_metrics)
        st.table(comparison_df.set_index('Metric'))

        # Winner Banner
        chall_roi = chall_metrics.get('cumulative_roi', 0)
        champ_roi = champ_metrics.get('cumulative_roi', 0)

        if chall_roi > champ_roi:
            st.success("**CHALLENGER WINS!**")
        elif chall_roi < champ_roi:
            st.warning("**CHAMPION HOLDS!**")
        else:
            st.info("**TIE!**")

        # === ALLOCATION COMPARISON: VERTICAL STACK (Styled DataFrames with Downloads) ===
        st.markdown("#### Interval Allocation Comparison")

        # Get grids from stored results (may differ between Champion and Challenger)
        champ_grids = champ.get('grids', selected_grids)
        chall_grids = chall.get('grids', selected_grids)

        # Check if grids differ between portfolios
        grids_differ = set(champ_grids) != set(chall_grids)
        if grids_differ:
            st.info(f"Note: Champion uses {len(champ_grids)} grids, Challenger uses {len(chall_grids)} grids. "
                    "Tables show each portfolio's grids; change analysis shows the union of all grids.")

        st.caption("Tables show allocation percentages by interval. Row Sum should be 100%.")

        # Champion Table (styled dataframe with download)
        with st.container(border=True):
            st.markdown("**🏆 Champion (Baseline)**")
            champ_styled, champ_df = create_optimized_allocation_table(
                champ['allocations'], champ_grids, grid_acres=champ['acres'],
                label="CHAMPION AVERAGE"
            )
            st.dataframe(champ_styled, use_container_width=True, hide_index=True)
            st.download_button(
                label="Download Champion CSV",
                data=champ_df.to_csv(index=False),
                file_name="champion_allocations.csv",
                mime="text/csv",
                key="download_champion_csv"
            )

        st.markdown("")  # Spacer

        # Challenger 1 Table (styled dataframe with download)
        with st.container(border=True):
            st.markdown("**⚔️ Challenger 1 (Optimized Intervals)**")
            chall_styled, chall_df = create_optimized_allocation_table(
                chall['allocations'], chall_grids, grid_acres=chall['acres'],
                label="OPTIMIZED AVERAGE"
            )
            st.dataframe(chall_styled, use_container_width=True, hide_index=True)
            st.download_button(
                label="Download Challenger CSV",
                data=chall_df.to_csv(index=False),
                file_name="challenger_allocations.csv",
                mime="text/csv",
                key="download_challenger_csv"
            )

        st.markdown("")  # Spacer

        # Change Analysis Table (styled dataframe with download)
        # Use UNION of grids for change analysis
        all_grids_union = list(set(champ_grids) | set(chall_grids))
        all_grids_union.sort()  # Sort for consistent ordering

        st.markdown("**Allocation Changes by Grid and Interval**")
        st.caption("Changes show +/- percentage shifts. Grids in only one portfolio show as N/A for the other. PORTFOLIO TOTALS shows average shifts and total acres.")
        change_styled, change_df = create_change_analysis_table(
            champ['allocations'], chall['allocations'],
            champ['acres'], chall['acres'],
            all_grids_union
        )
        st.dataframe(change_styled, use_container_width=True, hide_index=True)
        st.download_button(
            label="Download Changes CSV",
            data=change_df.to_csv(index=False),
            file_name="allocation_changes.csv",
            mime="text/csv",
            key="download_changes_csv"
        )

        # ==========================================================================
        # PHASE 2: WEATHER VIEW CHALLENGERS (Only shows after Challenger 1 exists)
        # ==========================================================================
        st.divider()
        st.markdown("### 🌦️ Add Weather View Challengers")
        st.caption("Generate additional challengers optimized for a specific weather outlook. These use portfolio-aggregated analog years for apples-to-apples comparison.")

        enable_weather_challengers = st.checkbox(
            "Enable Weather View Challengers",
            value=st.session_state.get('ps_enable_weather', False),
            key="ps_enable_weather"
        )

        if enable_weather_challengers:
            # --- Grid Selection for Weather Portfolio ---
            st.markdown("**Step 2.1: Select Grids for Weather Portfolio**")
            st.caption("Weather portfolio can use different grids than Champion. Default is Champion grids + Incremental grids with original user-specified acres.")

            # Default to Champion grids + Incremental grids (using original inputs, not MVO results)
            default_weather_grids = list(selected_grids) + list(st.session_state.get('ps_incremental_grids', []))

            # Filter to only valid options
            default_weather_grids = [g for g in default_weather_grids if g in all_grids]

            weather_grids = st.multiselect(
                "Weather Portfolio Grids",
                options=all_grids,
                default=default_weather_grids,
                max_selections=20,
                key="ps_weather_grids"
            )

            if weather_grids:
                # PRE-POPULATE acres in session state from ORIGINAL user inputs (not MVO results)
                # This must happen BEFORE the number_input widgets render
                # Only set if the key doesn't exist yet (first time)

                # Get the incremental grids for source determination
                incremental_grids = st.session_state.get('ps_incremental_grids', [])

                # Pre-populate session state for each weather grid
                # Calculate default acres for fallback (distributed evenly across grids)
                default_per_grid = max(1, total_insured_acres // len(weather_grids))

                for gid in weather_grids:
                    key = f"ps_weather_acres_{gid}"
                    # Only set if key doesn't exist yet
                    if key not in st.session_state:
                        # Use ORIGINAL inputs, not MVO-optimized results
                        # For Champion grids, use original Champion acres
                        if gid in selected_grids:
                            source_acres = st.session_state.get(f"ps_champ_acres_{gid}", default_per_grid)
                            st.session_state[key] = max(1, int(source_acres)) if source_acres >= 1 else default_per_grid
                        # For Incremental grids, use original incremental acres
                        elif gid in incremental_grids:
                            source_acres = st.session_state.get(f"ps_incr_acres_{gid}", default_per_grid)
                            st.session_state[key] = max(1, int(source_acres)) if source_acres >= 1 else default_per_grid
                        else:
                            # Fallback for any other grids
                            st.session_state[key] = default_per_grid

                # --- Acres per Grid ---
                st.markdown("**Acres per Grid (Starting Population)**")
                st.caption("Set acres for each grid. MVO optimization may adjust these.")

                # Add sync button to refresh acres from ORIGINAL inputs (not MVO results)
                if st.button("🔄 Sync Acres from Original Inputs", key="ps_sync_weather_acres"):
                    for gid in weather_grids:
                        key = f"ps_weather_acres_{gid}"
                        # Use ORIGINAL inputs, not MVO-optimized results
                        # For Champion grids, use original Champion acres
                        if gid in selected_grids:
                            source_acres = st.session_state.get(f"ps_champ_acres_{gid}", default_per_grid)
                            st.session_state[key] = max(1, int(source_acres)) if source_acres >= 1 else default_per_grid
                        # For Incremental grids, use original incremental acres
                        elif gid in incremental_grids:
                            source_acres = st.session_state.get(f"ps_incr_acres_{gid}", default_per_grid)
                            st.session_state[key] = max(1, int(source_acres)) if source_acres >= 1 else default_per_grid
                        else:
                            # Fallback for any other grids
                            st.session_state[key] = default_per_grid
                    st.rerun()

                weather_acres = {}
                acre_cols = st.columns(min(4, len(weather_grids)))
                for idx, gid in enumerate(weather_grids):
                    with acre_cols[idx % 4]:
                        # Session state was pre-populated above, so this will use those values
                        # Final safeguard: max(1, value) ensures valid input even if session state has 0
                        weather_acres[gid] = st.number_input(
                            f"{gid}",
                            min_value=1,
                            value=max(1, st.session_state.get(f"ps_weather_acres_{gid}", 10000)),
                            step=10,
                            key=f"ps_weather_acres_{gid}"
                        )

                st.markdown("---")

                # --- Market View Definition ---
                st.markdown("**Step 2.2: Define Market View**")
                st.caption("Select the weather scenario you want to optimize for.")

                mv_col1, mv_col2, mv_col3 = st.columns(3)

                with mv_col1:
                    enso_regime = st.selectbox(
                        "ENSO Regime",
                        options=["La Nina", "El Nino", "Neutral", "Some La Nina", "Some El Nino", "Any"],
                        index=0,
                        key="ps_weather_enso",
                        help="Strict filters (La Nina/El Nino/Neutral) require majority of intervals. 'Some' filters include years with at least one interval of that phase."
                    )

                with mv_col2:
                    historical_context = st.selectbox(
                        "Historical Context",
                        options=["Dry", "Normal", "Wet", "Any"],
                        index=0,
                        key="ps_weather_hist_context"
                    )

                with mv_col3:
                    trajectory = st.selectbox(
                        "Expected Trajectory",
                        options=["Get Wetter", "Stay Stable", "Get Drier", "Any"],
                        index=0,
                        key="ps_weather_trajectory"
                    )

                st.markdown("---")

                # --- Optimization Settings ---
                st.markdown("**Step 2.3: Optimization Settings**")

                opt_col1, opt_col2 = st.columns(2)

                with opt_col1:
                    weather_rank_by = st.selectbox(
                        "Rank Best Allocations By",
                        options=["Portfolio Return", "Risk-Adjusted Return", "Median ROI", "Win Rate"],
                        index=0,
                        key="ps_weather_rank_by"
                    )

                with opt_col2:
                    weather_interval_range = st.slider(
                        "Allowed Intervals per Grid",
                        min_value=2,
                        max_value=6,
                        value=(2, 6),
                        key="ps_weather_interval_range"
                    )

                st.markdown("---")

                # --- Find Analog Years Button ---
                if st.button("🔍 Find Analog Years", key="ps_find_analog_years", type="secondary"):
                    with st.spinner("Analyzing historical years across portfolio..."):
                        analog_years = calculate_portfolio_aggregated_analog_years(
                            session,
                            weather_grids,
                            enso_regime,
                            historical_context,
                            trajectory
                        )

                        st.session_state.ps_analog_years = analog_years
                        st.session_state.ps_weather_config = {
                            'grids': weather_grids,
                            'acres': weather_acres,
                            'enso_regime': enso_regime,
                            'historical_context': historical_context,
                            'trajectory': trajectory,
                            'rank_by': weather_rank_by,
                            'interval_range': weather_interval_range
                        }

                # Display analog years results
                if 'ps_analog_years' in st.session_state:
                    analog_years = st.session_state.ps_analog_years

                    if analog_years and len(analog_years) > 0:
                        # Build descriptive label with all three criteria
                        mv_enso = st.session_state.get('ps_weather_enso', 'Any')
                        mv_context = st.session_state.get('ps_weather_hist_context', 'Any')
                        mv_trend = st.session_state.get('ps_weather_trajectory', 'Any')

                        criteria_parts = []
                        if mv_enso != 'Any':
                            criteria_parts.append(mv_enso)
                        if mv_context != 'Any':
                            criteria_parts.append(mv_context)
                        if mv_trend != 'Any':
                            criteria_parts.append(mv_trend)

                        criteria_label = " + ".join(criteria_parts) if criteria_parts else "All Conditions"
                        st.success(f"**Found {len(analog_years)} analog years** matching: {criteria_label}")

                        # Show details in expander
                        with st.expander("📅 View Analog Year Details", expanded=False):
                            analog_df = pd.DataFrame(analog_years)

                            # Drop phase_agreement column before renaming
                            if 'phase_agreement' in analog_df.columns:
                                analog_df = analog_df.drop(columns=['phase_agreement'])

                            analog_df = analog_df.rename(columns={
                                'year': 'Year',
                                'dominant_phase': 'Dominant ENSO Phase',
                                'portfolio_avg_z': 'Portfolio Avg Z',
                                'portfolio_trajectory': 'Avg Trajectory',
                                'grids_with_data': 'Grids w/ Data',
                                'la_nina_intervals': 'Avg La Nina Int/Grid',
                                'el_nino_intervals': 'Avg El Nino Int/Grid'
                            })

                            # Get current ENSO filter to determine which columns to display
                            current_enso_filter = st.session_state.get('ps_weather_enso', 'Any')

                            # Select columns to display based on filter
                            base_columns = ['Year', 'Dominant ENSO Phase', 'Portfolio Avg Z', 'Avg Trajectory', 'Grids w/ Data']

                            if current_enso_filter == 'Some La Nina' and 'Avg La Nina Int/Grid' in analog_df.columns:
                                display_columns = ['Year', 'Dominant ENSO Phase', 'Avg La Nina Int/Grid', 'Portfolio Avg Z', 'Avg Trajectory', 'Grids w/ Data']
                            elif current_enso_filter == 'Some El Nino' and 'Avg El Nino Int/Grid' in analog_df.columns:
                                display_columns = ['Year', 'Dominant ENSO Phase', 'Avg El Nino Int/Grid', 'Portfolio Avg Z', 'Avg Trajectory', 'Grids w/ Data']
                            else:
                                display_columns = base_columns

                            # Filter to only existing columns
                            display_columns = [c for c in display_columns if c in analog_df.columns]
                            analog_df_display = analog_df[display_columns]

                            # Build format dict dynamically
                            format_dict = {
                                'Year': '{:.0f}',
                                'Portfolio Avg Z': '{:.3f}',
                                'Avg Trajectory': '{:.3f}',
                                'Grids w/ Data': '{:.0f}'
                            }
                            if 'Avg La Nina Int/Grid' in analog_df_display.columns:
                                format_dict['Avg La Nina Int/Grid'] = '{:.1f}'
                            if 'Avg El Nino Int/Grid' in analog_df_display.columns:
                                format_dict['Avg El Nino Int/Grid'] = '{:.1f}'

                            st.dataframe(
                                analog_df_display.style.format(format_dict),
                                use_container_width=True,
                                hide_index=True
                            )

                            # Download button (includes all columns for full data)
                            csv_data = analog_df.to_csv(index=False)
                            st.download_button(
                                label="📥 Download Analog Years CSV",
                                data=csv_data,
                                file_name="analog_years.csv",
                                mime="text/csv"
                            )

                        st.markdown("---")

                        # --- Generate Weather Challenger Button ---
                        weather2_iteration_map = {'Fast': 500, 'Standard': 3000, 'Thorough': 7000, 'Maximum': 15000}
                        weather2_depth_key = st.select_slider(
                            "Search Depth",
                            options=list(weather2_iteration_map.keys()),
                            value='Standard',
                            key="ps_weather2_depth"
                        )
                        weather2_iterations = weather2_iteration_map[weather2_depth_key]

                        if st.button("🌦️ Generate Weather Challenger 2", key="ps_generate_weather", type="primary"):
                            weather_config = st.session_state.get('ps_weather_config', {})
                            weather_grids_gen = weather_config.get('grids', [])
                            weather_acres_gen = weather_config.get('acres', {})
                            interval_range_gen = weather_config.get('interval_range', (2, 6))

                            # Extract year list from analog years
                            analog_year_list = [y['year'] for y in analog_years]

                            if len(weather_grids_gen) == 0:
                                st.error("No grids configured for weather portfolio.")
                            elif len(analog_year_list) == 0:
                                st.error("No analog years found. Please run 'Find Analog Years' first.")
                            else:
                                with st.spinner(f"Optimizing {len(weather_grids_gen)} grids for {len(analog_year_list)} analog years..."):
                                    weather_challenger_allocations = {}
                                    weather_challenger_acres = {}
                                    optimization_iterations = weather2_iterations

                                    progress_bar = st.progress(0)
                                    for idx, gid in enumerate(weather_grids_gen):
                                        progress_bar.progress((idx + 1) / len(weather_grids_gen))

                                        grid_acres = weather_acres_gen.get(gid, total_insured_acres // len(weather_grids_gen))

                                        best_alloc, best_roi, _ = run_analog_year_optimization(
                                            session=session,
                                            grid_id=gid,
                                            analog_years=analog_year_list,
                                            plan_code=plan_code,
                                            productivity_factor=productivity_factor,
                                            acres=grid_acres,
                                            intended_use=intended_use,
                                            coverage_level=coverage_level,
                                            iterations=optimization_iterations,
                                            interval_range_opt=interval_range_gen
                                        )

                                        weather_challenger_allocations[gid] = best_alloc
                                        weather_challenger_acres[gid] = grid_acres

                                    progress_bar.empty()

                                # Backtest on ALL years for fair comparison
                                with st.spinner("Backtesting weather challenger on all years..."):
                                    weather_df, weather_grid_results, weather_metrics = run_portfolio_backtest(
                                        session=session,
                                        selected_grids=weather_grids_gen,
                                        grid_allocations=weather_challenger_allocations,
                                        grid_acres=weather_challenger_acres,
                                        start_year=start_year,
                                        end_year=end_year,
                                        coverage_level=coverage_level,
                                        productivity_factor=productivity_factor,
                                        intended_use=intended_use,
                                        plan_code=plan_code,
                                        scenario='All Years (except Current Year)'
                                    )

                                # Store results
                                st.session_state.weather_challenger_results = {
                                    'allocations': weather_challenger_allocations,
                                    'acres': weather_challenger_acres,
                                    'grids': weather_grids_gen,
                                    'metrics': weather_metrics,
                                    'df': weather_df,
                                    'analog_years_used': analog_year_list,
                                    'methodology': 'naive'
                                }

                                st.success(f"Weather Challenger 2 generated! Optimized on {len(analog_year_list)} analog years, tested on all years.")

                        # Display Weather Challenger 2 Results (if they exist)
                        if 'weather_challenger_results' in st.session_state and st.session_state.weather_challenger_results:
                            weather_results = st.session_state.weather_challenger_results

                            st.markdown("---")
                            with st.container(border=True):
                                st.markdown("### 🌦️ Challenger 2: Weather Naive")

                                # Methodology note
                                analog_count = len(weather_results.get('analog_years_used', []))
                                # Build full criteria label
                                config = st.session_state.get('ps_weather_config', {})
                                criteria_parts = []
                                if config.get('enso_regime', 'Any') != 'Any':
                                    criteria_parts.append(config.get('enso_regime'))
                                if config.get('historical_context', 'Any') != 'Any':
                                    criteria_parts.append(config.get('historical_context'))
                                if config.get('trajectory', 'Any') != 'Any':
                                    criteria_parts.append(config.get('trajectory'))
                                full_criteria = " + ".join(criteria_parts) if criteria_parts else "All Conditions"

                                st.info(f"**Methodology:** Weather Challenger 2 optimizes each grid independently for the **{analog_count} analog years** matching: **{full_criteria}**. No cross-grid correlation is considered - this is a 'naive' approach.")

                                # 3-Way Performance Comparison
                                if ('champion_results' in st.session_state and st.session_state.champion_results and
                                    'challenger_results' in st.session_state and st.session_state.challenger_results):

                                    st.markdown("#### 3-Way Performance Comparison")
                                    champ_metrics = st.session_state.champion_results.get('metrics', {})
                                    chall1_metrics = st.session_state.challenger_results.get('metrics', {})
                                    weather_metrics_disp = weather_results.get('metrics', {})

                                    comparison_3way_df = create_3way_comparison_table(champ_metrics, chall1_metrics, weather_metrics_disp)
                                    st.table(comparison_3way_df.set_index('Metric'))

                                    # Winner determination
                                    champ_roi = champ_metrics.get('cumulative_roi', 0)
                                    chall1_roi = chall1_metrics.get('cumulative_roi', 0)
                                    weather_roi = weather_metrics_disp.get('cumulative_roi', 0)

                                    best_roi = max(champ_roi, chall1_roi, weather_roi)
                                    if weather_roi == best_roi and weather_roi > champ_roi:
                                        st.success(f"**WEATHER CHALLENGER 2 WINS!** ROI: {weather_roi:.1%}")
                                    elif chall1_roi == best_roi and chall1_roi > champ_roi:
                                        st.success(f"**CHALLENGER 1 WINS!** ROI: {chall1_roi:.1%}")
                                    elif champ_roi == best_roi:
                                        st.warning(f"**CHAMPION HOLDS!** ROI: {champ_roi:.1%}")

                                # Weather Challenger Allocation Table
                                st.markdown("#### Weather Challenger 2 Allocations")
                                weather_styled, weather_alloc_df = create_optimized_allocation_table(
                                    weather_results['allocations'],
                                    weather_results['grids'],
                                    grid_acres=weather_results['acres'],
                                    label="WEATHER AVERAGE"
                                )
                                st.dataframe(weather_styled, use_container_width=True, hide_index=True)

                                st.download_button(
                                    label="📥 Download Weather Challenger Allocations CSV",
                                    data=weather_alloc_df.to_csv(index=False),
                                    file_name="weather_challenger_allocations.csv",
                                    mime="text/csv",
                                    key="download_weather_csv"
                                )

                        # === WEATHER CHALLENGER 3: Correlated Weather Portfolio ===
                        st.markdown("---")
                        st.markdown("#### 🌪️ Challenger 3: Correlation-Aware (MVO)")
                        st.caption("Two-stage optimization: (1) Optimize intervals per grid on analog years, (2) MVO redistributes acres based on cross-grid correlations.")

                        # --- Premium Budget ---
                        st.markdown("**Premium Budget**")
                        weather3_enable_budget = st.checkbox(
                            "Set annual premium budget",
                            value=False,
                            help="Limit total annual premium spending for Challenger 3",
                            key="ps_weather3_budget"
                        )

                        weather3_annual_budget = 50000
                        weather3_allow_scale_up = False
                        if weather3_enable_budget:
                            budget_col1, budget_col2 = st.columns(2)
                            with budget_col1:
                                weather3_annual_budget = st.number_input(
                                    "Maximum Annual Premium ($)",
                                    min_value=1000,
                                    value=50000,
                                    step=1000,
                                    help="Maximum producer premium (after subsidy) per year",
                                    key="ps_weather3_budget_amt"
                                )
                            with budget_col2:
                                weather3_allow_scale_up = st.checkbox(
                                    "Auto-fill budget (scale up if under)",
                                    value=False,
                                    help="If the optimized strategy comes in under budget, automatically scale up acres to utilize the full budget.",
                                    key="ps_weather3_autofill"
                                )

                        st.markdown("---")

                        # --- Interval Strategy ---
                        st.markdown("**Interval Strategy**")

                        interval_strat_col1, interval_strat_col2 = st.columns(2)

                        with interval_strat_col1:
                            weather3_iteration_map = {'Fast': 500, 'Standard': 3000, 'Thorough': 7000, 'Maximum': 15000}
                            weather3_search_depth_key = st.select_slider(
                                "Search Depth",
                                options=list(weather3_iteration_map.keys()),
                                value='Standard',
                                key="ps_weather3_depth"
                            )
                            weather3_search_iterations = weather3_iteration_map[weather3_search_depth_key]

                        with interval_strat_col2:
                            weather3_objective = st.selectbox(
                                "Optimization Objective",
                                options=["Cumulative ROI", "Median ROI", "Win Rate", "Risk-Adjusted Return"],
                                index=0,
                                help="What metric to maximize when selecting the best interval allocation. Cumulative ROI = total returns, Median ROI = typical year, Win Rate = % profitable years, Risk-Adjusted = return per unit risk.",
                                key="ps_weather3_objective"
                            )
                            # Map display names to internal values
                            weather3_objective_map = {
                                "Cumulative ROI": "cumulative_roi",
                                "Median ROI": "median_roi",
                                "Win Rate": "profitable_pct",
                                "Risk-Adjusted Return": "risk_adj_ret"
                            }
                            weather3_objective_value = weather3_objective_map[weather3_objective]

                        # --- Diversification Constraints ---
                        st.markdown("**Diversification Constraints**")
                        div_col1, div_col2 = st.columns(2)

                        with div_col1:
                            weather3_require_full_coverage = st.checkbox(
                                "Ensure Full Calendar Coverage",
                                value=False,
                                help="Creates staggered portfolio with Pattern A (6 intervals) and Pattern B (5 intervals) to cover all 11 intervals while maintaining non-adjacency. Requires at least 2 grids.",
                                key="ps_weather3_full_coverage"
                            )

                        with div_col2:
                            if not weather3_require_full_coverage:
                                weather3_interval_range = st.slider(
                                    "Number of Active Intervals Range",
                                    min_value=2,
                                    max_value=6,
                                    value=(2, 6),
                                    help="Optimizer can choose ANY number of intervals within this range. Each interval must have 10%-50%.",
                                    key="ps_weather3_interval_range"
                                )
                            else:
                                weather3_interval_range = (5, 6)  # Full coverage uses Pattern A/B
                                st.info("Full coverage mode: Pattern A (6) + Pattern B (5)")

                        st.markdown("---")

                        # --- Acreage Strategy ---
                        st.markdown("**Acreage Strategy**")

                        weather3_optimize_acreage = st.checkbox(
                            "Optimize acreage distribution (Mean-Variance Optimization)",
                            value=True,  # Default ON for Challenger 3 since that's its purpose
                            help="Use Markowitz portfolio theory to redistribute acres based on risk-adjusted returns during analog years",
                            key="ps_weather3_mvo"
                        )

                        weather3_risk_aversion = 1.0
                        weather3_max_turnover = 0.20
                        if weather3_optimize_acreage:
                            mvo_col1, mvo_col2 = st.columns(2)
                            with mvo_col1:
                                weather3_risk_profile = st.select_slider(
                                    "Risk Profile",
                                    options=["Aggressive", "Growth", "Balanced", "Conservative", "Defensive"],
                                    value="Balanced",
                                    help="Aggressive = chase highest returns, accept volatility. Defensive = prioritize stability and diversification.",
                                    key="ps_weather3_risk_profile"
                                )

                                weather3_risk_map = {
                                    "Aggressive": 0.5,
                                    "Growth": 0.75,
                                    "Balanced": 1.0,
                                    "Conservative": 1.5,
                                    "Defensive": 2.0
                                }
                                weather3_risk_aversion = weather3_risk_map[weather3_risk_profile]

                            with mvo_col2:
                                weather3_max_turnover = st.slider(
                                    "Max Turnover",
                                    min_value=0,
                                    max_value=100,
                                    value=20,
                                    step=5,
                                    format="%d%%",
                                    help="How much each grid's acre allocation can change. 0% = no changes, 100% = full reallocation allowed.",
                                    key="ps_weather3_turnover"
                                ) / 100.0  # Convert back to decimal for calculations

                        st.markdown("---")

                        # --- Generate Button ---
                        if st.button("🌪️ Generate Challenger 3 (MVO)", key="ps_generate_weather_3", type="primary"):
                            weather_config = st.session_state.get('ps_weather_config', {})
                            weather_grids_gen = weather_config.get('grids', [])
                            weather_acres_gen = weather_config.get('acres', {})

                            analog_year_list = [y['year'] for y in st.session_state.get('ps_analog_years', [])]

                            if len(weather_grids_gen) < 2:
                                st.error("Challenger 3 requires at least 2 grids for correlation analysis.")
                            elif len(analog_year_list) == 0:
                                st.error("No analog years found. Please run 'Find Analog Years' first.")
                            else:
                                with st.spinner(f"Running 2-stage optimization for {len(weather_grids_gen)} grids..."):

                                    # === STAGE 1: Interval Optimization ===
                                    st.write("**Stage 1: Optimizing Interval Allocations...**")
                                    progress_bar = st.progress(0, text="Starting interval optimization...")

                                    weather3_allocations = {}
                                    weather3_interval_stats = {}

                                    for idx, gid in enumerate(weather_grids_gen):
                                        progress_bar.progress(
                                            (idx + 1) / len(weather_grids_gen),
                                            text=f"Optimizing intervals for {gid}..."
                                        )

                                        grid_acres = weather_acres_gen.get(gid, 1000)

                                        # Use run_analog_year_optimization with full coverage support
                                        if weather3_require_full_coverage:
                                            # For full coverage, we need to use the pattern-based approach
                                            best_alloc, best_roi, tested = run_fast_optimization_core(
                                                session, gid, min(analog_year_list), max(analog_year_list), plan_code,
                                                productivity_factor, grid_acres, intended_use,
                                                coverage_level, weather3_search_iterations, 'global',
                                                require_full_coverage=True,
                                                interval_range_opt=weather3_interval_range,
                                                grid_index=idx
                                            )
                                        else:
                                            best_alloc, best_roi, tested = run_analog_year_optimization(
                                                session=session,
                                                grid_id=gid,
                                                analog_years=analog_year_list,
                                                plan_code=plan_code,
                                                productivity_factor=productivity_factor,
                                                acres=grid_acres,
                                                intended_use=intended_use,
                                                coverage_level=coverage_level,
                                                iterations=weather3_search_iterations,
                                                interval_range_opt=weather3_interval_range,
                                                objective=weather3_objective_value
                                            )

                                        weather3_allocations[gid] = best_alloc
                                        weather3_interval_stats[gid] = {'roi': best_roi, 'tested': tested}

                                    progress_bar.empty()
                                    st.success(f"Stage 1 complete! Tested {sum(s['tested'] for s in weather3_interval_stats.values()):,} strategies.")

                                    # === STAGE 2: Acre Optimization using optimize_grid_allocation() ===
                                    # This uses the same proven 3-stage optimization as Challenger 1:
                                    # Stage 1: Budget Scaling - scale all grids proportionally to meet budget
                                    # Stage 2: MVO Rebalancing - optimize within turnover bounds of scaled baseline
                                    # Stage 3: Budget Fill-up - scale up to hit 99.95% of budget when auto-fill enabled

                                    weather3_acres = weather_acres_gen.copy()
                                    analog_roi_correlation = pd.DataFrame()
                                    initial_total_acres = sum(weather_acres_gen.values())
                                    budget_scaled_acres = weather_acres_gen.copy()
                                    opt_info = {}

                                    # Build grid_results structure for MVO (same format as Challenger 1)
                                    grid_results_for_mvo = {
                                        gid: {'best_strategy': {'allocation': weather3_allocations[gid], 'coverage_level': coverage_level}}
                                        for gid in weather_grids_gen
                                    }

                                    # Build ROI data for analog years (needed for MVO optimization)
                                    if weather3_optimize_acreage or weather3_enable_budget:
                                        st.write("**Stage 2: Building ROI correlation data for analog years...**")

                                        analog_roi_data = []

                                        for gid in weather_grids_gen:
                                            allocation = weather3_allocations.get(gid, {})
                                            if not allocation:
                                                st.warning(f"Grid {gid} skipped: empty allocation from Stage 1")
                                                continue

                                            try:
                                                subsidy = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                                                county_base_value = load_county_base_value(session, gid)
                                                current_rate_year = get_current_rate_year(session)
                                                premium_rates = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]

                                                dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)

                                                all_indices_df = load_all_indices(session, gid)
                                                all_indices_df = all_indices_df[all_indices_df['YEAR'].isin(analog_year_list)]

                                                for year in analog_year_list:
                                                    year_data = all_indices_df[all_indices_df['YEAR'] == year]
                                                    if year_data.empty:
                                                        continue

                                                    year_indemnity = 0
                                                    year_premium = 0

                                                    for interval, pct in allocation.items():
                                                        if pct == 0:
                                                            continue

                                                        index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                                                        index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

                                                        premium_rate = premium_rates.get(interval, 0)
                                                        interval_protection = int(round_half_up(dollar_protection * 1 * pct, 0))
                                                        total_prem = int(round_half_up(interval_protection * premium_rate, 0))
                                                        prem_subsidy = int(round_half_up(total_prem * subsidy, 0))
                                                        producer_premium = total_prem - prem_subsidy

                                                        trigger = coverage_level * 100
                                                        shortfall = max(0, (trigger - index_value) / trigger)
                                                        raw_indemnity = shortfall * interval_protection
                                                        # Convert to int immediately to ensure exact integer arithmetic when summing
                                                        indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                                                        year_indemnity += indemnity
                                                        year_premium += producer_premium

                                                    # Both indemnity and premium are already exact integer sums

                                                    roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                                                    analog_roi_data.append({'year': year, 'grid': gid, 'roi': roi})

                                            except Exception as e:
                                                st.warning(f"Grid {gid} skipped: {e}")
                                                continue

                                        if len(analog_roi_data) > 0:
                                            analog_roi_df = pd.DataFrame(analog_roi_data)

                                            # Build correlation matrix for display
                                            pivot_df = analog_roi_df.pivot_table(
                                                values='roi',
                                                index='year',
                                                columns='grid'
                                            )
                                            analog_roi_correlation = pivot_df.corr()

                                            # Now use the proven optimize_grid_allocation function
                                            if weather3_enable_budget and weather3_optimize_acreage:
                                                # Full 3-stage optimization: Budget + MVO + Fill-up
                                                st.write("**Running 3-stage optimization (Budget Scaling → MVO → Budget Fill-up)...**")

                                                weather3_acres, analog_roi_correlation, opt_info = optimize_grid_allocation(
                                                    analog_roi_df, grid_results_for_mvo, weather_acres_gen,
                                                    weather3_annual_budget, session, productivity_factor, intended_use, plan_code,
                                                    weather_grids_gen, weather3_risk_aversion, weather3_max_turnover, weather3_allow_scale_up
                                                )

                                                # Display optimization feedback
                                                scale_factor = opt_info.get('budget_scale_factor', 1.0)
                                                initial_total = opt_info.get('initial_total_acres', 0)
                                                scaled_total = opt_info.get('budget_scaled_acres', 0)
                                                if scale_factor != 1.0:
                                                    st.info(f"**Stage 1 - Budget scaling:** {initial_total:,.0f} → {scaled_total:,.0f} acres ({scale_factor:.1%} of original)")
                                                st.info(f"**Stage 2 - MVO rebalancing:** Grid weights adjusted within ±{weather3_max_turnover:.0%} turnover bounds")
                                                if opt_info.get('stage3_applied', False):
                                                    stage3_scale = opt_info.get('stage3_scale_factor', 1.0)
                                                    st.success(f"**Stage 3 - Budget fill-up:** Acres scaled up by {stage3_scale:.1%} to utilize full budget")

                                                # Save budget-scaled acres (before MVO) for display
                                                budget_scaled_acres = {
                                                    gid: weather_acres_gen.get(gid, 0) * scale_factor
                                                    for gid in weather_grids_gen
                                                }

                                            elif weather3_enable_budget and not weather3_optimize_acreage:
                                                # Budget constraint only (no MVO) - proportional scaling
                                                st.write("**Applying budget constraint (proportional scaling)...**")

                                                initial_cost, _ = calculate_annual_premium_cost(
                                                    session, weather_grids_gen, weather_acres_gen, grid_results_for_mvo,
                                                    productivity_factor, intended_use, plan_code
                                                )

                                                if initial_cost > weather3_annual_budget:
                                                    budget_scale_factor = (weather3_annual_budget * 0.9995) / initial_cost
                                                elif weather3_allow_scale_up and initial_cost < weather3_annual_budget:
                                                    budget_scale_factor = (weather3_annual_budget * 0.9995) / initial_cost
                                                else:
                                                    budget_scale_factor = 1.0

                                                weather3_acres = {
                                                    gid: acres * budget_scale_factor
                                                    for gid, acres in weather_acres_gen.items()
                                                }
                                                budget_scaled_acres = weather3_acres.copy()

                                                if budget_scale_factor != 1.0:
                                                    st.info(f"**Budget scaling:** {initial_total_acres:,.0f} → {sum(weather3_acres.values()):,.0f} acres ({budget_scale_factor:.1%} of original)")

                                            elif weather3_optimize_acreage and not weather3_enable_budget:
                                                # MVO only (no budget constraint)
                                                st.write("**Running MVO optimization (no budget constraint)...**")

                                                weather3_acres = optimize_without_budget(
                                                    analog_roi_df, grid_results_for_mvo, initial_total_acres,
                                                    weather_grids_gen, weather3_risk_aversion, weather3_max_turnover, weather_acres_gen
                                                )
                                                budget_scaled_acres = weather3_acres.copy()

                                                st.info(f"**MVO rebalancing:** Grid weights adjusted within ±{weather3_max_turnover:.0%} turnover bounds")

                                        else:
                                            st.warning("No ROI data available for analog years. Using original acre allocation.")

                                    # === STAGE 4: Final Backtest ===
                                    st.write("**Stage 4: Running Final Backtest...**")

                                    weather3_df, weather3_grid_results, weather3_metrics = run_portfolio_backtest(
                                        session=session,
                                        selected_grids=weather_grids_gen,
                                        grid_allocations=weather3_allocations,
                                        grid_acres=weather3_acres,
                                        start_year=start_year,
                                        end_year=end_year,
                                        coverage_level=coverage_level,
                                        productivity_factor=productivity_factor,
                                        intended_use=intended_use,
                                        plan_code=plan_code,
                                        scenario='All Years (except Current Year)'
                                    )

                                    st.session_state.weather_challenger_3_results = {
                                        'allocations': weather3_allocations,
                                        'acres': weather3_acres,
                                        'initial_acres': weather_acres_gen.copy(),
                                        'budget_scaled_acres': budget_scaled_acres,
                                        'grids': weather_grids_gen,
                                        'metrics': weather3_metrics,
                                        'df': weather3_df,
                                        'analog_years_used': analog_year_list,
                                        'analog_roi_correlation': analog_roi_correlation if weather3_optimize_acreage else pd.DataFrame(),
                                        'interval_stats': weather3_interval_stats,
                                        'methodology': 'mvo' if weather3_optimize_acreage else 'naive',
                                        'risk_aversion': weather3_risk_aversion,
                                        'max_turnover': weather3_max_turnover,
                                        'budget_enabled': weather3_enable_budget,
                                        'budget_amount': weather3_annual_budget if weather3_enable_budget else None
                                    }

                                    st.success(f"Challenger 3 (MVO) generated! Optimized on {len(analog_year_list)} analog years.")

                        # Display Weather Challenger 3 Results
                        if 'weather_challenger_3_results' in st.session_state and st.session_state.weather_challenger_3_results:
                            weather3 = st.session_state.weather_challenger_3_results

                            st.markdown("---")
                            with st.container(border=True):
                                st.markdown("### 🌪️ Challenger 3: Weather MVO")

                                # Methodology note
                                analog_count_3 = len(weather3.get('analog_years_used', []))
                                budget_text = ""
                                if weather3.get('budget_enabled'):
                                    budget_text = f" | Budget: ${weather3.get('budget_amount', 0):,.0f}"

                                # Build full criteria label for Challenger 3
                                config3 = st.session_state.get('ps_weather_config', {})
                                criteria_parts_3 = []
                                if config3.get('enso_regime', 'Any') != 'Any':
                                    criteria_parts_3.append(config3.get('enso_regime'))
                                if config3.get('historical_context', 'Any') != 'Any':
                                    criteria_parts_3.append(config3.get('historical_context'))
                                if config3.get('trajectory', 'Any') != 'Any':
                                    criteria_parts_3.append(config3.get('trajectory'))
                                full_criteria_3 = " + ".join(criteria_parts_3) if criteria_parts_3 else "All Conditions"

                                st.info(f"**Methodology:** Challenger 3 uses two-stage optimization on **{analog_count_3} analog years** matching: **{full_criteria_3}**. (1) Optimize intervals per grid, (2) MVO redistributes acres based on cross-grid correlations. Risk aversion: {weather3.get('risk_aversion', 1.0)}{budget_text}")

                                # 4-Way Performance Comparison
                                if ('champion_results' in st.session_state and st.session_state.champion_results and
                                    'challenger_results' in st.session_state and st.session_state.challenger_results and
                                    'weather_challenger_results' in st.session_state and st.session_state.weather_challenger_results):

                                    st.markdown("#### 4-Way Performance Comparison")

                                    champ_m = st.session_state.champion_results.get('metrics', {})
                                    chall1_m = st.session_state.challenger_results.get('metrics', {})
                                    weather2_m = st.session_state.weather_challenger_results.get('metrics', {})
                                    weather3_m = weather3.get('metrics', {})

                                    comparison_4way = {
                                        'Metric': ['Cumulative ROI', 'Risk-Adjusted Return', 'Est. Annual Premium', 'Win Rate'],
                                        'Champion': [
                                            f"{champ_m.get('cumulative_roi', 0):.1%}",
                                            f"{champ_m.get('risk_adj_return', 0):.2f}",
                                            f"${champ_m.get('avg_annual_premium', 0):,.0f}",
                                            f"{champ_m.get('profitable_pct', 0):.0%}"
                                        ],
                                        'Challenger 1': [
                                            f"{chall1_m.get('cumulative_roi', 0):.1%}",
                                            f"{chall1_m.get('risk_adj_return', 0):.2f}",
                                            f"${chall1_m.get('avg_annual_premium', 0):,.0f}",
                                            f"{chall1_m.get('profitable_pct', 0):.0%}"
                                        ],
                                        'Challenger 2 (Naive)': [
                                            f"{weather2_m.get('cumulative_roi', 0):.1%}",
                                            f"{weather2_m.get('risk_adj_return', 0):.2f}",
                                            f"${weather2_m.get('avg_annual_premium', 0):,.0f}",
                                            f"{weather2_m.get('profitable_pct', 0):.0%}"
                                        ],
                                        'Challenger 3 (MVO)': [
                                            f"{weather3_m.get('cumulative_roi', 0):.1%}",
                                            f"{weather3_m.get('risk_adj_return', 0):.2f}",
                                            f"${weather3_m.get('avg_annual_premium', 0):,.0f}",
                                            f"{weather3_m.get('profitable_pct', 0):.0%}"
                                        ]
                                    }

                                    st.table(pd.DataFrame(comparison_4way).set_index('Metric'))

                                    # Winner determination
                                    all_rois = {
                                        'Champion': champ_m.get('cumulative_roi', 0),
                                        'Challenger 1': chall1_m.get('cumulative_roi', 0),
                                        'Challenger 2 (Naive)': weather2_m.get('cumulative_roi', 0),
                                        'Challenger 3 (MVO)': weather3_m.get('cumulative_roi', 0)
                                    }
                                    winner = max(all_rois, key=all_rois.get)
                                    winner_roi = all_rois[winner]

                                    if winner == 'Challenger 3 (MVO)':
                                        st.success(f"**CHALLENGER 3 (MVO) WINS!** ROI: {winner_roi:.1%}")
                                    elif winner == 'Challenger 2 (Naive)':
                                        st.success(f"**CHALLENGER 2 (NAIVE) WINS!** ROI: {winner_roi:.1%}")
                                    elif winner == 'Challenger 1':
                                        st.success(f"**CHALLENGER 1 WINS!** ROI: {winner_roi:.1%}")
                                    else:
                                        st.warning(f"**CHAMPION HOLDS!** ROI: {winner_roi:.1%}")

                                # Acre Redistribution Analysis - Show all 3 stages
                                st.markdown("#### Acre Redistribution (MVO Impact)")

                                max_turnover_used = weather3.get('max_turnover', 0.20)
                                st.caption(f"MVO turnover constraint: ±{max_turnover_used:.0%} relative to budget-scaled baseline")

                                acre_comparison = []
                                for gid in weather3['grids']:
                                    initial = weather3['initial_acres'].get(gid, 0)
                                    budget_scaled = weather3.get('budget_scaled_acres', weather3['initial_acres']).get(gid, initial)
                                    mvo_optimized = weather3['acres'].get(gid, 0)

                                    # Budget Scale % = change from Initial to Budget Scaled
                                    budget_scale_pct = ((budget_scaled - initial) / initial * 100) if initial > 0 else 0

                                    # MVO Turnover % = change from Budget Scaled to MVO Optimized
                                    # This should respect the ±max_turnover bound
                                    mvo_turnover_pct = ((mvo_optimized - budget_scaled) / budget_scaled * 100) if budget_scaled > 0 else 0

                                    acre_comparison.append({
                                        'Grid': gid,
                                        'Initial Acres': f"{initial:,.0f}",
                                        'Budget Scaled': f"{budget_scaled:,.0f}",
                                        'MVO Optimized': f"{mvo_optimized:,.0f}",
                                        'Budget Scale %': f"{budget_scale_pct:+.1f}%",
                                        'MVO Turnover %': f"{mvo_turnover_pct:+.1f}%"
                                    })

                                acre_df = pd.DataFrame(acre_comparison)
                                st.dataframe(acre_df, use_container_width=True, hide_index=True)

                                # Analog Year Correlation Matrix
                                if not weather3.get('analog_roi_correlation', pd.DataFrame()).empty:
                                    with st.expander("📉 Analog Year ROI Correlations", expanded=False):
                                        st.caption("Correlations based on performance during analog years only. Lower = better diversification.")

                                        corr_df = weather3['analog_roi_correlation']

                                        # Generate heatmap matching Historical Grid Correlations style
                                        fig, ax = plt.subplots(figsize=(10, 6))
                                        sns.heatmap(
                                            corr_df,
                                            annot=True,
                                            cmap='RdYlGn',
                                            fmt=".3f",
                                            vmin=-1,
                                            vmax=1,
                                            center=0,
                                            ax=ax,
                                            square=True,
                                            linewidths=0.5
                                        )
                                        ax.set_title("Analog Year ROI Correlations", fontsize=12)
                                        plt.tight_layout()

                                        st.pyplot(fig)
                                        plt.close(fig)

                                # Allocation Table
                                st.markdown("#### Weather Challenger 3 Allocations")
                                weather3_styled, weather3_alloc_df = create_optimized_allocation_table(
                                    weather3['allocations'],
                                    weather3['grids'],
                                    grid_acres=weather3['acres'],
                                    label="WEATHER MVO AVERAGE"
                                )
                                st.dataframe(weather3_styled, use_container_width=True, hide_index=True)

                                st.download_button(
                                    label="📥 Download Weather Challenger 3 CSV",
                                    data=weather3_alloc_df.to_csv(index=False),
                                    file_name="weather_challenger_3_allocations.csv",
                                    mime="text/csv",
                                    key="download_weather3_csv"
                                )

                        # ==============================================================
                        # SCENARIO STRESS TEST
                        # ==============================================================
                        # Only show if all 4 strategies exist
                        if ('champion_results' in st.session_state and st.session_state.champion_results and
                            'challenger_results' in st.session_state and st.session_state.challenger_results and
                            'weather_challenger_results' in st.session_state and st.session_state.weather_challenger_results and
                            'weather_challenger_3_results' in st.session_state and st.session_state.weather_challenger_3_results):

                            st.divider()
                            st.markdown("### 📊 Scenario Stress Test")
                            st.caption("Re-run the 4-way comparison under different conditions to validate your weather thesis.")

                            # Get analog years from session state for the radio option
                            analog_year_list_stress = [y['year'] for y in st.session_state.get('ps_analog_years', [])]
                            weather_config_stress = st.session_state.get('ps_weather_config', {})

                            # Build scenario description for analog years with ALL THREE criteria
                            enso_text_stress = weather_config_stress.get('enso_regime', 'Any')
                            context_text_stress = weather_config_stress.get('historical_context', 'Any')
                            trend_text_stress = weather_config_stress.get('trajectory', 'Any')

                            criteria_stress = []
                            if enso_text_stress != 'Any':
                                criteria_stress.append(enso_text_stress)
                            if context_text_stress != 'Any':
                                criteria_stress.append(context_text_stress)
                            if trend_text_stress != 'Any':
                                criteria_stress.append(trend_text_stress)

                            analog_label = f"Analog Years Only ({' + '.join(criteria_stress)})" if criteria_stress else "Analog Years Only (All Conditions)"

                            stress_scenario = st.radio(
                                "Select scenario to test:",
                                options=[
                                    "All Years (except Current Year)",
                                    "La Nina Years Only",
                                    "El Nino Years Only",
                                    "Neutral Years Only",
                                    analog_label,
                                    "Custom Range"
                                ],
                                index=0,
                                key="ps_stress_scenario"
                            )

                            # Custom range inputs (only show if selected)
                            stress_start_year = start_year
                            stress_end_year = end_year
                            if stress_scenario == "Custom Range":
                                col1, col2 = st.columns(2)
                                stress_start_year = col1.selectbox("Start Year", list(range(1948, 2026)), index=62, key="ps_stress_start")
                                stress_end_year = col2.selectbox("End Year", list(range(1948, 2026)), index=76, key="ps_stress_end")

                            if st.button("🔬 Run Stress Test", key="ps_run_stress_test", type="secondary"):
                                # Map radio selection to scenario filter
                                scenario_map = {
                                    "All Years (except Current Year)": "All Years (except Current Year)",
                                    "La Nina Years Only": "ENSO Phase: La Nina",
                                    "El Nino Years Only": "ENSO Phase: El Nino",
                                    "Neutral Years Only": "ENSO Phase: Neutral",
                                    analog_label: "Analog Years",
                                    "Custom Range": "Custom Range"
                                }

                                test_scenario = scenario_map.get(stress_scenario, "All Years (except Current Year)")

                                with st.spinner(f"Running stress test: {stress_scenario}..."):

                                    # Get stored allocations and acres for each strategy
                                    champ_stress = st.session_state.champion_results
                                    chall1_stress = st.session_state.challenger_results
                                    chall2_stress = st.session_state.weather_challenger_results
                                    chall3_stress = st.session_state.weather_challenger_3_results

                                    # Helper function to run backtest for a strategy under the stress scenario
                                    def run_stress_backtest(allocations, acres, grids, scenario_type):
                                        df, grid_results, metrics = run_portfolio_backtest(
                                            session=session,
                                            selected_grids=grids,
                                            grid_allocations=allocations,
                                            grid_acres=acres,
                                            start_year=stress_start_year if scenario_type == "Custom Range" else 1948,
                                            end_year=stress_end_year if scenario_type == "Custom Range" else 2024,
                                            coverage_level=coverage_level,
                                            productivity_factor=productivity_factor,
                                            intended_use=intended_use,
                                            plan_code=plan_code,
                                            scenario=test_scenario if scenario_type != analog_label else "Analog Years"
                                        )

                                        return metrics

                                    # Run stress test for each strategy
                                    stress_results = {}

                                    stress_results['Champion'] = run_stress_backtest(
                                        champ_stress['allocations'], champ_stress['acres'], champ_stress.get('grids', selected_grids),
                                        stress_scenario
                                    )

                                    stress_results['Challenger 1'] = run_stress_backtest(
                                        chall1_stress['allocations'], chall1_stress['acres'], chall1_stress.get('grids', selected_grids),
                                        stress_scenario
                                    )

                                    stress_results['Challenger 2'] = run_stress_backtest(
                                        chall2_stress['allocations'], chall2_stress['acres'], chall2_stress.get('grids', []),
                                        stress_scenario
                                    )

                                    stress_results['Challenger 3'] = run_stress_backtest(
                                        chall3_stress['allocations'], chall3_stress['acres'], chall3_stress.get('grids', []),
                                        stress_scenario
                                    )

                                    st.session_state.stress_test_results = {
                                        'scenario': stress_scenario,
                                        'results': stress_results,
                                        'analog_years_count': len(analog_year_list_stress) if stress_scenario == analog_label else None
                                    }

                                st.success("Stress test complete!")

                            # Display Stress Test Results
                            if 'stress_test_results' in st.session_state and st.session_state.stress_test_results:
                                stress = st.session_state.stress_test_results
                                results = stress['results']

                                st.markdown(f"#### Results: {stress['scenario']}")

                                # Get baseline (All Years) metrics for delta calculation
                                champ_baseline = st.session_state.champion_results.get('metrics', {})
                                chall1_baseline = st.session_state.challenger_results.get('metrics', {})
                                chall2_baseline = st.session_state.weather_challenger_results.get('metrics', {})
                                chall3_baseline = st.session_state.weather_challenger_3_results.get('metrics', {})

                                # Build comparison table
                                def format_roi(val):
                                    return f"{val:.1%}" if val else "N/A"

                                def format_delta(stress_val, baseline_val):
                                    if stress_val is None or baseline_val is None:
                                        return "N/A"
                                    delta = stress_val - baseline_val
                                    if delta > 0:
                                        return f"+{delta:.1%}"
                                    else:
                                        return f"{delta:.1%}"

                                champ_roi = results['Champion'].get('cumulative_roi', 0)
                                chall1_roi = results['Challenger 1'].get('cumulative_roi', 0)
                                chall2_roi = results['Challenger 2'].get('cumulative_roi', 0)
                                chall3_roi = results['Challenger 3'].get('cumulative_roi', 0)

                                # Create scenario label for table
                                scenario_short = stress['scenario'].replace(" Only", "").replace(" (Current Results)", "")

                                # Get the baseline scenario from session state (what Champion was trained on)
                                baseline_scenario = st.session_state.get('ps_scenario', 'All Years (except Current Year)')
                                baseline_label = baseline_scenario.replace('ENSO Phase: ', '')

                                comparison_data = {
                                    'Metric': [
                                        f'ROI ({baseline_label})',
                                        f'ROI ({scenario_short})',
                                        f'Change vs {baseline_label}',
                                        'Change vs Champion'
                                    ],
                                    'Champion': [
                                        format_roi(champ_baseline.get('cumulative_roi', 0)),
                                        format_roi(champ_roi),
                                        format_delta(champ_roi, champ_baseline.get('cumulative_roi', 0)),
                                        "—"
                                    ],
                                    'Challenger 1': [
                                        format_roi(chall1_baseline.get('cumulative_roi', 0)),
                                        format_roi(chall1_roi),
                                        format_delta(chall1_roi, chall1_baseline.get('cumulative_roi', 0)),
                                        format_delta(chall1_roi, champ_roi)
                                    ],
                                    'Challenger 2': [
                                        format_roi(chall2_baseline.get('cumulative_roi', 0)),
                                        format_roi(chall2_roi),
                                        format_delta(chall2_roi, chall2_baseline.get('cumulative_roi', 0)),
                                        format_delta(chall2_roi, champ_roi)
                                    ],
                                    'Challenger 3': [
                                        format_roi(chall3_baseline.get('cumulative_roi', 0)),
                                        format_roi(chall3_roi),
                                        format_delta(chall3_roi, chall3_baseline.get('cumulative_roi', 0)),
                                        format_delta(chall3_roi, champ_roi)
                                    ]
                                }

                                comparison_df = pd.DataFrame(comparison_data)
                                st.table(comparison_df.set_index('Metric'))

                                # Generate thesis validation insight
                                st.markdown("#### 💡 Thesis Validation")

                                # Find best performer under stress scenario
                                stress_rois = {
                                    'Champion': champ_roi,
                                    'Challenger 1': chall1_roi,
                                    'Challenger 2': chall2_roi,
                                    'Challenger 3': chall3_roi
                                }
                                best_stress = max(stress_rois, key=stress_rois.get)
                                best_stress_roi = stress_rois[best_stress]

                                # Find best performer under all years
                                baseline_rois = {
                                    'Champion': champ_baseline.get('cumulative_roi', 0),
                                    'Challenger 1': chall1_baseline.get('cumulative_roi', 0),
                                    'Challenger 2': chall2_baseline.get('cumulative_roi', 0),
                                    'Challenger 3': chall3_baseline.get('cumulative_roi', 0)
                                }
                                best_baseline = max(baseline_rois, key=baseline_rois.get)

                                # Calculate Weather Challenger advantage in stress vs baseline
                                weather_advantage_stress = max(chall2_roi, chall3_roi) - champ_roi
                                weather_advantage_baseline = max(
                                    chall2_baseline.get('cumulative_roi', 0),
                                    chall3_baseline.get('cumulative_roi', 0)
                                ) - champ_baseline.get('cumulative_roi', 0)

                                # Determine which weather challenger is better
                                better_weather = "Challenger 3 (MVO)" if chall3_roi > chall2_roi else "Challenger 2 (Naive)"

                                # Build insight message with full criteria for analog years
                                if "Analog" in stress['scenario']:
                                    # Use the full criteria label from weather config
                                    config_insight = st.session_state.get('ps_weather_config', {})
                                    criteria_insight = []
                                    if config_insight.get('enso_regime', 'Any') != 'Any':
                                        criteria_insight.append(config_insight.get('enso_regime'))
                                    if config_insight.get('historical_context', 'Any') != 'Any':
                                        criteria_insight.append(config_insight.get('historical_context'))
                                    if config_insight.get('trajectory', 'Any') != 'Any':
                                        criteria_insight.append(config_insight.get('trajectory'))
                                    scenario_display = " + ".join(criteria_insight) if criteria_insight else "selected analog"
                                else:
                                    scenario_display = stress['scenario'].replace(" Only", "").replace(" (except Current Year)", "").replace("Years", "years")

                                if weather_advantage_stress > weather_advantage_baseline + 0.02:  # Weather strategy shines
                                    insight_msg = f"""
**Your weather thesis is validated.** {better_weather} outperforms Champion by
**{weather_advantage_stress:+.1%}** in {scenario_display} conditions, compared to only
{weather_advantage_baseline:+.1%} across all years.

This confirms your strategy is well-tuned for the selected market view.
"""
                                    st.success(insight_msg)
                                elif weather_advantage_stress < weather_advantage_baseline - 0.02:  # Weather strategy underperforms
                                    insight_msg = f"""
**Caution: Your weather strategy underperforms in this scenario.** {better_weather} beats
Champion by only {weather_advantage_stress:+.1%} in {scenario_display} conditions, worse than
the {weather_advantage_baseline:+.1%} advantage across all years.

This may indicate the strategy is not optimally tuned for these specific conditions,
or that the analog years used for optimization don't represent this scenario well.
"""
                                    st.warning(insight_msg)
                                else:  # Roughly equivalent
                                    insight_msg = f"""
**Performance is consistent across scenarios.** {better_weather} outperforms Champion
by {weather_advantage_stress:+.1%} in {scenario_display} conditions, similar to the
{weather_advantage_baseline:+.1%} advantage across all years.

Your weather strategy appears robust and not overly dependent on specific conditions.
"""
                                    st.info(insight_msg)

                                # Add scenario-specific warnings
                                if "El Nino" in stress['scenario'] and weather_advantage_stress < 0:
                                    st.warning(f"""
⚠️ **Hedging Alert:** Your Weather Challengers underperform Champion by
{abs(weather_advantage_stress):.1%} if El Nino conditions materialize.
Consider whether your conviction in the La Nina thesis justifies this downside risk.
""")

                                # Download buttons for stress test
                                download_col1, download_col2 = st.columns(2)

                                with download_col1:
                                    stress_csv = comparison_df.to_csv(index=False)
                                    st.download_button(
                                        label="📥 Download Stress Test Results",
                                        data=stress_csv,
                                        file_name=f"stress_test_{stress['scenario'].replace(' ', '_').lower()}.csv",
                                        mime="text/csv",
                                        key="download_stress_csv"
                                    )

                                with download_col2:
                                    # Generate Word report
                                    weather_config = st.session_state.get('ps_weather_config', {})
                                    analog_count = stress.get('analog_years_count')

                                    docx_buffer = generate_strategy_report_docx(
                                        comparison_df=comparison_df,
                                        stress_scenario=stress['scenario'],
                                        start_year=start_year,
                                        end_year=end_year,
                                        coverage_level=coverage_level,
                                        productivity_factor=productivity_factor,
                                        intended_use=intended_use,
                                        analog_years_count=analog_count,
                                        weather_config=weather_config
                                    )

                                    st.download_button(
                                        label="📄 Export Strategy Report (Word)",
                                        data=docx_buffer,
                                        file_name=f"prf_strategy_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_strategy_docx"
                                    )

                    elif 'ps_analog_years' in st.session_state:
                        st.warning("No analog years found matching your criteria. Try broadening your market view (e.g., set some filters to 'Any').")

        # ==========================================================================
        # AGGREGATE PORTFOLIO DETAILS (Audit View)
        # ==========================================================================
        st.divider()
        st.markdown("### 📂 Aggregate Portfolio Details (Audit View)")
        st.caption("Complete year-by-year data for all strategies. Updates as you run each portfolio.")

        # Collect data from all available portfolios
        all_portfolio_data = []

        # Helper function to format allocation as string
        def format_allocation_string(alloc_dict):
            """Convert allocation dict to readable string like 'Jan-Feb: 20%, Mar-Apr: 30%'"""
            if not alloc_dict:
                return "N/A"
            parts = [f"{k}: {v*100:.0f}%" for k, v in sorted(alloc_dict.items(), key=lambda x: INTERVAL_ORDER_11.index(x[0])) if v > 0]
            return ", ".join(parts) if parts else "N/A"

        # --- Champion Data ---
        if 'champion_results' in st.session_state and st.session_state.champion_results:
            champ_data = st.session_state.champion_results
            champ_grids = champ_data.get('grids', [])
            champ_allocs = champ_data.get('allocations', {})
            champ_acres = champ_data.get('acres', {})
            champ_grid_results = champ_data.get('grid_results', {})

            for gid in champ_grids:
                if gid in champ_grid_results:
                    grid_result = champ_grid_results[gid]
                    if 'year_results' in grid_result and isinstance(grid_result['year_results'], pd.DataFrame):
                        year_df = grid_result['year_results'].copy()
                        year_df['Portfolio'] = 'Champion'
                        year_df['Grid'] = gid
                        year_df['Allocation'] = format_allocation_string(champ_allocs.get(gid, {}))
                        year_df['Acres'] = champ_acres.get(gid, 0)
                        all_portfolio_data.append(year_df)

        # --- Challenger 1 Data ---
        if 'challenger_results' in st.session_state and st.session_state.challenger_results:
            chall1_data = st.session_state.challenger_results
            chall1_grids = chall1_data.get('grids', [])
            chall1_allocs = chall1_data.get('allocations', {})
            chall1_acres = chall1_data.get('acres', {})
            chall1_grid_results = chall1_data.get('grid_results', {})

            for gid in chall1_grids:
                if gid in chall1_grid_results:
                    grid_result = chall1_grid_results[gid]
                    if 'year_results' in grid_result and isinstance(grid_result['year_results'], pd.DataFrame):
                        year_df = grid_result['year_results'].copy()
                        year_df['Portfolio'] = 'Challenger 1'
                        year_df['Grid'] = gid
                        year_df['Allocation'] = format_allocation_string(chall1_allocs.get(gid, {}))
                        year_df['Acres'] = chall1_acres.get(gid, 0)
                        all_portfolio_data.append(year_df)

        # --- Challenger 2 (Weather Naive) Data ---
        if 'weather_challenger_results' in st.session_state and st.session_state.weather_challenger_results:
            chall2_data = st.session_state.weather_challenger_results
            chall2_grids = chall2_data.get('grids', [])
            chall2_allocs = chall2_data.get('allocations', {})
            chall2_acres = chall2_data.get('acres', {})

            # Weather Challenger 2 stores results differently - need to recalculate or extract from df
            chall2_df = chall2_data.get('df', pd.DataFrame())
            if not chall2_df.empty:
                # The df has portfolio-level results, we need grid-level
                # Re-run calculation to get grid-level detail
                for gid in chall2_grids:
                    try:
                        allocation = chall2_allocs.get(gid, {})
                        acres = chall2_acres.get(gid, 0)

                        if not allocation:
                            continue

                        subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                        county_base_value = load_county_base_value(session, gid)
                        current_rate_year = get_current_rate_year(session)
                        premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
                        dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                        total_protection = dollar_protection * acres

                        all_indices_df = load_all_indices(session, gid)
                        all_indices_df = all_indices_df[(all_indices_df['YEAR'] >= start_year) & (all_indices_df['YEAR'] <= end_year)]

                        year_results = []
                        for year in sorted(all_indices_df['YEAR'].unique()):
                            year_data = all_indices_df[all_indices_df['YEAR'] == year]
                            if year_data.empty:
                                continue

                            total_indemnity = 0
                            total_producer_premium = 0

                            for interval, pct in allocation.items():
                                if pct == 0:
                                    continue

                                index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                                index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

                                premium_rate = premium_rates_df.get(interval, 0)
                                interval_protection = int(round_half_up(total_protection * pct, 0))
                                total_premium = int(round_half_up(interval_protection * premium_rate, 0))
                                premium_subsidy = int(round_half_up(total_premium * subsidy_percent, 0))
                                producer_premium = total_premium - premium_subsidy

                                trigger = coverage_level * 100
                                shortfall_pct = max(0, (trigger - index_value) / trigger)
                                raw_indemnity = shortfall_pct * interval_protection
                                # Convert to int immediately to ensure exact integer arithmetic when summing
                                indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                                total_indemnity += indemnity
                                total_producer_premium += producer_premium

                            # Both indemnity and premium are already exact integer sums

                            net_return = total_indemnity - total_producer_premium
                            roi = net_return / total_producer_premium if total_producer_premium > 0 else 0

                            year_results.append({
                                'year': year,
                                'indemnity': total_indemnity,
                                'premium': total_producer_premium,
                                'net': net_return,
                                'roi': roi
                            })

                        if year_results:
                            year_df = pd.DataFrame(year_results)
                            year_df['Portfolio'] = 'Challenger 2 (Weather)'
                            year_df['Grid'] = gid
                            year_df['Allocation'] = format_allocation_string(allocation)
                            year_df['Acres'] = acres
                            all_portfolio_data.append(year_df)

                    except Exception as e:
                        continue

        # --- Challenger 3 (Weather MVO) Data ---
        if 'weather_challenger_3_results' in st.session_state and st.session_state.weather_challenger_3_results:
            chall3_data = st.session_state.weather_challenger_3_results
            chall3_grids = chall3_data.get('grids', [])
            chall3_allocs = chall3_data.get('allocations', {})
            chall3_acres = chall3_data.get('acres', {})

            for gid in chall3_grids:
                try:
                    allocation = chall3_allocs.get(gid, {})
                    acres = chall3_acres.get(gid, 0)

                    if not allocation:
                        continue

                    subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                    county_base_value = load_county_base_value(session, gid)
                    current_rate_year = get_current_rate_year(session)
                    premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
                    dollar_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                    total_protection = dollar_protection * acres

                    all_indices_df = load_all_indices(session, gid)
                    all_indices_df = all_indices_df[(all_indices_df['YEAR'] >= start_year) & (all_indices_df['YEAR'] <= end_year)]

                    year_results = []
                    for year in sorted(all_indices_df['YEAR'].unique()):
                        year_data = all_indices_df[all_indices_df['YEAR'] == year]
                        if year_data.empty:
                            continue

                        total_indemnity = 0
                        total_producer_premium = 0

                        for interval, pct in allocation.items():
                            if pct == 0:
                                continue

                            index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                            index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100

                            premium_rate = premium_rates_df.get(interval, 0)
                            interval_protection = int(round_half_up(total_protection * pct, 0))
                            total_premium = int(round_half_up(interval_protection * premium_rate, 0))
                            premium_subsidy = int(round_half_up(total_premium * subsidy_percent, 0))
                            producer_premium = total_premium - premium_subsidy

                            trigger = coverage_level * 100
                            shortfall_pct = max(0, (trigger - index_value) / trigger)
                            raw_indemnity = shortfall_pct * interval_protection
                            # Convert to int immediately to ensure exact integer arithmetic when summing
                            indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                            total_indemnity += indemnity
                            total_producer_premium += producer_premium

                        # Both indemnity and premium are already exact integer sums

                        net_return = total_indemnity - total_producer_premium
                        roi = net_return / total_producer_premium if total_producer_premium > 0 else 0

                        year_results.append({
                            'year': year,
                            'indemnity': total_indemnity,
                            'premium': total_producer_premium,
                            'net': net_return,
                            'roi': roi
                        })

                    if year_results:
                        year_df = pd.DataFrame(year_results)
                        year_df['Portfolio'] = 'Challenger 3 (MVO)'
                        year_df['Grid'] = gid
                        year_df['Allocation'] = format_allocation_string(allocation)
                        year_df['Acres'] = acres
                        all_portfolio_data.append(year_df)

                except Exception as e:
                    continue

        # --- Combine and Display ---
        if all_portfolio_data:
            master_audit_df = pd.concat(all_portfolio_data, ignore_index=True)

            # Standardize column names
            column_rename = {
                'year': 'Year',
                'indemnity': 'Total Indemnity',
                'premium': 'Producer Premium',
                'net': 'Net Return',
                'roi': 'Total ROI'
            }
            master_audit_df = master_audit_df.rename(columns=column_rename)

            # Reorder columns: Portfolio, Grid, Year, Allocation, Acres, financials
            column_order = ['Portfolio', 'Grid', 'Year', 'Allocation', 'Acres',
                           'Total Indemnity', 'Producer Premium', 'Net Return', 'Total ROI']

            # Only include columns that exist
            column_order = [c for c in column_order if c in master_audit_df.columns]
            master_audit_df = master_audit_df[column_order]

            # Sort by Portfolio, Grid, Year
            master_audit_df = master_audit_df.sort_values(['Portfolio', 'Grid', 'Year'])

            # Count portfolios present
            portfolios_present = master_audit_df['Portfolio'].nunique()
            total_rows = len(master_audit_df)

            with st.expander(f"📊 Aggregate Portfolio Details ({portfolios_present} portfolios, {total_rows:,} records)", expanded=False):
                st.caption("Year-by-year breakdown for all grids across all portfolios. Use this to verify calculations.")

                # Filter options
                filter_col1, filter_col2 = st.columns(2)
                with filter_col1:
                    portfolio_filter = st.multiselect(
                        "Filter by Portfolio:",
                        options=sorted(master_audit_df['Portfolio'].unique()),
                        default=sorted(master_audit_df['Portfolio'].unique()),
                        key="audit_portfolio_filter"
                    )
                with filter_col2:
                    grid_filter = st.multiselect(
                        "Filter by Grid:",
                        options=sorted(master_audit_df['Grid'].unique()),
                        default=sorted(master_audit_df['Grid'].unique()),
                        key="audit_grid_filter"
                    )

                # Apply filters
                filtered_df = master_audit_df[
                    (master_audit_df['Portfolio'].isin(portfolio_filter)) &
                    (master_audit_df['Grid'].isin(grid_filter))
                ]

                # Display
                st.dataframe(
                    filtered_df.style.format({
                        'Year': '{:.0f}',
                        'Acres': '{:,.0f}',
                        'Total Indemnity': '${:,.0f}',
                        'Producer Premium': '${:,.0f}',
                        'Net Return': '${:,.0f}',
                        'Total ROI': '{:.2%}'
                    }),
                    use_container_width=True,
                    height=400,
                    hide_index=True
                )

                # Download button
                csv_data = filtered_df.to_csv(index=False)
                st.download_button(
                    label="📥 Download Aggregate Audit CSV",
                    data=csv_data,
                    file_name="champion_vs_challenger_audit_details.csv",
                    mime="text/csv",
                    key="download_aggregate_audit"
                )

                # Summary stats per portfolio
                st.markdown("**Portfolio Summary:**")
                summary_data = []
                for portfolio in sorted(filtered_df['Portfolio'].unique()):
                    port_df = filtered_df[filtered_df['Portfolio'] == portfolio]
                    total_indem = port_df['Total Indemnity'].sum()
                    total_prem = port_df['Producer Premium'].sum()
                    total_net = port_df['Net Return'].sum()
                    cum_roi = total_net / total_prem if total_prem > 0 else 0
                    num_grids = port_df['Grid'].nunique()
                    num_years = port_df['Year'].nunique()

                    summary_data.append({
                        'Portfolio': portfolio,
                        'Grids': num_grids,
                        'Years': num_years,
                        'Total Indemnity': total_indem,
                        'Total Premium': total_prem,
                        'Net Return': total_net,
                        'Cumulative ROI': cum_roi
                    })

                summary_df = pd.DataFrame(summary_data)
                st.dataframe(
                    summary_df.style.format({
                        'Total Indemnity': '${:,.0f}',
                        'Total Premium': '${:,.0f}',
                        'Net Return': '${:,.0f}',
                        'Cumulative ROI': '{:.2%}'
                    }),
                    use_container_width=True,
                    hide_index=True
                )
        else:
            st.info("Run Champion and/or Challenger strategies to see aggregate details here.")


# =============================================================================
# === 5. TAB 3: BACKTESTING ENGINE (S3) ===
# =============================================================================
def render_tab3(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code):
    st.subheader("Parameters")
    
    col1, col2, col3 = st.columns(3)
    start_year = col1.selectbox("Start Year", list(range(1948, 2026)), index=62, key="s3_start")
    end_year = col2.selectbox("End Year", list(range(1948, 2026)), index=77, key="s3_end")
    coverage_level = col3.selectbox("Coverage Level", [0.70, 0.75, 0.80, 0.85, 0.90], index=2, format_func=lambda x: f"{x:.0%}", key="s3_coverage")

    with st.expander("Step 2: Set Interval Allocations", expanded=True):
        pct_of_value_alloc, is_valid = render_allocation_inputs("s3")
    
    st.divider()

    if 'tab3_run' not in st.session_state:
        st.session_state.tab3_run = False

    if st.button("Run Backtest", key="s3_run_button", disabled=not is_valid):
        st.session_state.tab3_run = True
        try:
            with st.spinner(f"Analyzing {start_year}-{end_year}..."):
                coverage_level_string = f"{coverage_level:.0%}"
                subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                county_base_value = load_county_base_value(session, grid_id)
                current_rate_year = get_current_rate_year(session)
                premium_rates_df = load_premium_rates(session, grid_id, intended_use, [coverage_level], current_rate_year)[coverage_level]
                dollar_amount_of_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                total_policy_protection = dollar_amount_of_protection * total_insured_acres
                all_indices_df = load_all_indices(session, grid_id)
                
                year_results = []
                for year in range(start_year, end_year + 1):
                    actuals_df = all_indices_df[all_indices_df['YEAR'] == year].set_index('INTERVAL_NAME')
                    if actuals_df.empty: continue
                    
                    roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                    roi_df['Percent of Value'] = roi_df.index.map(pct_of_value_alloc)
                    roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                    roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                    roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                    roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                    roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce').fillna(0)
                    roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                    roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                    roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                    trigger_level = coverage_level * 100
                    shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                    roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)
                    total_indemnity = roi_df['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                    total_producer_prem = roi_df['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                    year_roi = (total_indemnity - total_producer_prem) / total_producer_prem if total_producer_prem > 0 else 0.0
                    
                    year_results.append({
                        'Year': year, 'Total Indemnity': total_indemnity, 'Producer Premium': total_producer_prem,
                        'Net Return': total_indemnity - total_producer_prem, 'Total ROI': year_roi
                    })
            
            results_df = pd.DataFrame(year_results)

            st.session_state.tab3_results = {
                "results_df": results_df, "grid_id": grid_id, "start_year": start_year, "end_year": end_year,
                "current_rate_year": current_rate_year, "coverage_level": coverage_level, "productivity_factor": productivity_factor,
                "intended_use": intended_use, "total_insured_acres": total_insured_acres,
                "pct_of_value_alloc": pct_of_value_alloc
            }
            st.session_state.tab1_results = None
            st.session_state.tab2_results = None
            st.session_state.tab4_results = None
            st.session_state.tab5_results = None

        except Exception as e:
            st.error(f"Error: {e}")
            st.exception(e)
            st.session_state.tab3_results = None

    if 'tab3_results' in st.session_state and st.session_state.tab3_results:
        try:
            r = st.session_state.tab3_results
            st.header(f"Backtest Results - Grid {r['grid_id']}")
            st.caption(f"Coverage: {r['coverage_level']:.0%} | Productivity: {r['productivity_factor']:.0%} | Use: {r['intended_use']} | Acres: {r['total_insured_acres']:,}")
            
            st.subheader("Strategy")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Coverage", f"{r['coverage_level']:.0%}")
            c2.metric("Productivity", f"{r['productivity_factor']:.0%}")
            c3.metric("Use", r['intended_use'])
            c4.metric("Acres", f"{r['total_insured_acres']:,}")
            with st.expander("Show Allocation"):
                alloc_display = {k: f"{v*100:.0f}%" for k, v in r['pct_of_value_alloc'].items() if v > 0}
                st.dataframe(pd.Series(alloc_display, name="Allocation"), use_container_width=True)

            st.subheader("Summary")
            results_df = r['results_df']
            profitable_years = results_df[results_df['Total ROI'] > 0]
            total_years = len(results_df)
            profitable_count = len(profitable_years)
            
            s1, s2, s3 = st.columns(3)
            s1.metric("Years Analyzed", f"{total_years}")
            s2.metric("Win Rate", f"{profitable_count} ({profitable_count/total_years*100:.1f}%)")
            s3.metric("Median ROI", f"{results_df['Total ROI'].median():.2%}")
            
            s4, s5 = st.columns(2)
            s4.metric("Best Year", f"{results_df['Total ROI'].max():.2%} ({int(results_df.loc[results_df['Total ROI'].idxmax(), 'Year'])})")
            s5.metric("Worst Year", f"{results_df['Total ROI'].min():.2%} ({int(results_df.loc[results_df['Total ROI'].idxmin(), 'Year'])})")

            st.subheader("Cumulative")
            total_indemnity_all = results_df['Total Indemnity'].sum()
            total_premium_all = results_df['Producer Premium'].sum()
            net_return_all = results_df['Net Return'].sum()
            cumulative_roi = net_return_all / total_premium_all if total_premium_all > 0 else 0

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Total Indemnity", f"${total_indemnity_all:,.0f}")
            c2.metric("Total Premium", f"${total_premium_all:,.0f}")
            c3.metric("Net Return", f"${net_return_all:,.0f}")
            c4.metric("Cumulative ROI", f"{cumulative_roi:.2%}")

            st.subheader("Annual Results")
            
            csv = results_df.to_csv(index=False)
            st.download_button(
                label="Download CSV",
                data=csv,
                file_name=f"backtest_grid_{extract_numeric_grid_id(r['grid_id'])}_{r['start_year']}-{r['end_year']}.csv",
                mime="text/csv",
            )
            
            st.dataframe(results_df.style.format({
                'Year': '{:.0f}', 'Total Indemnity': '${:,.0f}',
                'Producer Premium': '${:,.0f}', 'Net Return': '${:,.0f}', 'Total ROI': '{:.2%}'
            }), use_container_width=True)
        except Exception as e:
            st.error(f"Error displaying results: {e}")
            st.session_state.tab3_results = None
    elif not st.session_state.tab3_run:
        st.info("Select parameters and click 'Run Backtest'")

# =============================================================================
# === 6. TAB 5: PORTFOLIO BACKTESTING ENGINE ===
# =============================================================================
def render_tab5(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code):
    st.subheader("Portfolio Backtesting Engine")

    # === PRESET LOADING BUTTON ===
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("Load King Ranch", key="tab5_load_kr"):
            try:
                all_grids_for_preset = load_distinct_grids(session)

                # Build mapping of numeric IDs to their proper county names from preset
                target_grid_mapping = {}
                for county, grid_ids in KING_RANCH_PRESET['counties'].items():
                    for grid_id_num in grid_ids:
                        target_grid_mapping[grid_id_num] = f"{grid_id_num} ({county} - TX)"

                # Match grids in the order from preset
                preset_grid_ids = []
                for numeric_id in KING_RANCH_PRESET['grids']:
                    target_str = target_grid_mapping[numeric_id]
                    # Try exact match first
                    if target_str in all_grids_for_preset:
                        preset_grid_ids.append(target_str)
                    else:
                        # Fallback: find any grid with this numeric ID
                        found = False
                        for grid_option in all_grids_for_preset:
                            if extract_numeric_grid_id(grid_option) == numeric_id:
                                preset_grid_ids.append(grid_option)
                                found = True
                                break
                        if not found:
                            st.warning(f"Could not find grid {numeric_id}")

                st.session_state.tab5_grids = preset_grid_ids
                st.session_state.tab5_use_custom_acres = True  # Enable custom acres for King Ranch

                # Set acres for each grid
                for gid in preset_grid_ids:
                    numeric_id = extract_numeric_grid_id(gid)
                    st.session_state[f"tab5_acres_{gid}"] = KING_RANCH_PRESET['acres'][numeric_id]

                # Set allocations via preset keys (NOT editor keys)
                for gid in preset_grid_ids:
                    numeric_id = extract_numeric_grid_id(gid)
                    alloc = KING_RANCH_PRESET['allocations'][numeric_id]
                    # Convert percentages to decimals and store in preset key
                    alloc_decimal = {interval: float(alloc.get(interval, 0.0)) / 100.0 for interval in INTERVAL_ORDER_11}
                    st.session_state[f"tab5_grid_{gid}_preset_allocation"] = alloc_decimal

                # Set King Ranch specific parameters
                st.session_state.productivity_factor = 1.35  # 135%
                st.session_state.tab5_coverage = 0.75  # 75% coverage

                st.success("King Ranch loaded! (8 grids, 135% productivity, 75% coverage)")
                # Removed st.rerun() - rely on Streamlit's natural reactivity

            except Exception as e:
                st.error(f"Error loading King Ranch: {e}")
                st.exception(e)

    with col2:
        st.caption("Auto-populate King Ranch strategy")

    st.divider()

    try:
        all_grids = load_distinct_grids(session)
    except:
        all_grids = [grid_id]

    default_grids = st.session_state.get('tab5_grids', [grid_id])

    selected_grids = st.multiselect(
        "Select Grids",
        options=all_grids,
        default=default_grids,
        max_selections=20,
        key="tab5_grids"
    )

    if not selected_grids:
        st.warning("Select at least one grid")
        return

    st.divider()

    # === SCENARIO DEFINITION SECTION (NEW) ===
    st.markdown("#### Scenario Definition")

    scenario_options = [
        'All Years (except Current Year)',
        'ENSO Phase: La Nina',
        'ENSO Phase: El Nino',
        'ENSO Phase: Neutral',
        'Select my own interval'
    ]

    selected_scenario = st.radio(
        "Select one scenario to backtest:",
        options=scenario_options,
        index=0,
        key='tab5_scenario_select'
    )

    # Conditional year range display - only show if "Select my own interval" is chosen
    start_year = 1948
    end_year = 2024
    if selected_scenario == 'Select my own interval':
        col1, col2 = st.columns(2)
        start_year = col1.selectbox("Start Year", list(range(1948, 2026)), index=62, key="tab5_start")
        end_year = col2.selectbox("End Year", list(range(1948, 2026)), index=76, key="tab5_end")

    st.divider()

    # === PARAMETERS SECTION ===
    st.subheader("Parameters")
    coverage_level = st.selectbox(
        "Coverage Level",
        [0.70, 0.75, 0.80, 0.85, 0.90],
        index=2,
        format_func=lambda x: f"{x:.0%}",
        key="tab5_coverage"
    )

    st.divider()

    # === ACRE CONFIGURATION SECTION ===
    st.subheader("Acre Configuration")

    use_custom_acres = st.checkbox(
        "Configure acres per grid",
        value=st.session_state.get('tab5_use_custom_acres', False),
        key="tab5_use_custom_acres"
    )

    grid_acres = {}
    if use_custom_acres:
        cols = st.columns(min(4, len(selected_grids)))
        for idx, gid in enumerate(selected_grids):
            with cols[idx % 4]:
                # Check if King Ranch preset exists for this grid
                numeric_id = extract_numeric_grid_id(gid)
                default_acres = st.session_state.get(f"tab5_acres_{gid}", KING_RANCH_PRESET['acres'].get(numeric_id, total_insured_acres))
                grid_acres[gid] = st.number_input(
                    f"{gid}",
                    min_value=1,
                    value=default_acres,
                    step=10,
                    key=f"tab5_acres_{gid}"
                )
    else:
        # Show info message about equal distribution
        acres_per_grid = total_insured_acres // len(selected_grids)
        st.info(f"Using {total_insured_acres:,} acres from Common Parameters, equally distributed across {len(selected_grids)} grids ({acres_per_grid:,} acres per grid)")
        for gid in selected_grids:
            grid_acres[gid] = acres_per_grid

    st.divider()

    # === ALLOCATIONS SECTION ===
    st.subheader(f"Allocations for {len(selected_grids)} Grid(s)")

    grid_allocations = {}
    all_valid = True

    for gid in selected_grids:
        with st.expander(f"{gid} ({grid_acres[gid]:,} acres)", expanded=len(selected_grids) == 1):
            alloc_dict, is_valid = render_allocation_inputs(f"tab5_grid_{gid}")
            grid_allocations[gid] = alloc_dict
            if not is_valid:
                all_valid = False

    st.divider()

    # === RUN BACKTEST SECTION ===
    if 'tab5_run' not in st.session_state:
        st.session_state.tab5_run = False

    if st.button("Run Portfolio Backtest", key="tab5_run_button", disabled=not all_valid):
        st.session_state.tab5_run = True
        try:
            grid_results = {}
            years_used = []

            with st.spinner(f"Running backtest for {len(selected_grids)} grids..."):
                for gid in selected_grids:
                    try:
                        subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                        county_base_value = load_county_base_value(session, gid)
                        current_rate_year = get_current_rate_year(session)
                        premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
                        dollar_amount_of_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                        total_policy_protection = dollar_amount_of_protection * grid_acres[gid]
                        all_indices_df = load_all_indices(session, gid)

                        # Apply scenario-based year filtering using cached helper
                        filtered_df = filter_indices_by_scenario(
                            all_indices_df, selected_scenario, start_year, end_year
                        )

                        # Warn if ENSO filtering was requested but column not available
                        if 'ENSO Phase' in selected_scenario and 'OPTICAL_MAPPING_CPC' not in all_indices_df.columns:
                            st.warning(f"ENSO data not available for {gid}, using all years")

                        # Get unique years for this grid
                        grid_years = filtered_df['YEAR'].unique()

                        year_results = []
                        for year in sorted(grid_years):
                            actuals_df = filtered_df[filtered_df['YEAR'] == year].set_index('INTERVAL_NAME')
                            if actuals_df.empty:
                                continue

                            roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                            roi_df['Percent of Value'] = roi_df.index.map(grid_allocations[gid])
                            roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                            roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                            roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                            roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                            roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce').fillna(0)
                            roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                            roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                            roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                            trigger_level = coverage_level * 100
                            shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                            roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)
                            total_indemnity = roi_df['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                            total_producer_prem = roi_df['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                            year_roi = (total_indemnity - total_producer_prem) / total_producer_prem if total_producer_prem > 0 else 0.0

                            year_results.append({
                                'Year': year, 'Total Indemnity': total_indemnity, 'Producer Premium': total_producer_prem,
                                'Net Return': total_indemnity - total_producer_prem, 'Total ROI': year_roi
                            })
                            if year not in years_used:
                                years_used.append(year)

                        results_df = pd.DataFrame(year_results)
                        grid_results[gid] = {
                            'results_df': results_df,
                            'allocation': grid_allocations[gid]
                        }

                    except Exception as e:
                        st.error(f"Grid {gid}: {str(e)}")

            # Determine display year range
            if years_used:
                display_start = min(years_used)
                display_end = max(years_used)
            else:
                display_start = start_year
                display_end = end_year

            st.session_state.tab5_results = {
                "grid_results": grid_results,
                "selected_grids": selected_grids,
                "grid_acres": grid_acres,
                "grid_allocations": grid_allocations,
                "start_year": display_start,
                "end_year": display_end,
                "coverage_level": coverage_level,
                "productivity_factor": productivity_factor,
                "intended_use": intended_use,
                "total_insured_acres": total_insured_acres,
                "current_rate_year": current_rate_year,
                "scenario": selected_scenario,
                "years_used": sorted(years_used)
            }
            st.session_state.tab2_results = None
            st.session_state.tab4_results = None

        except Exception as e:
            st.error(f"Error: {e}")
            st.exception(e)
            st.session_state.tab5_results = None

    if 'tab5_results' in st.session_state and st.session_state.tab5_results:
        try:
            r = st.session_state.tab5_results

            st.header(f"Portfolio Results ({r['start_year']}-{r['end_year']})")
            st.caption(f"Scenario: {r.get('scenario', 'All Years')} | Coverage: {r['coverage_level']:.0%} | Productivity: {r['productivity_factor']:.0%} | Use: {r['intended_use']}")

            if r.get('years_used'):
                st.caption(f"Years included: {len(r['years_used'])} ({min(r['years_used'])}-{max(r['years_used'])})")

            # === PORTFOLIO COVERAGE TABLE (NEW) ===
            if len(r['selected_grids']) > 1:
                st.subheader("Portfolio Coverage")

                coverage_data = []
                total_coverage = {interval: 0 for interval in INTERVAL_ORDER_11}
                total_portfolio_acres = sum(r['grid_acres'].values())

                for gid in r['selected_grids']:
                    if gid in r['grid_results']:
                        allocation = r['grid_results'][gid]['allocation']
                        row = {'Grid': str(gid)[:20]}
                        row_sum = 0

                        for interval in INTERVAL_ORDER_11:
                            pct = allocation.get(interval, 0) * 100
                            row_sum += pct
                            total_coverage[interval] += pct
                            row[interval] = f"{pct:.0f}%" if pct > 0 else "--"

                        row['Row Sum'] = f"{row_sum:.0f}%"
                        row['Acres'] = f"{r['grid_acres'].get(gid, 0):,.0f}"
                        coverage_data.append(row)

                # Add average row
                avg_row = {'Grid': 'AVERAGE'}
                avg_row_sum = 0
                valid_grids_count = len([gid for gid in r['selected_grids'] if gid in r['grid_results']])
                for interval in INTERVAL_ORDER_11:
                    pct = total_coverage[interval] / valid_grids_count if valid_grids_count > 0 else 0
                    avg_row_sum += pct
                    avg_row[interval] = f"{pct:.0f}%" if pct > 0.5 else "--"
                avg_row['Row Sum'] = f"{avg_row_sum:.0f}%"
                avg_row['Acres'] = f"{total_portfolio_acres:,.0f}"
                coverage_data.append(avg_row)

                coverage_df = pd.DataFrame(coverage_data)

                # CSV Download Button for coverage
                csv_coverage = coverage_df.to_csv(index=False)
                st.download_button(
                    label="Download Coverage CSV",
                    data=csv_coverage,
                    file_name=f"portfolio_coverage_{r['start_year']}-{r['end_year']}.csv",
                    mime="text/csv",
                    key="tab5_coverage_csv"
                )

                st.dataframe(coverage_df, use_container_width=True, hide_index=True)

                st.divider()

            st.subheader("Cumulative Results by Grid")

            combined_data = []
            portfolio_total_premium = 0
            portfolio_total_indemnity = 0
            portfolio_total_net_return = 0
            year_rois_all_grids = []

            for gid in r['selected_grids']:
                if gid in r['grid_results']:
                    results_df = r['grid_results'][gid]['results_df']

                    total_indemnity = results_df['Total Indemnity'].sum()
                    total_premium = results_df['Producer Premium'].sum()
                    net_return = results_df['Net Return'].sum()
                    cumulative_roi = net_return / total_premium if total_premium > 0 else 0

                    year_rois = results_df['Total ROI'].values
                    std_dev = np.std(year_rois) if len(year_rois) > 0 else 0
                    risk_adj_ret = cumulative_roi / std_dev if std_dev > 0 else 0

                    portfolio_total_premium += total_premium
                    portfolio_total_indemnity += total_indemnity
                    portfolio_total_net_return += net_return
                    year_rois_all_grids.extend(year_rois)

                    grid_acres_val = r['grid_acres'].get(gid, 0)

                    combined_data.append({
                        'Grid': str(gid)[:20],
                        'Acres': grid_acres_val,
                        'Total Premium': total_premium,
                        'Total Indemnity': total_indemnity,
                        'Net Return': net_return,
                        'Cumulative ROI': cumulative_roi,
                        'Std Dev': std_dev,
                        'Risk-Adj Return': risk_adj_ret
                    })

            # CSV Download Button
            csv_df = pd.DataFrame(combined_data)
            csv_export = csv_df.to_csv(index=False)

            st.download_button(
                label="Download Results CSV",
                data=csv_export,
                file_name=f"portfolio_results_{r['start_year']}-{r['end_year']}.csv",
                mime="text/csv",
                key="tab5_results_csv"
            )

            # === FORMATTED TABLE ===
            st.text("Grid                  Acres   Total Premium    Total Indemnity       Net Return        Cum ROI        Std Dev       Risk-Adj")
            st.text("-" * 145)

            for row in combined_data:
                line = f"{row['Grid']:<20} {row['Acres']:>7,} {row['Total Premium']:>16,.0f} {row['Total Indemnity']:>17,.0f} {row['Net Return']:>16,.0f} {row['Cumulative ROI']:>14.2%} {row['Std Dev']:>13.2%} {row['Risk-Adj Return']:>13.2f}"
                st.text(line)

            st.text("-" * 145)
            total_acres = sum(r['grid_acres'].values())
            totals_line = f"{'TOTAL':<20} {total_acres:>7,} {portfolio_total_premium:>16,.0f} {portfolio_total_indemnity:>17,.0f} {portfolio_total_net_return:>16,.0f} {'':>14} {'':>13} {'':>13}"
            st.text(totals_line)
            st.text("=" * 145)

            st.divider()
            st.subheader("Portfolio Metrics")

            portfolio_roi = portfolio_total_net_return / portfolio_total_premium if portfolio_total_premium > 0 else 0

            if len(year_rois_all_grids) > 0:
                portfolio_std_dev = np.std(year_rois_all_grids)
                portfolio_risk_adj = portfolio_roi / portfolio_std_dev if portfolio_std_dev > 0 else 0
            else:
                portfolio_std_dev = 0
                portfolio_risk_adj = 0

            c1, c2 = st.columns(2)
            c1.metric("Portfolio ROI", f"{portfolio_roi:.2%}")
            c2.metric("Risk-Adjusted Return", f"{portfolio_risk_adj:.2f}")

            st.divider()
            st.subheader("Details by Grid")

            # Initialize collection list for aggregate audit view
            all_grids_detail_data = []

            for gid in r['selected_grids']:
                if gid in r['grid_results']:
                    with st.expander(f"{gid} ({r['grid_acres'].get(gid, 0):,} acres)"):
                        results_df = r['grid_results'][gid]['results_df']
                        allocation = r['grid_results'][gid]['allocation']

                        alloc_display = {k: f"{v*100:.0f}%" for k, v in allocation.items() if v > 0}
                        st.text(f"Allocation: {', '.join([f'{k}: {v}' for k, v in alloc_display.items()])}")

                        st.dataframe(results_df.style.format({
                            'Year': '{:.0f}', 'Total Indemnity': '${:,.0f}',
                            'Producer Premium': '${:,.0f}', 'Net Return': '${:,.0f}', 'Total ROI': '{:.2%}'
                        }), use_container_width=True)

                    # Accumulate data for aggregate view
                    grid_df = results_df.copy()
                    grid_df.insert(0, 'Grid ID', gid)
                    all_grids_detail_data.append(grid_df)

            # Render Aggregate Portfolio Details Panel
            if all_grids_detail_data:
                master_audit_df = pd.concat(all_grids_detail_data, ignore_index=True)

                # Reorder columns so Grid ID and Year are at the front
                cols = master_audit_df.columns.tolist()
                priority_cols = ['Grid ID', 'Year']
                other_cols = [c for c in cols if c not in priority_cols]
                master_audit_df = master_audit_df[priority_cols + other_cols]

                with st.expander("📂 Aggregate Portfolio Details (Audit View)"):
                    st.caption(f"Complete backtesting data for all {len(r['selected_grids'])} grids across all years.")

                    st.dataframe(master_audit_df.style.format({
                        'Year': '{:.0f}', 'Total Indemnity': '${:,.0f}',
                        'Producer Premium': '${:,.0f}', 'Net Return': '${:,.0f}', 'Total ROI': '{:.2%}'
                    }), use_container_width=True, height=400)

                    # Download button for CSV export
                    csv_data = master_audit_df.to_csv(index=False)
                    st.download_button(
                        label="Download Master Audit CSV",
                        data=csv_data,
                        file_name="portfolio_audit_details.csv",
                        mime="text/csv",
                        key="download_master_audit"
                    )

        except Exception as e:
            st.error(f"Error: {e}")
            st.exception(e)
            st.session_state.tab5_results = None
    elif not st.session_state.tab5_run:
        st.info("Configure grids and click 'Run Portfolio Backtest'")

# =============================================================================
# === 7. TAB 4: OPTIMIZATION ENGINE ===
# =============================================================================

@st.cache_data(ttl=3600)
def run_optimization_s4(
    _session, grid_id, start_year, end_year, plan_code, prod_factor, 
    acres, use, coverage_levels, objective, min_intervals, max_intervals_to_test, search_depth
):
    # Load data
    county_base_value = load_county_base_value(_session, grid_id)
    all_indices_df = load_all_indices(_session, grid_id)
    all_indices_df = all_indices_df[(all_indices_df['YEAR'] >= start_year) & (all_indices_df['YEAR'] <= end_year)]
    
    current_rate_year = get_current_rate_year(_session)
    all_premiums = load_premium_rates(_session, grid_id, use, coverage_levels, current_rate_year)
    all_subsidies = load_subsidies(_session, plan_code, coverage_levels)

    # === MARGINAL MODE ===
    if search_depth == 'marginal':
        numeric_grid_id = extract_numeric_grid_id(grid_id)
        
        if numeric_grid_id in KING_RANCH_PRESET['allocations']:
            base_allocation = KING_RANCH_PRESET['allocations'][numeric_grid_id]
            candidates = generate_marginal_variations(base_allocation)
        else:
            return pd.DataFrame(columns=['coverage_level', 'allocation', 'average_roi', 'median_roi', 
                                         'cumulative_roi', 'profitable_pct', 'std_dev', 'min_roi', 
                                         'max_roi', 'risk_adj_ret']), current_rate_year, 0
    
    # === INCREMENTAL MODE ===
    elif search_depth == 'incremental':
        numeric_grid_id = extract_numeric_grid_id(grid_id)
        
        if numeric_grid_id in KING_RANCH_PRESET['allocations']:
            base_allocation = KING_RANCH_PRESET['allocations'][numeric_grid_id]
            candidates = generate_incremental_variations(base_allocation)
        else:
            return pd.DataFrame(columns=['coverage_level', 'allocation', 'average_roi', 'median_roi', 
                                         'cumulative_roi', 'profitable_pct', 'std_dev', 'min_roi', 
                                         'max_roi', 'risk_adj_ret']), current_rate_year, 0
    
    else:
        # === STANDARD MODE ===
        interval_scores = {}
        for interval in INTERVAL_ORDER_11:
            interval_data = all_indices_df[all_indices_df['INTERVAL_NAME'] == interval]['INDEX_VALUE']
            avg_shortage = (100 - interval_data).mean()
            interval_scores[interval] = avg_shortage
        sorted_intervals = sorted(interval_scores.items(), key=lambda x: x[1], reverse=True)
        
        search_depth_map = {'fast': 5, 'standard': 6, 'thorough': 7, 'maximum': 8}
        search_depth_num = search_depth_map.get(search_depth.lower(), 6)
        top_intervals = [x[0] for x in sorted_intervals[:search_depth_num]]

        candidates = []
        
        for num_intervals in range(min_intervals, max_intervals_to_test + 1):
            for combo in combinations(top_intervals, num_intervals):
                if has_adjacent_intervals_in_list(list(combo)):
                    continue
                    
                candidates.extend(generate_allocations(combo, num_intervals))
    
    unique_candidates = []
    seen = set()
    for candidate in candidates:
        key = tuple(sorted(candidate.items()))
        if key not in seen:
            seen.add(key)
            unique_candidates.append(candidate)
    
    results = []
    
    def calculate_roi_for_strategy(allocation, coverage_level):
        subsidy = all_subsidies[coverage_level]
        premiums = all_premiums[coverage_level]
        dollar_protection = calculate_protection(county_base_value, coverage_level, prod_factor)
        total_protection = dollar_protection * acres
        
        year_rois = []
        total_indemnity_all_years = 0
        total_producer_premium_all_years = 0
        
        for year in range(start_year, end_year + 1):
            year_data = all_indices_df[all_indices_df['YEAR'] == year]
            if year_data.empty: continue
            
            total_indemnity, total_producer_premium = 0, 0
            
            if abs(sum(allocation.values()) - 1.0) > 0.001: return None
            
            for interval, pct in allocation.items():
                if pct == 0: continue
                if pct > 0.501: return None
                
                index_row = year_data[year_data['INTERVAL_NAME'] == interval]
                index_value = float(index_row['INDEX_VALUE'].iloc[0]) if not index_row.empty else 100
                premium_rate = premiums.get(interval, 0)
                interval_protection = int(round_half_up(total_protection * pct, 0))
                total_premium = int(round_half_up(interval_protection * premium_rate, 0))
                premium_subsidy = int(round_half_up(total_premium * subsidy, 0))
                producer_premium = total_premium - premium_subsidy

                trigger = coverage_level * 100
                shortfall_pct = max(0, (trigger - index_value) / trigger)
                raw_indemnity = shortfall_pct * interval_protection
                # Convert to int immediately to ensure exact integer arithmetic when summing
                indemnity = int(round_half_up(raw_indemnity, 0)) if raw_indemnity >= 0.01 else 0

                total_indemnity += indemnity
                total_producer_premium += producer_premium

            # Both indemnity and premium are already exact integer sums

            year_roi = (total_indemnity - total_producer_premium) / total_producer_premium if total_producer_premium > 0 else 0
            year_rois.append(year_roi)

            total_indemnity_all_years += total_indemnity
            total_producer_premium_all_years += total_producer_premium

        if len(year_rois) == 0: return None
        year_rois_array = np.array(year_rois)

        average_roi = year_rois_array.mean()
        cumulative_roi = (total_indemnity_all_years - total_producer_premium_all_years) / total_producer_premium_all_years if total_producer_premium_all_years > 0 else 0
        std_dev = year_rois_array.std()
        
        if std_dev > 0 and not np.isnan(std_dev) and not np.isinf(cumulative_roi):
            risk_adj_ret = cumulative_roi / std_dev
            risk_adj_ret = np.clip(risk_adj_ret, -1000, 1000)
        else:
            risk_adj_ret = 0
        
        return {
            'average_roi': average_roi, 
            'median_roi': np.median(year_rois_array),
            'cumulative_roi': cumulative_roi,
            'profitable_pct': (year_rois_array > 0).sum() / len(year_rois_array),
            'std_dev': std_dev, 
            'min_roi': year_rois_array.min(), 
            'max_roi': year_rois_array.max(),
            'risk_adj_ret': risk_adj_ret
        }

    for coverage_level in coverage_levels:
        for allocation in unique_candidates:
            metrics = calculate_roi_for_strategy(allocation, coverage_level)
            if metrics is not None:
                results.append({'coverage_level': coverage_level, 'allocation': allocation, **metrics})
    
    if len(results) == 0:
        return pd.DataFrame(columns=['coverage_level', 'allocation', 'average_roi', 'median_roi', 
                                     'cumulative_roi', 'profitable_pct', 'std_dev', 'min_roi', 
                                     'max_roi', 'risk_adj_ret']), current_rate_year, len(unique_candidates)
    
    results_df = pd.DataFrame(results)
    
    if objective not in results_df.columns:
        objective = 'cumulative_roi'
    
    results_df = results_df.sort_values(objective, ascending=False)
    return results_df, current_rate_year, len(unique_candidates)

def render_tab4(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code):
    st.subheader("Optimizer")
    
    # === PRESET LOADING BUTTON ===
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("📋 Load King Ranch Grids", key="tab4_load_kr_grids"):
            try:
                all_grids_for_preset = load_distinct_grids(session)
                
                # Build mapping of numeric IDs to their proper county names from preset
                target_grid_mapping = {}
                for county, grid_ids in KING_RANCH_PRESET['counties'].items():
                    for grid_id_num in grid_ids:
                        target_grid_mapping[grid_id_num] = f"{grid_id_num} ({county} - TX)"
                
                # Match grids in the order from preset
                preset_grid_ids = []
                for numeric_id in KING_RANCH_PRESET['grids']:
                    target_str = target_grid_mapping[numeric_id]
                    # Try exact match first
                    if target_str in all_grids_for_preset:
                        preset_grid_ids.append(target_str)
                    else:
                        # Fallback: find any grid with this numeric ID
                        found = False
                        for grid_option in all_grids_for_preset:
                            if extract_numeric_grid_id(grid_option) == numeric_id:
                                preset_grid_ids.append(grid_option)
                                found = True
                                break
                        if not found:
                            st.warning(f"Could not find grid {numeric_id}")
                
                st.session_state.s4_grids = preset_grid_ids
                
                # Set King Ranch specific parameters
                st.session_state.productivity_factor = 1.35  # 135%
                st.session_state.s4_coverage = [0.75]  # 75% coverage only
                st.session_state.s4_king_ranch_comparison_mode = True  # Flag for comparison output
                
                if preset_grid_ids:
                    st.success("✅ King Ranch loaded! (8 grids, 135% productivity, 75% coverage)")

                # Removed st.rerun() - rely on Streamlit's natural reactivity

            except Exception as e:
                st.error(f"Error loading King Ranch: {e}")

    with col2:
        st.caption("Auto-select King Ranch grids")
    
    try:
        all_grids = load_distinct_grids(session)
    except:
        all_grids = [grid_id]
    
    default_grids = st.session_state.get('s4_grids', [grid_id])
    
    selected_grids = st.multiselect(
        "Select Grids",
        options=all_grids,
        default=default_grids,
        max_selections=20,
        key="s4_grids"
    )
    
    if not selected_grids:
        st.warning("Select at least one grid")
        return
    
    multi_grid_mode = len(selected_grids) > 1
    
    # === ACRES CONFIGURATION ===
    use_custom_acres = st.checkbox("Configure acres per grid", value=False, key="s4_use_acres")
    
    grid_acres = {}
    if use_custom_acres:
        st.subheader(f"Acres for {len(selected_grids)} Grid(s)")
        cols = st.columns(min(4, len(selected_grids)))
        for idx, gid in enumerate(selected_grids):
            with cols[idx % 4]:
                # Check if King Ranch preset exists
                numeric_id = extract_numeric_grid_id(gid)
                default_acres = KING_RANCH_PRESET['acres'].get(numeric_id, total_insured_acres)
                grid_acres[gid] = st.number_input(
                    f"{gid}",
                    min_value=1,
                    value=default_acres,
                    step=10,
                    key=f"s4_acres_{gid}"
                )
    else:
        # Use default acres for all grids
        for gid in selected_grids:
            grid_acres[gid] = total_insured_acres
    
    col1, col2 = st.columns(2)
    start_year = col1.selectbox("Start Year", list(range(1948, 2026)), index=50, key="s4_start")
    end_year = col2.selectbox("End Year", list(range(1948, 2026)), index=77, key="s4_end")
    
    coverage_levels = st.multiselect(
        "Coverage Levels", 
        [0.70, 0.75, 0.80, 0.85, 0.90], 
        default=[0.70, 0.75, 0.80, 0.85, 0.90],
        key="s4_coverage"
    )
    
    objective = st.selectbox(
        "Objective",
        ['cumulative_roi', 'median_roi', 'profitable_pct', 'risk_adj_ret'],
        index=0,
        key="s4_objective"
    )
    
    col1, col2 = st.columns(2)
    
    col1a, col1b = col1.columns(2)
    min_intervals = col1a.number_input(
        "Min Intervals",
        min_value=2,
        max_value=6,
        value=5,
        step=1,
        key="s4_min_intervals"
    )
    max_intervals_to_test = col1b.number_input(
        "Max Intervals",
        min_value=2,
        max_value=6,
        value=6,
        step=1,
        key="s4_max_intervals"
    )
    
    if min_intervals > max_intervals_to_test:
        col1.error("Min cannot exceed Max")
        min_intervals = max_intervals_to_test
    
    search_depth_map = {
        'Fast': 'fast', 
        'Marginal': 'marginal',
        'Incremental': 'incremental',
        'Standard': 'standard', 
        'Thorough': 'thorough', 
        'Maximum': 'maximum'
    }
    search_depth_key = col2.select_slider(
        "Search Depth",
        options=list(search_depth_map.keys()),
        value='Standard',
        key="s4_search_depth"
    )
    search_depth = search_depth_map[search_depth_key]
    
    # Add helpful descriptions for special modes
    if search_depth == 'marginal':
        col2.caption("Shifts intervals by 1-2 months")
    elif search_depth == 'incremental':
        col2.caption("Small fine-tuning: ±1-5% adjustments")

    # === ACRE ALLOCATION STRATEGY (NEW - Two-Stage Optimization) ===
    st.divider()
    st.subheader("Acre Allocation Strategy")

    allocation_mode = st.radio(
        "How should acres be distributed across grids?",
        ["Uniform Acres", "Optimized Acres"],
        help="Uniform: Equal/custom acres per grid. Optimized: Uses correlations for risk-adjusted allocation.",
        key="s4_allocation_mode"
    )

    risk_aversion = 1.0  # Default
    max_turnover = 0.20  # Default
    if allocation_mode == "Optimized Acres":
        st.info("The optimizer calculates which grids generate the highest returns on average and which grids tend to have bad years simultaneously. It then allocates more acres to high-profit grids that don't fail together, reducing the risk of being 'all-in' when drought hits multiple locations at once.")

        opt_col1, opt_col2 = st.columns(2)
        with opt_col1:
            risk_profile = st.select_slider(
                "Risk Profile",
                options=["Aggressive", "Growth", "Balanced", "Conservative", "Defensive"],
                value="Balanced",
                help="Aggressive = chase highest returns, accept volatility. Defensive = prioritize stability and diversification.",
                key="s4_risk_profile"
            )

            # Map to actual lambda values for the optimizer
            risk_aversion_map = {
                "Aggressive": 0.5,
                "Growth": 0.75,
                "Balanced": 1.0,
                "Conservative": 1.5,
                "Defensive": 2.0
            }
            risk_aversion = risk_aversion_map[risk_profile]

        with opt_col2:
            max_turnover = st.slider(
                "Max Turnover",
                min_value=0,
                max_value=100,
                value=20,
                step=5,
                format="%d%%",
                help="How much each grid's allocation can change. 0% = no changes allowed, 100% = full reallocation allowed.",
                key="s4_max_turnover"
            ) / 100.0  # Convert back to decimal for calculations

    st.divider()

    enable_budget = st.checkbox(
        "Set Premium Budget",
        value=False,
        help="Limit total annual premium spending",
        key="s4_enable_budget"
    )

    annual_budget = 50000  # Default
    budget_method = "Optimized Distribution"  # Default
    if enable_budget:
        annual_budget = st.number_input(
            "Maximum Annual Premium Budget ($)",
            min_value=1000,
            value=50000,
            step=1000,
            help="Maximum producer premium (after subsidy) per year",
            key="s4_annual_budget"
        )

        if allocation_mode == "Optimized Acres":
            budget_method = st.radio(
                "Budget Optimization Method:",
                ["Equal Scaling", "Optimized Distribution"],
                index=1,
                help="Equal Scaling: Scale all grids proportionally. Optimized Distribution: Redistribute acres to maximize returns within budget.",
                key="s4_budget_method"
            )
        else:
            st.caption("With Uniform Acres, budget constraint uses Equal Scaling only")
            budget_method = "Equal Scaling"

        # Auto-fill budget option (scale up if under budget)
        s4_allow_scale_up = st.checkbox(
            "Auto-fill budget (scale up if under)",
            value=False,
            help="When enabled, if the optimized allocation costs less than the budget, acres will be scaled up proportionally to utilize the full budget.",
            key="s4_allow_scale_up"
        )
    else:
        s4_allow_scale_up = False

    st.divider()

    if 'tab4_run' not in st.session_state:
        st.session_state.tab4_run = False

    if st.button("Run Optimization", key="s4_run_button"):
        st.session_state.tab4_run = True
        if not coverage_levels:
            st.error("Select at least one coverage level")
            return
            
        try:
            correlation_matrix = None
            if multi_grid_mode:
                with st.spinner("Calculating correlations..."):
                    all_grid_data = []
                    for gid in selected_grids:
                        df = load_all_indices(session, gid)
                        df = df[(df['YEAR'] >= start_year) & (df['YEAR'] <= end_year)]
                        df['GRID_ID'] = gid
                        all_grid_data.append(df)
                    
                    combined_df = pd.concat(all_grid_data, ignore_index=True)
                    
                    pivot_df = combined_df.pivot_table(
                        values='INDEX_VALUE',
                        index=['YEAR', 'INTERVAL_NAME'],
                        columns='GRID_ID'
                    )
                    
                    correlation_matrix = pivot_df.corr()
            
            grid_results = {}
            
            with st.spinner(f"Optimizing {len(selected_grids)} grids..."):
                for gid in selected_grids:
                    try:
                        # Use grid-specific acres
                        acres_for_grid = grid_acres.get(gid, total_insured_acres)
                        
                        results_df, rate_year, strategy_count = run_optimization_s4(
                            session, gid, start_year, end_year, plan_code, 
                            productivity_factor, acres_for_grid, intended_use, 
                            coverage_levels, objective, min_intervals, max_intervals_to_test, search_depth
                        )
                        
                        if not results_df.empty:
                            grid_results[gid] = {
                                'results_df': results_df,
                                'best_strategy': results_df.iloc[0],
                                'strategy_count': strategy_count,
                                'acres': acres_for_grid
                            }
                        else:
                            st.warning(f"Grid {gid}: No valid strategies")
                    except Exception as e:
                        st.error(f"Grid {gid}: {str(e)}")
            
            if not grid_results:
                st.error("No valid strategies found")
                return

            # === STAGE 2: Optimize ACRE distribution (NEW - Two-Stage Optimization) ===
            optimized_grid_acres = grid_acres.copy()
            roi_correlation_matrix = None
            stage2_info = None

            if allocation_mode == "Uniform Acres":
                # Simple equal distribution (EXISTING BEHAVIOR)
                if use_custom_acres:
                    # Use custom acres as specified
                    for gid in selected_grids:
                        optimized_grid_acres[gid] = st.session_state.get(f"s4_acres_{gid}", total_insured_acres / len(selected_grids))
                else:
                    # Use uniform acres
                    acres_per_grid = total_insured_acres / len(selected_grids)
                    for gid in selected_grids:
                        optimized_grid_acres[gid] = acres_per_grid

                # Apply budget constraint if enabled (Equal Scaling only)
                if enable_budget:
                    total_cost, _ = calculate_annual_premium_cost(
                        session, selected_grids, optimized_grid_acres, grid_results,
                        productivity_factor, intended_use, plan_code
                    )

                    if total_cost > annual_budget:
                        optimized_grid_acres, scale_factor = apply_budget_constraint(
                            optimized_grid_acres, total_cost, annual_budget
                        )
                        stage2_info = f"Acres scaled by {scale_factor:.1%} to meet budget"

            else:  # "Optimized Acres" - NEW FEATURE
                with st.spinner("Stage 2: Optimizing acre distribution using correlations..."):
                    # Build historical ROI series for each grid using their best allocation
                    all_grid_roi_data = []

                    for gid in selected_grids:
                        if gid not in grid_results:
                            continue

                        best_alloc = grid_results[gid]['best_strategy']['allocation']
                        best_cov = grid_results[gid]['best_strategy']['coverage_level']

                        # Calculate yearly ROIs for this grid's optimal strategy
                        for year in range(start_year, end_year + 1):
                            roi, indemnity, premium = calculate_yearly_roi_for_grid(
                                session, gid, year, best_alloc, best_cov,
                                productivity_factor, intended_use, plan_code,
                                acres=1  # Normalize to per-acre
                            )

                            all_grid_roi_data.append({
                                'grid': gid,
                                'year': year,
                                'roi': roi,
                                'indemnity': indemnity,
                                'premium': premium
                            })

                    base_data_df = pd.DataFrame(all_grid_roi_data)

                    # Initial acres guess (uniform or custom)
                    initial_acres = {}
                    if use_custom_acres:
                        for gid in selected_grids:
                            initial_acres[gid] = st.session_state.get(f"s4_acres_{gid}", total_insured_acres / len(selected_grids))
                    else:
                        for gid in selected_grids:
                            initial_acres[gid] = total_insured_acres / len(selected_grids)

                    if enable_budget and budget_method == "Optimized Distribution":
                        # Use mean-variance optimizer (two-stage: budget scaling then MVO rebalancing)
                        optimized_grid_acres, roi_correlation_matrix, opt_info = optimize_grid_allocation(
                            base_data_df=base_data_df,
                            grid_results=grid_results,
                            initial_acres_per_grid=initial_acres,
                            annual_budget=annual_budget,
                            session=session,
                            productivity_factor=productivity_factor,
                            intended_use=intended_use,
                            plan_code=plan_code,
                            selected_grids=selected_grids,
                            risk_aversion=risk_aversion,
                            max_turnover=max_turnover,
                            allow_scale_up=s4_allow_scale_up
                        )
                        scale_factor = opt_info.get('budget_scale_factor', 1.0)
                        initial_total = opt_info.get('initial_total_acres', 0)
                        scaled_total = opt_info.get('budget_scaled_acres', 0)
                        if scale_factor != 1.0:
                            stage2_info = f"Budget scaled {initial_total:,.0f} → {scaled_total:,.0f} acres, then MVO rebalanced within ±{max_turnover:.0%}"
                        else:
                            stage2_info = f"Acre distribution optimized within ±{max_turnover:.0%} turnover bounds"

                        # Add Stage 3 info if it was applied
                        if opt_info.get('stage3_applied', False):
                            stage3_scale = opt_info.get('stage3_scale_factor', 1.0)
                            stage2_info += f", then scaled up by {stage3_scale:.1%} to fill budget"

                    elif enable_budget and budget_method == "Equal Scaling":
                        # Calculate cost with initial acres, then scale
                        total_cost, _ = calculate_annual_premium_cost(
                            session, selected_grids, initial_acres, grid_results,
                            productivity_factor, intended_use, plan_code
                        )

                        if total_cost > annual_budget:
                            optimized_grid_acres, scale_factor = apply_budget_constraint(
                                initial_acres, total_cost, annual_budget
                            )
                            stage2_info = f"Acres scaled by {scale_factor:.1%} to meet budget"
                        else:
                            optimized_grid_acres = initial_acres.copy()
                            stage2_info = "Budget constraint not binding"

                    else:
                        # No budget - optimize for pure risk-adjusted return
                        max_acres = sum(initial_acres.values())
                        optimized_grid_acres = optimize_without_budget(
                            base_data_df=base_data_df,
                            grid_results=grid_results,
                            max_total_acres=max_acres,
                            selected_grids=selected_grids,
                            risk_aversion=risk_aversion,
                            max_turnover=max_turnover,
                            initial_acres_per_grid=initial_acres
                        )
                        stage2_info = f"Acre distribution optimized within ±{max_turnover:.0%} turnover bounds"

            # Update grid_results with optimized acres
            for gid in selected_grids:
                if gid in grid_results:
                    grid_results[gid]['acres'] = optimized_grid_acres.get(gid, 0)

            st.session_state.tab4_results = {
                "grid_results": grid_results,
                "selected_grids": selected_grids,
                "start_year": start_year,
                "end_year": end_year,
                "objective": objective,
                "coverage_levels": coverage_levels,
                "multi_grid_mode": multi_grid_mode,
                "correlation_matrix": correlation_matrix,
                "search_depth": search_depth,
                "grid_acres": optimized_grid_acres,
                "allocation_mode": allocation_mode,
                "budget_enabled": enable_budget,
                "annual_budget": annual_budget if enable_budget else None,
                "budget_method": budget_method if enable_budget else None,
                "risk_aversion": risk_aversion,
                "stage2_info": stage2_info,
                "roi_correlation_matrix": roi_correlation_matrix
            }
            st.session_state.tab2_results = None
            st.session_state.tab5_results = None

        except Exception as e:
            st.error(f"Error: {e}")
            st.exception(e)
            st.session_state.tab4_results = None

    if 'tab4_results' in st.session_state and st.session_state.tab4_results:
        try:
            r = st.session_state.tab4_results
            
            objective_display = {
                'cumulative_roi': 'Cumulative ROI',
                'median_roi': 'Median ROI',
                'profitable_pct': 'Win Rate',
                'risk_adj_ret': 'Risk-Adjusted Return'
            }
            
            st.header(f"Results ({r['start_year']}-{r['end_year']})")
            st.caption(f"Productivity: {productivity_factor:.0%} | Use: {intended_use}")
            
            if r.get('search_depth') == 'marginal':
                st.info("🔧 Marginal Mode: Testing variations of existing allocations")
            elif r.get('search_depth') == 'incremental':
                st.info("📊 Incremental Mode: Testing percentage adjustments within current allocations")

            # === ACRE ALLOCATION MODE DISPLAY ===
            if r.get('allocation_mode') == "Optimized Acres":
                st.success(f"Optimized Acres Mode | Risk Tolerance: {r.get('risk_aversion', 1.0):.1f}")
                if r.get('stage2_info'):
                    st.caption(r.get('stage2_info'))
            elif r.get('budget_enabled') and r.get('stage2_info'):
                st.info(f"Budget Constraint Active: {r.get('stage2_info')}")

            # === ACRE DISTRIBUTION ANALYSIS (for Optimized Acres mode) ===
            if r.get('allocation_mode') == "Optimized Acres" and r['multi_grid_mode']:
                st.divider()
                st.subheader("Acre Distribution Analysis")

                # Show how acres were allocated
                acre_data = []
                total_acres = sum(r['grid_acres'].values())
                for gid in r['selected_grids']:
                    if gid in r['grid_results']:
                        grid_acres_val = r['grid_acres'].get(gid, 0)
                        acre_data.append({
                            'Grid': str(gid),
                            'Optimized Acres': grid_acres_val,
                            'Percent of Total': (grid_acres_val / total_acres * 100) if total_acres > 0 else 0
                        })

                acre_df = pd.DataFrame(acre_data)

                # Add totals row
                total_row = pd.DataFrame([{
                    'Grid': 'TOTAL',
                    'Optimized Acres': total_acres,
                    'Percent of Total': 100.0
                }])
                acre_df = pd.concat([acre_df, total_row], ignore_index=True)

                # Format and display
                st.dataframe(
                    acre_df.style.format({
                        'Optimized Acres': '{:,.0f}',
                        'Percent of Total': '{:.1f}%'
                    }),
                    use_container_width=True,
                    hide_index=True
                )

                # Download button
                st.download_button(
                    label="📥 Download Acre Distribution CSV",
                    data=acre_df.to_csv(index=False),
                    file_name=f"acre_distribution_{r['start_year']}-{r['end_year']}.csv",
                    mime="text/csv",
                    key="acre_dist_csv"
                )

                # Budget utilization display
                if r.get('budget_enabled'):
                    actual_cost, grid_costs = calculate_annual_premium_cost(
                        session, r['selected_grids'], r['grid_acres'], r['grid_results'],
                        productivity_factor, intended_use, plan_code
                    )

                    col1, col2, col3 = st.columns(3)
                    col1.metric("Annual Premium Cost", f"${actual_cost:,.0f}")
                    col2.metric("Budget", f"${r['annual_budget']:,.0f}")
                    col3.metric("Budget Utilization", f"{(actual_cost / r['annual_budget'] * 100):.1f}%")

                # ROI Correlation Matrix (if available)
                if r.get('roi_correlation_matrix') is not None and not r['roi_correlation_matrix'].empty:
                    st.divider()
                    st.subheader("ROI Correlation Matrix")
                    st.caption("Correlations based on historical ROI performance (lower = better diversification)")

                    # Format correlation matrix for display
                    roi_corr = r['roi_correlation_matrix']
                    corr_csv = roi_corr.to_csv()
                    st.download_button(
                        label="📥 Download ROI Correlations CSV",
                        data=corr_csv,
                        file_name=f"roi_correlations_{r['start_year']}-{r['end_year']}.csv",
                        mime="text/csv",
                        key="roi_corr_csv"
                    )

                    # Display as text table for alignment
                    roi_header = "Grid                  " + "".join([f"{str(gid)[:20]:>21}" for gid in roi_corr.columns])
                    st.text(roi_header)
                    st.text("─" * len(roi_header))

                    for idx, row in roi_corr.iterrows():
                        line = f"{str(idx):<20} "
                        for val in row:
                            line += f"{val:>21.3f}"
                        st.text(line)

                    st.text("═" * len(roi_header))

                st.divider()

            if r['multi_grid_mode'] and r['correlation_matrix'] is not None:
                st.subheader("Grid Correlations")
                
                # CSV Download Button
                corr_csv = r['correlation_matrix'].to_csv()
                st.download_button(
                    label="📥 Download CSV",
                    data=corr_csv,
                    file_name=f"grid_correlations_{r['start_year']}-{r['end_year']}.csv",
                    mime="text/csv",
                )
                
                # Fixed: Widen grid column to 20 chars (matching financial summary)
                header = "Grid                  " + "".join([f"{str(gid)[:20]:>21}" for gid in r['correlation_matrix'].columns])
                st.text(header)
                st.text("─" * len(header))
                
                for idx, row in r['correlation_matrix'].iterrows():
                    line = f"{str(idx):<20} "
                    for val in row:
                        line += f"{val:>21.3f}"
                    st.text(line)
                
                st.text("═" * len(header))
                st.caption("Lower correlation = better diversification")
                
                st.divider()
            
            for gid in r['selected_grids']:
                if gid not in r['grid_results']:
                    st.warning(f"No results for {gid}")
                    continue
                
                grid_data = r['grid_results'][gid]
                results_df = grid_data['results_df']
                best = grid_data['best_strategy']
                
                st.subheader(f"{gid} - Top Strategy")
                
                if r.get('search_depth') == 'marginal':
                    st.caption(f"Tested {grid_data['strategy_count']} marginal variations")
                elif r.get('search_depth') == 'incremental':
                    st.caption(f"Tested {grid_data['strategy_count']} incremental percentage variations")
                else:
                    st.caption(f"Tested {grid_data['strategy_count']} strategies")
                
                c1, c2, c3, c4, c5 = st.columns(5)
                c1.metric("Coverage", f"{best['coverage_level']:.0%}")
                c2.metric("Median ROI", f"{best['median_roi']:.2%}")
                c3.metric("Cumulative ROI", f"{best['cumulative_roi']:.2%}")
                c4.metric("Win Rate", f"{best['profitable_pct']:.1%}")
                c5.metric("Risk-Adj", f"{best['risk_adj_ret']:.2f}")

                # Show best allocation
                alloc_str = ", ".join([f"{k}: {v*100:.0f}%" for k, v in sorted(best['allocation'].items(), key=lambda x: x[1], reverse=True) if v > 0])
                st.caption(f"Best Allocation: {alloc_str}")

                st.divider()
            
            if r['multi_grid_mode']:
                st.subheader("Portfolio Coverage")
                
                # Build coverage data as DataFrame for perfect alignment
                coverage_data = []
                total_coverage = {interval: 0 for interval in INTERVAL_ORDER_11}
                valid_grids_count = len([gid for gid in r['selected_grids'] if gid in r['grid_results']])
                
                for gid in r['selected_grids']:
                    if gid in r['grid_results']:
                        best = r['grid_results'][gid]['best_strategy']
                        row = {'Grid': str(gid)}
                        row_sum = 0
                        
                        for interval in INTERVAL_ORDER_11:
                            pct = best['allocation'].get(interval, 0) * 100
                            row_sum += pct
                            total_coverage[interval] += pct
                            row[interval] = f"{pct:.0f}%" if pct > 0 else "--"
                        
                        row['Row Sum'] = f"{row_sum:.0f}%"
                        coverage_data.append(row)
                
                # Add average row
                avg_row = {'Grid': 'AVERAGE'}
                avg_row_sum = 0
                for interval in INTERVAL_ORDER_11:
                    pct = total_coverage[interval] / valid_grids_count if valid_grids_count > 0 else 0
                    avg_row_sum += pct
                    avg_row[interval] = f"{pct:.0f}%" if pct > 0.5 else "--"
                avg_row['Row Sum'] = f"{avg_row_sum:.0f}%"
                coverage_data.append(avg_row)
                
                coverage_df = pd.DataFrame(coverage_data)
                
                # CSV Download Button
                csv_export = coverage_df.to_csv(index=False)
                st.download_button(
                    label="📥 Download CSV",
                    data=csv_export,
                    file_name=f"portfolio_coverage_{r['start_year']}-{r['end_year']}.csv",
                    mime="text/csv",
                )
                
                st.dataframe(coverage_df, use_container_width=True, hide_index=True)
                
                st.divider()
                st.subheader("Financial Summary")
                
                with st.spinner("Calculating..."):
                    num_years = r['end_year'] - r['start_year'] + 1
                    
                    if num_years <= 0:
                        st.error("Invalid year range")
                        return
                    
                    # Annual - Added Risk-Adj column
                    st.text(f"ANNUAL AVERAGE ({r['start_year']}-{r['end_year']}: {num_years} years)")
                    st.text("Grid                     Coverage   Avg Premium/Yr    Avg Indemnity/Yr   Avg Net Return/Yr       Risk-Adj")
                    st.text("─" * 120)
                    
                    portfolio_total_premium = 0
                    portfolio_total_indemnity = 0
                    portfolio_year_rois_all = []
                    
                    for gid in r['selected_grids']:
                        if gid not in r['grid_results']:
                            continue
                            
                        best = r['grid_results'][gid]['best_strategy']
                        coverage_level = best['coverage_level']
                        allocation = best['allocation']
                        
                        try:
                            subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                            county_base_value = load_county_base_value(session, gid)
                            current_rate_year = get_current_rate_year(session)
                            premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
                            dollar_amount_of_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                            
                            # Use grid-specific acres
                            acres_for_grid = r['grid_acres'].get(gid, total_insured_acres)
                            total_policy_protection = dollar_amount_of_protection * acres_for_grid
                            all_indices_df = load_all_indices(session, gid)
                            all_indices_df = all_indices_df[(all_indices_df['YEAR'] >= r['start_year']) & (all_indices_df['YEAR'] <= r['end_year'])]
                            
                            grid_total_premium = 0
                            grid_total_indemnity = 0
                            grid_year_rois = []
                            
                            for year in range(r['start_year'], r['end_year'] + 1):
                                actuals_df = all_indices_df[all_indices_df['YEAR'] == year].set_index('INTERVAL_NAME')
                                if actuals_df.empty:
                                    continue
                                
                                roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                                roi_df['Percent of Value'] = roi_df.index.map(allocation)
                                roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                                roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                                roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                                roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                                roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce').fillna(0)
                                roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                                roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                                roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                                trigger_level = coverage_level * 100
                                shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                                roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)

                                year_premium = roi_df['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                                year_indemnity = roi_df['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()

                                grid_total_premium += year_premium
                                grid_total_indemnity += year_indemnity

                                year_roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                                grid_year_rois.append(year_roi)

                            grid_annual_premium = grid_total_premium / num_years
                            grid_annual_indemnity = grid_total_indemnity / num_years
                            grid_annual_net = grid_annual_indemnity - grid_annual_premium
                            
                            # Calculate risk-adjusted return
                            grid_cumulative_roi = (grid_total_indemnity - grid_total_premium) / grid_total_premium if grid_total_premium > 0 else 0
                            grid_std_dev = np.std(grid_year_rois) if len(grid_year_rois) > 0 else 0
                            grid_risk_adj = grid_cumulative_roi / grid_std_dev if grid_std_dev > 0 else 0
                            
                            portfolio_total_premium += grid_total_premium
                            portfolio_total_indemnity += grid_total_indemnity
                            portfolio_year_rois_all.extend(grid_year_rois)
                            
                            line = f"{str(gid):<25} {coverage_level:>7.0%} {grid_annual_premium:>18,.0f} {grid_annual_indemnity:>19,.0f} {grid_annual_net:>19,.0f} {grid_risk_adj:>18.2f}"
                            st.text(line)
                            
                        except Exception as e:
                            st.text(f"{str(gid):<25} ERROR: {str(e)[:60]}")
                    
                    st.text("═" * 120)
                    portfolio_annual_premium = portfolio_total_premium / num_years if num_years > 0 else 0
                    portfolio_annual_indemnity = portfolio_total_indemnity / num_years if num_years > 0 else 0
                    portfolio_annual_net = portfolio_annual_indemnity - portfolio_annual_premium
                    portfolio_annual_roi = portfolio_annual_net / portfolio_annual_premium if portfolio_annual_premium > 0 else 0
                    
                    # Portfolio risk-adjusted return
                    portfolio_cumulative_roi = (portfolio_total_indemnity - portfolio_total_premium) / portfolio_total_premium if portfolio_total_premium > 0 else 0
                    portfolio_std_dev = np.std(portfolio_year_rois_all) if len(portfolio_year_rois_all) > 0 else 0
                    portfolio_risk_adj = portfolio_cumulative_roi / portfolio_std_dev if portfolio_std_dev > 0 else 0
                    
                    line = f"{'AVERAGE':<25} {'--':>7} {portfolio_annual_premium:>18,.0f} {portfolio_annual_indemnity:>19,.0f} {portfolio_annual_net:>19,.0f} {portfolio_risk_adj:>18.2f}"
                    st.text(line)
                    st.text(f"Average Annual ROI: {portfolio_annual_roi:.2%}")
                    
                    # Cumulative - Added Risk-Adj column
                    st.text(f"\nCUMULATIVE TOTAL ({r['start_year']}-{r['end_year']}: {num_years} years)")
                    st.text("Grid                     Coverage     Total Premium      Total Indemnity       Total Net Return       Risk-Adj")
                    st.text("─" * 120)
                    
                    for gid in r['selected_grids']:
                        if gid in r['grid_results']:
                            best = r['grid_results'][gid]['best_strategy']
                            coverage_level = best['coverage_level']
                            allocation = best['allocation']
                            
                            try:
                                subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                                county_base_value = load_county_base_value(session, gid)
                                current_rate_year = get_current_rate_year(session)
                                premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
                                dollar_amount_of_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                                
                                # Use grid-specific acres
                                acres_for_grid = r['grid_acres'].get(gid, total_insured_acres)
                                total_policy_protection = dollar_amount_of_protection * acres_for_grid
                                all_indices_df = load_all_indices(session, gid)
                                all_indices_df = all_indices_df[(all_indices_df['YEAR'] >= r['start_year']) & (all_indices_df['YEAR'] <= r['end_year'])]
                                
                                grid_total_premium = 0
                                grid_total_indemnity = 0
                                grid_year_rois = []
                                
                                for year in range(r['start_year'], r['end_year'] + 1):
                                    actuals_df = all_indices_df[all_indices_df['YEAR'] == year].set_index('INTERVAL_NAME')
                                    if actuals_df.empty:
                                        continue
                                    
                                    roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                                    roi_df['Percent of Value'] = roi_df.index.map(allocation)
                                    roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                                    roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                                    roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                                    roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                                    roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce').fillna(0)
                                    roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                                    roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                                    roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                                    trigger_level = coverage_level * 100
                                    shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                                    roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)

                                    year_premium = roi_df['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                                    year_indemnity = roi_df['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()

                                    grid_total_premium += year_premium
                                    grid_total_indemnity += year_indemnity

                                    year_roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                                    grid_year_rois.append(year_roi)

                                grid_cumulative_net = grid_total_indemnity - grid_total_premium
                                
                                # Calculate risk-adjusted return
                                grid_cumulative_roi = grid_cumulative_net / grid_total_premium if grid_total_premium > 0 else 0
                                grid_std_dev = np.std(grid_year_rois) if len(grid_year_rois) > 0 else 0
                                grid_risk_adj = grid_cumulative_roi / grid_std_dev if grid_std_dev > 0 else 0
                                
                                line = f"{str(gid):<25} {coverage_level:>7.0%} {grid_total_premium:>18,.0f} {grid_total_indemnity:>19,.0f} {grid_cumulative_net:>20,.0f} {grid_risk_adj:>18.2f}"
                                st.text(line)
                                
                            except Exception as e:
                                st.text(f"{str(gid):<25} ERROR: {str(e)[:60]}")
                    
                    st.text("═" * 120)
                    portfolio_cumulative_net = portfolio_total_indemnity - portfolio_total_premium
                    portfolio_cumulative_roi = portfolio_cumulative_net / portfolio_total_premium if portfolio_total_premium > 0 else 0
                    
                    line = f"{'TOTAL':<25} {'--':>7} {portfolio_total_premium:>18,.0f} {portfolio_total_indemnity:>19,.0f} {portfolio_cumulative_net:>20,.0f} {portfolio_risk_adj:>18.2f}"
                    st.text(line)
                    st.text(f"Cumulative ROI: {portfolio_cumulative_roi:.2%}")
                    st.text("═" * 120)
                
                # === KING RANCH COMPARISON OUTPUT ===
                if st.session_state.get('s4_king_ranch_comparison_mode', False) and use_custom_acres:
                    st.divider()
                    st.subheader("King Ranch Comparison: Current vs. Suggested")
                    st.caption(f"Comparing current allocations against optimizer-recommended allocations | Productivity: {productivity_factor:.0%} | Use: {intended_use}")
                    
                    # Build comparison tables
                    current_data = []
                    suggested_data = []
                    change_data = []
                    
                    current_total_premium = 0
                    current_total_indemnity = 0
                    current_year_rois = []
                    
                    suggested_total_premium = 0
                    suggested_total_indemnity = 0
                    suggested_year_rois = []
                    
                    for gid in r['selected_grids']:
                        if gid not in r['grid_results']:
                            continue
                        
                        numeric_id = extract_numeric_grid_id(gid)
                        
                        # Current allocation from King Ranch preset
                        if numeric_id in KING_RANCH_PRESET['allocations']:
                            current_alloc_pct = KING_RANCH_PRESET['allocations'][numeric_id]
                            current_alloc = {k: v/100.0 for k, v in current_alloc_pct.items()}
                        else:
                            continue
                        
                        # Suggested allocation from optimizer
                        suggested_alloc = r['grid_results'][gid]['best_strategy']['allocation']
                        coverage_level = r['grid_results'][gid]['best_strategy']['coverage_level']
                        
                        # Build rows for tables
                        current_row = {'Grid': str(gid)}
                        suggested_row = {'Grid': str(gid)}
                        change_row = {'Grid': str(gid)}
                        
                        for interval in INTERVAL_ORDER_11:
                            curr_pct = current_alloc.get(interval, 0) * 100
                            sugg_pct = suggested_alloc.get(interval, 0) * 100
                            change_pct = sugg_pct - curr_pct
                            
                            current_row[interval] = f"{curr_pct:.0f}%" if curr_pct > 0 else "--"
                            suggested_row[interval] = f"{sugg_pct:.0f}%" if sugg_pct > 0 else "--"
                            
                            if abs(change_pct) < 0.5:
                                change_row[interval] = "--"
                            elif change_pct > 0:
                                change_row[interval] = f"+{change_pct:.0f}%"
                            else:
                                change_row[interval] = f"{change_pct:.0f}%"
                        
                        current_row['Total'] = "100%"
                        suggested_row['Total'] = "100%"
                        change_row['Total'] = "0%"
                        
                        current_data.append(current_row)
                        suggested_data.append(suggested_row)
                        change_data.append(change_row)
                        
                        # Calculate performance for both allocations
                        try:
                            subsidy_percent = load_subsidies(session, plan_code, [coverage_level])[coverage_level]
                            county_base_value = load_county_base_value(session, gid)
                            current_rate_year = get_current_rate_year(session)
                            premium_rates_df = load_premium_rates(session, gid, intended_use, [coverage_level], current_rate_year)[coverage_level]
                            dollar_amount_of_protection = calculate_protection(county_base_value, coverage_level, productivity_factor)
                            
                            acres_for_grid = r['grid_acres'].get(gid, total_insured_acres)
                            total_policy_protection = dollar_amount_of_protection * acres_for_grid
                            all_indices_df = load_all_indices(session, gid)
                            all_indices_df = all_indices_df[(all_indices_df['YEAR'] >= r['start_year']) & (all_indices_df['YEAR'] <= r['end_year'])]
                            
                            # Calculate for CURRENT allocation
                            curr_grid_premium = 0
                            curr_grid_indemnity = 0
                            curr_grid_rois = []
                            
                            for year in range(r['start_year'], r['end_year'] + 1):
                                actuals_df = all_indices_df[all_indices_df['YEAR'] == year].set_index('INTERVAL_NAME')
                                if actuals_df.empty:
                                    continue
                                
                                roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                                roi_df['Percent of Value'] = roi_df.index.map(current_alloc)
                                roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                                roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                                roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                                roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                                roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce').fillna(0)
                                roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                                roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                                roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                                trigger_level = coverage_level * 100
                                shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                                roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)

                                year_premium = roi_df['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                                year_indemnity = roi_df['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()

                                curr_grid_premium += year_premium
                                curr_grid_indemnity += year_indemnity

                                year_roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                                curr_grid_rois.append(year_roi)

                            current_total_premium += curr_grid_premium
                            current_total_indemnity += curr_grid_indemnity
                            current_year_rois.extend(curr_grid_rois)

                            # Calculate for SUGGESTED allocation (already calculated in grid_results)
                            sugg_grid_premium = 0
                            sugg_grid_indemnity = 0
                            sugg_grid_rois = []
                            
                            for year in range(r['start_year'], r['end_year'] + 1):
                                actuals_df = all_indices_df[all_indices_df['YEAR'] == year].set_index('INTERVAL_NAME')
                                if actuals_df.empty:
                                    continue
                                
                                roi_df = pd.DataFrame(index=INTERVAL_ORDER_11)
                                roi_df['Percent of Value'] = roi_df.index.map(suggested_alloc)
                                roi_df['Policy Protection Per Unit'] = (total_policy_protection * roi_df['Percent of Value']).apply(lambda x: round_half_up(x, 0))
                                roi_df = roi_df.join(pd.Series(premium_rates_df, name='PREMIUM_RATE'))
                                roi_df = roi_df.join(actuals_df.rename(columns={'INDEX_VALUE': 'Actual Index Value'}))
                                roi_df['PREMIUM_RATE'] = pd.to_numeric(roi_df['PREMIUM_RATE'], errors='coerce').fillna(0)
                                roi_df['Actual Index Value'] = pd.to_numeric(roi_df['Actual Index Value'], errors='coerce').fillna(0)
                                roi_df['Total Premium'] = (roi_df['Policy Protection Per Unit'] * roi_df['PREMIUM_RATE']).apply(lambda x: round_half_up(x, 0))
                                roi_df['Premium Subsidy'] = (roi_df['Total Premium'] * subsidy_percent).apply(lambda x: round_half_up(x, 0))
                                roi_df['Producer Premium'] = roi_df['Total Premium'] - roi_df['Premium Subsidy']
                                trigger_level = coverage_level * 100
                                shortfall_pct = (trigger_level - roi_df['Actual Index Value']) / trigger_level
                                roi_df['Estimated Indemnity'] = (shortfall_pct * roi_df['Policy Protection Per Unit']).clip(lower=0).apply(lambda x: round_half_up(x, 0) if abs(x) >= 0.01 else 0.0)

                                year_premium = roi_df['Producer Premium'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()
                                year_indemnity = roi_df['Estimated Indemnity'].apply(lambda x: round_half_up(x, 0) if pd.notna(x) else 0).sum()

                                sugg_grid_premium += year_premium
                                sugg_grid_indemnity += year_indemnity

                                year_roi = (year_indemnity - year_premium) / year_premium if year_premium > 0 else 0
                                sugg_grid_rois.append(year_roi)

                            suggested_total_premium += sugg_grid_premium
                            suggested_total_indemnity += sugg_grid_indemnity
                            suggested_year_rois.extend(sugg_grid_rois)
                            
                        except Exception as e:
                            st.error(f"Error calculating performance for {gid}: {e}")
                    
                    # Display three tables side by side
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown("**Current Allocation**")
                        current_df = pd.DataFrame(current_data)
                        st.dataframe(current_df, use_container_width=True, hide_index=True)
                        st.download_button(
                            label="📥 CSV",
                            data=current_df.to_csv(index=False),
                            file_name=f"kr_current_{r['start_year']}-{r['end_year']}.csv",
                            mime="text/csv",
                            key="kr_current_csv"
                        )
                    
                    with col2:
                        st.markdown("**Suggested Allocation**")
                        suggested_df = pd.DataFrame(suggested_data)
                        st.dataframe(suggested_df, use_container_width=True, hide_index=True)
                        st.download_button(
                            label="📥 CSV",
                            data=suggested_df.to_csv(index=False),
                            file_name=f"kr_suggested_{r['start_year']}-{r['end_year']}.csv",
                            mime="text/csv",
                            key="kr_suggested_csv"
                        )
                    
                    with col3:
                        st.markdown("**Change**")
                        change_df = pd.DataFrame(change_data)
                        st.dataframe(change_df, use_container_width=True, hide_index=True)
                        st.download_button(
                            label="📥 CSV",
                            data=change_df.to_csv(index=False),
                            file_name=f"kr_change_{r['start_year']}-{r['end_year']}.csv",
                            mime="text/csv",
                            key="kr_change_csv"
                        )
                    
                    # Calculate summary metrics
                    st.divider()
                    st.subheader("Performance Comparison")
                    
                    current_roi = (current_total_indemnity - current_total_premium) / current_total_premium if current_total_premium > 0 else 0
                    suggested_roi = (suggested_total_indemnity - suggested_total_premium) / suggested_total_premium if suggested_total_premium > 0 else 0
                    roi_improvement = suggested_roi - current_roi
                    
                    current_std = np.std(current_year_rois) if len(current_year_rois) > 0 else 0
                    suggested_std = np.std(suggested_year_rois) if len(suggested_year_rois) > 0 else 0
                    
                    current_risk_adj = current_roi / current_std if current_std > 0 else 0
                    suggested_risk_adj = suggested_roi / suggested_std if suggested_std > 0 else 0
                    risk_adj_improvement = suggested_risk_adj - current_risk_adj
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown("**Current Portfolio**")
                        st.metric("Cumulative ROI", f"{current_roi:.2%}")
                        st.metric("Risk-Adjusted Return", f"{current_risk_adj:.2f}")
                    
                    with col2:
                        st.markdown("**Suggested Portfolio**")
                        st.metric("Cumulative ROI", f"{suggested_roi:.2%}", delta=f"{roi_improvement:.2%}")
                        st.metric("Risk-Adjusted Return", f"{suggested_risk_adj:.2f}", delta=f"{risk_adj_improvement:.2f}")
                    
                    with col3:
                        st.markdown("**Improvement**")
                        st.metric("ROI Difference", f"{roi_improvement:.2%}")
                        st.metric("Risk-Adj Difference", f"{risk_adj_improvement:.2f}")
            else:
                st.divider()
                st.success("Recommended Strategy")
                gid = r['selected_grids'][0]
                best = r['grid_results'][gid]['best_strategy']
                
                c1, c2 = st.columns(2)
                c1.metric("Coverage", f"{best['coverage_level']:.0%}")
                c2.metric(f"{objective_display.get(r['objective'])}", f"{best[r['objective']]:.2%}" if r['objective'] != 'risk_adj_ret' else f"{best[r['objective']]:.2f}")
                
                st.text("Optimal Allocation")
                st.text("─" * 40)
                st.text(f"{'Interval':<15} {'Allocation':<15}")
                st.text("─" * 40)
                for interval, pct in sorted(best['allocation'].items(), key=lambda x: x[1], reverse=True):
                    if pct > 0:
                        st.text(f"{interval:<15} {pct*100:>6.0f}%")
                st.text("═" * 40)
        
        except Exception as e:
            st.error(f"Error: {e}")
            st.exception(e)
            st.session_state.tab4_results = None
    elif not st.session_state.tab4_run:
        st.info("Select parameters and click 'Run Optimization'")

# =============================================================================
# === 8. MAIN APP LAYOUT ===
# =============================================================================

def main():
    st.title("PRF Backtesting Tool")
    session = get_active_session()

    try:
        valid_grids = load_distinct_grids(session)
    except Exception as e:
        st.sidebar.error("Fatal Error: Could not load Grid ID list")
        st.error(f"Could not load Grid ID list: {e}")
        st.stop()

    st.sidebar.header("Common Parameters")
    
    if 'grid_id' not in st.session_state:
        st.session_state.grid_id = valid_grids[0] if valid_grids else "7928 (Nueces - TX)" 
    if 'productivity_factor' not in st.session_state:
        st.session_state.productivity_factor = 1.0 
    if 'total_insured_acres' not in st.session_state:
        st.session_state.total_insured_acres = 1000
    if 'intended_use' not in st.session_state:
        st.session_state.intended_use = 'Grazing'
    if 'insurance_plan_code' not in st.session_state:
        st.session_state.insurance_plan_code = 13
    
    if 'tab2_run' not in st.session_state:
        st.session_state.tab2_run = False
    if 'tab4_run' not in st.session_state:
        st.session_state.tab4_run = False
    if 'tab5_run' not in st.session_state:
        st.session_state.tab5_run = False

    if 'tab2_results' not in st.session_state:
        st.session_state.tab2_results = None
    if 'tab4_results' not in st.session_state:
        st.session_state.tab4_results = None
    if 'tab5_results' not in st.session_state:
        st.session_state.tab5_results = None
    
    try:
        default_grid_index = valid_grids.index(st.session_state.grid_id)
    except (ValueError, AttributeError):
        default_grid_index = 0
    grid_id = st.sidebar.selectbox(
        "Grid ID", options=valid_grids, index=default_grid_index,
        key="sidebar_grid_id"
    )
    
    prod_options = list(range(60, 151))
    prod_options_formatted = [f"{x}%" for x in prod_options]
    try:
        current_prod_index = prod_options.index(int(st.session_state.productivity_factor * 100))
    except ValueError:
        current_prod_index = 40
    selected_prod_str = st.sidebar.selectbox(
        "Productivity Factor", options=prod_options_formatted, index=current_prod_index,
        key="sidebar_prod_factor"
    )
    productivity_factor = int(selected_prod_str.replace('%', '')) / 100.0
    
    total_insured_acres = st.sidebar.number_input(
        "Total Insured Acres", value=st.session_state.total_insured_acres, step=10,
        key="sidebar_acres"
    )
    intended_use = st.sidebar.selectbox(
        "Intended Use", ['Grazing', 'Haying'], 
        index=0 if st.session_state.intended_use == 'Grazing' else 1,
        key="sidebar_use"
    )
    
    plan_code = st.sidebar.number_input(
        "Insurance Plan Code", 
        value=st.session_state.insurance_plan_code, 
        disabled=True
    )
    
    st.session_state.grid_id = grid_id
    st.session_state.productivity_factor = productivity_factor
    st.session_state.total_insured_acres = total_insured_acres
    st.session_state.intended_use = intended_use
    
    # King Ranch Comparison Mode notification
    if st.session_state.get('s4_king_ranch_comparison_mode', False):
        st.sidebar.info("📊 King Ranch Comparison Mode Active\n\nGo to Optimizer tab → Enable 'Configure acres per grid' → Run to see Current vs. Suggested analysis")

    with st.sidebar.expander("📊 Z-Score Translation Key"):
        st.markdown("""
        **Historical Context (Year Avg Z-Score):**
        - **Dry**: Z < -0.25 (~40th percentile or drier)
        - **Normal**: -0.25 ≤ Z ≤ 0.25 (~40th to ~60th percentile)
        - **Wet**: Z > 0.25 (~60th percentile or wetter)

        **Expected Trajectory (SOY 11P vs EOY 5P):**
        - **Get Drier**: Change < -0.05
        - **Stay Stable**: -0.05 <= Change <= 0.05
        - **Get Wetter**: Change > 0.05

        *Trajectory Change = Nov-Dec 5P Z minus Jan-Feb 11P Z*

        *11P = 11-period rolling avg (stable baseline)*
        *5P = 5-period rolling avg (recent trend)*
        *Positive Change = year trending wetter*
        """)

    with st.sidebar.expander("🏆 Strategy Key"):
        st.markdown("""
        **Champion**
        Your current portfolio allocation

        **Challenger 1**
        Your portfolio with optimized interval timing

        **Challenger 2 (Naive)**
        Weather-focused portfolio - each grid optimized independently for analog years

        **Challenger 3 (MVO)**
        Weather-optimized portfolio - considers cross-grid correlations during analog years, redistributes acres via Mean-Variance Optimization
        """)

    with st.sidebar.expander("⚙️ Budget & Turnover Logic"):
        st.markdown("""
        **Two-Stage Optimization Process:**

        When both Budget and Max Turnover are set, the optimizer works in two stages:

        **Stage 1: Budget Scaling**
        - All grids scale proportionally to meet the premium budget
        - Example: If budget requires 50% reduction, every grid drops 50%
        - This preserves your original portfolio distribution

        **Stage 2: MVO Rebalancing**
        - Mean-Variance Optimization adjusts the distribution
        - Max Turnover constrains how much MVO can shift each grid
        - Turnover bounds apply to the *scaled* baseline, not original acres

        **Example with 20% Turnover:**
        ```
        Original:     Grid A = 80,000 acres
        After Budget: Grid A = 40,000 acres (50% scale)
        MVO Range:    32,000 - 48,000 (±20% of scaled)
        ```

        *This ensures budget compliance while preserving portfolio intent.*
        """)

    st.sidebar.divider()
    st.sidebar.caption("*2025 Rates are used for this application")
    st.sidebar.caption("*Common Parameters are secondary to parameters on each tab")

    st.sidebar.divider()
    if st.sidebar.button("🔄 Clear Champion vs Challenger", key="clear_champ_chall"):
        # Clear all Champion vs Challenger session state
        keys_to_clear = [
            'champion_results',
            'challenger_results',
            'weather_challenger_results',
            'weather_challenger_3_results',
            'stress_test_results',
            'ps_analog_years',
            'ps_weather_config',
            'ps_enable_weather',
            'ps_kr_load_requested',
            'ps_acres_expander_opened',
            'ps_alloc_expander_opened'
        ]

        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]

        # Also clear any ps_ prefixed keys for grid allocations/acres
        keys_to_remove = [k for k in st.session_state.keys() if k.startswith('ps_')]
        for key in keys_to_remove:
            del st.session_state[key]

        st.sidebar.success("Champion vs Challenger cleared!")
        st.rerun()

    if st.sidebar.button("🏠 Load King Ranch (Champion)", key="sidebar_load_kr_champ"):
        st.session_state.sidebar_kr_champion_requested = True
        st.rerun()

    if st.sidebar.button("📈 Load King Ranch Incrementals", key="sidebar_load_kr_incr"):
        st.session_state.sidebar_kr_incrementals_requested = True
        st.rerun()

    tab1, tab2, tab3 = st.tabs([
        "Decision Support (Audit)",
        "Portfolio Backtest (Audit)",
        "Champion vs Challenger(s)"
    ])

    with tab1:
        render_tab2(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code)

    with tab2:
        render_tab5(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code)

    with tab3:
        render_portfolio_strategy_tab(session, grid_id, intended_use, productivity_factor, total_insured_acres, plan_code)

if __name__ == "__main__":
    main()